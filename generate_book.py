#!/usr/bin/env python3
"""
Generate "The Oldest Sell: How Alcohol Mastered Every Medium, from Clay Tablet
to the Feed" by Jacob Edward Thomas — a popular-history book adapted from the
historical chapters (Part A) of the author's doctoral dissertation.

The book is generated as a Word document (.docx). Historical advertisements and
photographs are extracted at build time directly from the source dissertation
PDF (THOMAS-PRIMARY-2025.pdf) and embedded with captions, so the deliverable is
fully reproducible from the repository.

Dependencies:  pip install python-docx PyMuPDF
"""

import os
import tempfile

import fitz  # PyMuPDF — used to pull the historical figures out of the source PDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PDF_PATH = os.path.join(HERE, "THOMAS-PRIMARY-2025.pdf")
IMG_CACHE = os.path.join(tempfile.gettempdir(), "oldest_sell_figures")

ACCENT = (0xBF, 0x57, 0x00)   # UT-Austin burnt orange, used sparingly

# ---------------------------------------------------------------------------
# LOW-LEVEL TYPOGRAPHY HELPERS
# ---------------------------------------------------------------------------

def set_run_font(run, name="Garamond", size=12, bold=False, italic=False,
                 color=None, small_caps=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if small_caps:
        run.font.small_caps = True
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_paragraph(doc, text, font_size=12, bold=False, italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
                  space_before=0, first_line_indent=None, line_spacing=1.3,
                  color=None, small_caps=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold, italic=italic, color=color,
                 small_caps=small_caps)
    return p


def add_chapter_title(doc, kicker, title):
    """Start a new chapter on its own page with a small-caps kicker + title."""
    doc.add_page_break()
    k = doc.add_paragraph()
    k.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k.paragraph_format.space_before = Pt(60)
    k.paragraph_format.space_after = Pt(6)
    krun = k.add_run(kicker.upper())
    set_run_font(krun, size=12, bold=True, color=ACCENT, small_caps=True)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(28)
    trun = t.add_run(title)
    set_run_font(trun, size=20, bold=True)


# ---------------------------------------------------------------------------
# HISTORICAL FIGURES — extracted from the dissertation PDF at build time
# ---------------------------------------------------------------------------
# Each illustration is keyed to the exact image object (xref) inside
# THOMAS-PRIMARY-2025.pdf. The xrefs were verified by rendering the source
# pages and matching every figure to its caption.

ILLUSTRATIONS = {
    1:  dict(xref=156, caption="A view of Southwark on the south bank of the Thames, 1616. The riverside district was famous for its taverns; the hanging sign of the Bear Inn is picked out at the center. After H. C. Shelley (1909)."),
    2:  dict(xref=157, caption="The George Inn, Southwark — one of the last galleried coaching inns in London, still trading beneath a painted sign."),
    3:  dict(xref=167, caption="The Stella Artois horn descends from Den Hoorn, a brewery first recorded in the 1300s — arguably the oldest trademark still in daily use."),
    4:  dict(xref=177, caption="Enslaved Africans at work on an Antiguan sugar plantation — the first link in the chain that turned cane into molasses, and molasses into rum. From Ten Views in the Island of Antigua, after William Clark; Yale Center for British Art."),
    5:  dict(xref=187, caption="A Caribbean boiling house and distillery worked by enslaved laborers, where the trade's sugar became spirit. After William Clark; Yale Center for British Art."),
    6:  dict(xref=198, caption="An early American broadside advertising alcohol, 1868. Duke University Libraries."),
    7:  dict(xref=199, caption="A painted alcohol billboard, 1916 — outdoor advertising at industrial scale. Duke University Libraries."),
    8:  dict(xref=208, caption="A distiller's mail-order card, c. 1910 — alcohol solicited by post. Duke University Libraries."),
    9:  dict(xref=216, caption="A color newspaper advertisement for alcohol, 1912. Duke University Libraries."),
    10: dict(xref=231, caption="Anheuser-Busch's Bevo “near beer” in The Alamance Gleaner, 1919 — keeping the brand in front of the public all through Prohibition. North Carolina Digital Heritage Center."),
    11: dict(xref=237, caption="A broadside for Schlitz's FAMO near beer. Smithsonian Institution."),
    12: dict(xref=238, caption="A delivery truck advertising malt syrup — the wink-and-nod ingredient for home brewing during Prohibition. After Klein (2019)."),
    13: dict(xref=249, caption="U.S. Signal Corps photograph: “American soldiers in a captured German trench drinking beer out of steins and smoking cigars.” National WWI Museum and Memorial."),
    14: dict(xref=254, caption="An Anheuser-Busch Bevo advertisement wrapping near beer in wartime patriotism. Anheuser-Busch."),
    15: dict(xref=265, caption="Rodeo Lager sells beer through the figure of the cowboy, 1936. Lehmann Printing & Lithographing Co."),
    16: dict(xref=273, caption="A frontier explorer fronts this 1930s bourbon advertisement — masculinity as a sales pitch in the depths of the Depression. Lehmann Printing & Lithographing Co."),
    17: dict(xref=274, caption="California Belle port wine, 1935, sold to women through the beauty ideals of the day. Lehmann Printing & Lithographing Co."),
    18: dict(xref=289, caption="A 1943 Schenley advertisement urging frugality and the purchase of war bonds — patriotism inside a whiskey ad. American Century Shop."),
    19: dict(xref=295, caption="Schenley ties its whiskey to the war effort, 1943. American Foreign Service Association."),
    20: dict(xref=303, caption="American soldiers celebrate with French wine in the streets of liberated Paris, 1944. Popperfoto."),
    21: dict(xref=309, caption="U.S. infantry receive their beer ration on Sterling Island, 1944. National WWII Museum."),
    22: dict(xref=310, caption="A Royal Navy seaman draws his rum ration, c. 1942 — a tradition that survived into the 1970s. Central Press."),
    23: dict(xref=327, caption="Budweiser on television, three decades apart. Left: “Where there's life…there's Bud” (1956). Right: “Wassup” (1999)."),
    24: dict(xref=339, caption="Six decades of product placement: James Bond and Smirnoff, from Dr. No (1962) to No Time to Die (2021)."),
    25: dict(xref=351, caption="The “Lite House” outside the Dallas Cowboys' stadium — a Miller Lite zone built to keep fans drinking before and after the hours the stadium itself can sell beer. Dallas Cowboys."),
    26: dict(xref=359, caption="George Clooney and Casamigos. Celebrity and brand became nearly indistinguishable; the tequila later sold to Diageo for about a billion dollars."),
    27: dict(xref=371, caption="Jack Daniel's online, then and now. Top: c. 1997. Bottom: c. 2023."),
    28: dict(xref=387, caption="A Smirnoff post engineered to read like a friend's: selfies, hashtags, emojis, likes, comments, shares. Smirnoff US, c. 2023."),
}

_PDF = None


def _extract_image(xref):
    """Pull one image object out of the source PDF, normalize to RGB PNG, cache."""
    global _PDF
    os.makedirs(IMG_CACHE, exist_ok=True)
    out = os.path.join(IMG_CACHE, f"x{xref}.png")
    if _PDF is None:
        _PDF = fitz.open(PDF_PATH)
    pix = fitz.Pixmap(_PDF, xref)
    if pix.n - pix.alpha >= 4:          # CMYK -> RGB
        pix = fitz.Pixmap(fitz.csRGB, pix)
    if pix.alpha:                        # drop alpha for clean print
        pix = fitz.Pixmap(pix, 0)
    pix.save(out)
    return out, pix.width, pix.height


def add_illustration(doc, n, max_w=4.9, max_h=7.0):
    """Embed illustration *n* centered, scaled to fit, with its caption."""
    info = ILLUSTRATIONS[n]
    path, w, h = _extract_image(info["xref"])
    ar = h / float(w)
    disp_w = max_w
    if disp_w * ar > max_h:
        disp_w = max_h / ar
    doc.add_picture(path, width=Inches(disp_w))
    pic_par = doc.paragraphs[-1]
    pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_par.paragraph_format.space_before = Pt(10)
    pic_par.paragraph_format.space_after = Pt(4)
    add_paragraph(doc, info["caption"], font_size=9.5, italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14,
                  space_before=0, line_spacing=1.1)


def add_body(doc, text):
    """Render body text. Paragraphs are separated by blank lines.
    A line of the form [[IMG:n]] embeds illustration n."""
    first = True
    for para in text.strip().split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith("[[IMG:"):
            n = int(para[len("[[IMG:"):-2])
            add_illustration(doc, n)
            first = True  # paragraph after an image reads as a fresh start
            continue
        indent = 0.0 if first else 0.3
        add_paragraph(doc, para, first_line_indent=indent)
        first = False


# ---------------------------------------------------------------------------
# FRONT MATTER
# ---------------------------------------------------------------------------

def build_title_page(doc):
    for _ in range(5):
        add_paragraph(doc, "", space_after=0)
    add_paragraph(doc, "THE OLDEST SELL", font_size=34, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_paragraph(doc, "How Alcohol Mastered Every Medium,\nfrom Clay Tablet to the Feed",
                  font_size=15, italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
    add_paragraph(doc, "JACOB EDWARD THOMAS", font_size=14, small_caps=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_paragraph(
        doc,
        "Adapted from doctoral research at The University of Texas at Austin",
        font_size=11, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


def build_copyright_page(doc):
    doc.add_page_break()
    for _ in range(2):
        add_paragraph(doc, "", space_after=0)
    lines = [
        "The Oldest Sell: How Alcohol Mastered Every Medium, from Clay Tablet to the Feed",
        "",
        "Copyright © 2025 Jacob Edward Thomas.",
        "",
        "This book is adapted from the historical chapters of the author's doctoral "
        "dissertation, Alcohol Marketing on Social Media During the COVID-19 Pandemic: "
        "Historical Perspectives, Modern Evidence, and Future Regulation (The University "
        "of Texas at Austin, 2025).",
        "",
        "The historical advertisements, photographs, and other figures reproduced here "
        "are drawn from public archives and library collections and are credited in their "
        "captions. Rights to those images remain with their respective holders; they appear "
        "here for historical and educational commentary.",
        "",
        "The narrative text is released under the MIT License together with the rest of the "
        "project repository.",
    ]
    for ln in lines:
        add_paragraph(doc, ln, font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      space_after=6, line_spacing=1.2)


def build_epigraph_page(doc):
    doc.add_page_break()
    for _ in range(6):
        add_paragraph(doc, "", space_after=0)
    add_paragraph(doc,
                  "“Every civilization that has left a record has also left "
                  "evidence of drinking. Alcohol is older than banking, older than "
                  "organized religion, older than writing itself.”",
                  font_size=13, italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24,
                  line_spacing=1.4)


def build_contents_page(doc, entries):
    doc.add_page_break()
    add_paragraph(doc, "CONTENTS", font_size=16, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24,
                  space_before=24, small_caps=True)
    for kicker, title in entries:
        line = title if kicker is None else f"{kicker} · {title}"
        add_paragraph(doc, line, font_size=11.5,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
                      line_spacing=1.2)


# ---------------------------------------------------------------------------
# PROSE
# ---------------------------------------------------------------------------

INTRODUCTION = """
Every civilization that has left a record has also left evidence of drinking. The Sumerians brewed beer before they invented the wheel. The Egyptians paid the laborers who raised the pyramids in rations of ale. The Romans planted vineyards at the farthest edges of their empire, and when the legions finally marched home, the vines stayed behind and kept on growing. Alcohol is not merely old. It is one of the most resilient enterprises in the history of commerce — older than banking, older than organized religion, older than writing itself.

This is a book about how that enterprise became an industry, and how the industry came to occupy so much of modern life. It is a story that begins with a residue of rice and honey at the bottom of a nine-thousand-year-old pot in the Yellow River valley, winds through ancient temples and medieval taverns, crosses the Atlantic in the holds of slave ships loaded with rum, and arrives in the twenty-first century as a notification on a phone screen at two in the morning. Along the way it touches war and politics, race and gender, science and salesmanship, prohibition and excess.

If there is a single thread that runs through all of it, it is this: the alcohol industry has understood, better than almost any other commercial enterprise in history, that selling a drink is really about selling a story — and that the best way to tell a story is to seize whatever new medium people have just begun to trust. The clay tablet. The printed page. The painted sign. The newspaper, the billboard, the radio, the television screen. The website, the inbox, the social feed. Each time human beings invented a new way to talk to one another, the alcohol industry was among the first to figure out how to talk to them through it, and often the very first to do it well.

That pattern is not an accident. It is a strategy, and it is the spine of this book. Over and over, across millennia and across technologies, three habits recur. The industry attaches itself to each new communication medium as it appears. It reaches hardest for people at their most vulnerable — soldiers, the grieving, the frightened, the poor. And it builds portrayals of drinking that long outlive the campaigns that made them, until they harden into something that feels less like advertising than like common sense. By the end of these pages those three habits will be impossible to unsee.

A word about how the book is built. It moves through the long history one medium at a time — the painted sign, the printed page, the broadcast screen, the digital feed — because that is how the industry itself has moved, vaulting from each communication technology to the next as soon as it appeared. Woven through that chronological journey is the human story of who did the drinking and who did the selling: the enslaved and the soldiers, the grieving and the frightened, the founders and the brewers and the admen. The two threads are really one. The history of how alcohol was sold cannot be separated from the history of how people were reached, and the history of how people were reached is, to a remarkable degree, the history of communication itself.

America occupies a special place in the story. No other country has swung so violently between drunkenness and temperance, between celebration and criminalization. The United States tried to outlaw alcohol outright, failed spectacularly, and then spent the following century building the most sophisticated alcohol-marketing apparatus the world has ever seen — an apparatus so thoroughly woven into the culture that most people no longer notice it at all. Beer sponsors the biggest broadcast of the year. Wine anchors a thousand lifestyle brands. Spirits power the celebrity-endorsement machine. And beneath the glossy surface runs a history far stranger and darker than the advertisements would ever let on.

The numbers are easy to recite and hard to feel. Americans pour roughly fifty billion dollars a year into alcohol, and the advertising that sustains that spending has become so thoroughly woven into ordinary life that it has gone effectively invisible — which is precisely the condition the industry has spent centuries working to achieve. A successful advertisement announces itself. A perfected one disappears, until the message it carries feels less like a pitch than like something everyone already knows: that a hard day earns a drink, that celebration requires one, that this brand belongs to people like you. To write the history of alcohol marketing is therefore to make the invisible visible again — to slow down a story that has been engineered to slip past notice, and to look hard at the machinery underneath.

I came to that story by an unusual route. I am not a brewer or a bartender or an advertising man; I am a public-health researcher who spent the better part of a decade studying how the alcohol industry behaved on social media during the COVID-19 pandemic, and how that behavior rhymed with a century of earlier adaptations. The deeper I dug into the present, the further back the roots ran — to television and radio, to newspapers and broadsides, to printed manuals and painted signs, and finally to a residue at the bottom of a nine-thousand-year-old pot. This book is the long view that my research kept demanding, written for anyone, and it ends roughly where my own studies begin: in the spring of 2020, with a frightened country locked indoors and an ancient industry waiting in the feed.

This is not a temperance tract, and it is not a toast to drinking culture. It is an attempt to tell the story plainly — to trace how an ancient commodity became a modern juggernaut, and to keep an honest accounting of the human costs along the way. The chapters move roughly in order of time, from the first fermented beverages to the isolated, screen-lit households of the COVID-19 pandemic, when the oldest sell met the newest medium under the worst possible conditions. The stories change. The selling never stops.
"""

CHAPTER_1 = """
Somewhere in the Yellow River valley, around seven thousand years before the common era, someone left a mixture of rice, honey, and hawthorn fruit sitting in a clay jar and went away. When they came back, the liquid had changed. It fizzed. It smelled sharp and sweet. It made them feel different. That anonymous moment — reconstructed thousands of years later from the chemical residue dried into pottery fragments at a site called Jiahu, in present-day northern China — is the oldest hard evidence we have of people deliberately making something to get drunk on. It is also, if you squint, the opening scene of a nine-thousand-year commercial enterprise.

The residue at Jiahu dates to somewhere between 7000 and 6600 B.C.E., which makes intentional fermentation older than the wheel, older than bronze, older than the first cities. And it did not happen only once, in only one place. Alcohol emerged independently on nearly every inhabited continent, wherever there was sugar and warmth and time. The reason is almost banal: fermentation is what happens naturally when certain yeasts and bacteria find sugar. Ethanol — the intoxicating compound in every alcoholic drink — is simply their waste product. Our ancestors did not invent alcohol so much as notice it, chase it, and learn to make it on purpose. Once they could, they almost never stopped.

The evidence of how thoroughly they embraced it is scattered across the ancient world. The Sumerians, who gave us some of the earliest writing, also gave us one of the earliest recipes: a hymn to Ninkasi, their goddess of beer, that doubles as a set of brewing instructions, so that to praise the goddess was to recite the method for making her drink. The Egyptians refined the process into an industry and used it as a currency of labor, distributing beer as part of the wages paid to the workers who raised the monuments at Giza — a civilization, in other words, partly built on the promise of a daily ration. Far to the west, in Mesoamerica, the fermented sap of the agave became pulque, a drink so sacred and so socially dangerous that the Aztecs hedged it with elaborate law, in some cases punishing public drunkenness with death. In sub-Saharan Africa, communities brewed sorghum beer for communal ceremony. Wherever people settled, fermented something, and gathered to drink it, the same double character appeared: alcohol as a gift of the gods and a threat to good order, to be both celebrated and contained.

Alcohol was also, from very early on, medicine. For most of human history it was among the few reliably effective drugs available — a painkiller, a sedative, an antiseptic, a base in which to dissolve and preserve other remedies. A physician with little else in his bag could at least dull suffering with wine or spirits, and the world's oldest texts reflect this. The Hebrew scriptures advise giving strong drink to the dying and wine to those in bitter distress; Greek and Roman medicine prescribed wine for an enormous range of complaints. This medicinal authority gave alcohol a moral cover that no mere intoxicant could claim, and the industry would lean on it for thousands of years. The fifteenth-century apothecaries who first distilled brandy sold it explicitly as a health tonic. American patent-medicine men of the nineteenth century would do the same, peddling high-proof "cures" by the bottle. Even the language survives in our own time, in the wellness gloss on a glass of red wine. The claim that a drink is good for you is one of the oldest sales pitches in the human repertoire, and it works because, for a very long stretch of history, it was sometimes even true.

And almost as soon as there was a trade in drink, there was an attempt to regulate it — which tells us that the commerce in alcohol was substantial enough, and troublesome enough, to demand the attention of the law. The Code of Hammurabi, carved in Babylon nearly four thousand years ago, is among the oldest legal texts humanity possesses, and it finds room amid its famous pronouncements to govern the alehouse. It set rules for what a tavern-keeper could charge, prescribed harsh punishment for those who cheated their customers on the measure, and even made the tavern responsible for reporting criminals who gathered under its roof. One notorious provision threatened a dishonest alewife with death by drowning. We need not take every detail literally to grasp the larger point: by the time of Hammurabi, the selling of drink was a recognized commercial institution, embedded enough in daily life that a king thought it worth his statutes.

That detail reframes the whole ancient story. We tend to picture early drinking as a purely communal or sacred affair, and much of it was. But behind the ritual there was already a business — a tavern with prices and a proprietor, a measure that could be honest or short, a customer who could be served fairly or cheated. The alehouse was one of civilization's first commercial spaces, a place where strangers exchanged money for a manufactured pleasure, and like every such space it bred both conviviality and fraud. The long contest this book describes — between an industry forever seeking advantage and a society forever trying to set limits on it — is not a modern development. It is nearly as old as the city itself, and the alcohol trade has been refining its side of the argument for four thousand years longer than its regulators have.

What almost every one of these early cultures shared was the conviction that drinking meant something. It was woven into religion and kinship and the rhythm of the year. The drink was not yet a product to be branded and sold to strangers; it was a sacrament, a wage, a rite. The long transformation this book traces is, in one sense, the story of how that sacred, communal substance was slowly pried loose from ritual and remade into a commodity — and of who profited from the prying.

For most of that long prehistory, drinking was not casual. The earliest evidence points overwhelmingly to ritual and spirit. In Neolithic Europe, pagan communities drank during elaborate ceremonies marking the turning of the seasons. Around Stonehenge, archaeologists have found hearths packed with animal bones and the broken remains of pottery vessels — the debris of great feasts where roasting and drinking went together. Drinking was how you marked a death, sealed a bargain, honored a god, or survived a winter. It belonged to the calendar and the altar long before it belonged to the marketplace.

It was in the Mediterranean world that drinking first became something we would recognize as an industry. The Greeks did not merely drink wine; they organized their social and intellectual lives around it. They mixed it with water for everyday fortification, and they built one of their central civic institutions, the symposium — literally "drinking together" — into the venue where politics, philosophy, and poetry actually happened. Their oldest literature is soaked in it. In the opening book of the Iliad, the Greek commander Agamemnon stages an elaborate feast, complete with poured offerings of wine, in an effort to win the favor of the gods and steel his army for the war at Troy. The poem treats the soldiers' rations of wine as utterly unremarkable, which is the most revealing detail of all: by the eighth century B.C.E., wine was already a logistical staple of organized warfare. The link between alcohol and the fighting man, which will return again and again in this book, is very nearly as old as written war itself.

The symposium deserves a closer look, because it shows how thoroughly a culture can organize itself around a drink. This was not a tavern crawl; it was a structured institution with its own etiquette and offices. The wine was mixed with water in a great bowl, the krater, in proportions set by a chosen master of ceremonies, the symposiarch, whose job was to keep the gathering on the right side of the line between conviviality and chaos. Around that bowl, reclining on couches, the citizens of a Greek city conducted a remarkable share of their public life — debating politics, trading verses, testing philosophical arguments, forging the alliances and rivalries that shaped the city's affairs. Some of the foundational texts of Western thought are set at drinking parties; the philosophers who shaped two millennia of argument did much of their arguing with cups in their hands. To the Greek mind, wine properly handled did not dull the intellect but loosened and sharpened it, opening a space where men spoke more freely and thought more boldly than the sober daylight allowed.

The Greeks also bequeathed a habit of mind that the alcohol industry would exploit forever after: the conviction that the right drink, taken the right way, was a mark of cultivation and civilization itself. To drink well was to be a certain kind of person — refined, social, fully a citizen. The barbarian, by contrast, drank his wine unmixed and fell into a stupor; the civilized man diluted it, paced himself, and let it elevate the conversation. Here, more than two thousand years before the wine brand and the lifestyle commercial, is the basic promise that alcohol marketing would sell again and again: that what is in your cup says something flattering about who you are. The drink was already more than a drink. It was a costume, a credential, a story the drinker told about himself — and stories, as the rest of this book will show, are the thing the industry has always truly sold.

Rome took what the Greeks had built and industrialized it. Wine was a staple in the diet of nearly every Roman citizen and part of the standard ration for the legions. As Roman power and Roman engineering spread, so did organized viticulture. The Romans standardized the cultivation of grapes and the formulas for wine the way they standardized roads and aqueducts and law, and that impulse produced what may be the first piece of alcohol-industry literature in history: De Agri Cultura, written by the statesman Marcus Porcius Cato around 160 B.C.E. It was a practical manual for landowners, full of unsentimental instruction on which slopes to plant, how to press and ferment, how to store, and how to ship. Cato's advice assumed something new in the world — that wine was not only a sacred drink or a household staple but a commodity, to be produced at scale on estates worked by enslaved laborers and moved through an elaborate distribution network for profit.

The manual's specificity is what makes it feel so modern. It told the would-be wine magnate to seek out land with south-facing slopes and good drainage, to invest heavily in enslaved labor, to build the necessary infrastructure of presses and storehouses, and to arrange the shipping and distribution that would carry the product to market. This was not advice for a farmer making a little wine for his household and his gods; it was a business plan. And because it circulated widely among Roman landowners, it helped standardize an entire industry, turning winemaking from a scatter of local customs into a repeatable, scalable enterprise. Cato had grasped a principle the alcohol trade would rediscover in every later age: that the most powerful thing you can put into a new medium is not a single advertisement but a method — a way of doing business that others can copy and spread until it becomes the water everyone swims in.

That assumption changed how people drank. With supply pouring out of industrial-scale Roman estates, the act of drinking drifted away from the temple and toward the tavern, from the seasonal rite toward the everyday and, increasingly, the excessive. And as soon as drinking became ordinary, the warnings began. Drunkenness — as opposed to drinking — came to be seen as a public nuisance and, by some, as a genuine danger. The Stoic philosopher Seneca wrote that habitual intoxication weakens the mind and that its damage lingers long after the last cup. The irony was that Seneca lived in a society where wine was as common as water, and very often safer to drink, given the state of Roman plumbing.

The Roman achievement was not merely to drink a great deal but to spread the means of drinking everywhere the empire reached. As the legions advanced into Gaul, Iberia, and the valley of the Rhine, vineyards advanced with them, planted to supply the soldiers and the settlers and then left behind to flourish. Many of the great wine regions of modern Europe trace their roots to this military viticulture; the empire receded, but the vines stayed in the ground and kept producing, an agricultural infrastructure outliving the political one that planted it. Wine became one of the threads that bound the Roman world together — a portable piece of Roman identity that a legionary could carry to the edge of the known world and a conquered people could be taught to want.

And it was sold and served through an infrastructure that would look familiar to any modern city-dweller. Roman towns were dense with taverns and wine shops, the popinae and tabernae where ordinary people drank, ate, gambled, and gossiped, and where the wealthy were generally too proud to be seen. Wine was graded and priced by quality, from the cheap stuff watered down for the masses to the prized vintages of celebrated estates, so that what a person drank announced where they stood. The drink had become, in other words, a full consumer market, complete with brands of a sort, price tiers, and the social signaling that the alcohol industry would monetize for the next two thousand years. When Rome fell, this apparatus did not vanish so much as fragment and survive, ready to be reassembled by the medieval world that rose from the ruins — a world that, as we will see, drank even harder than Rome.

What the ancient world established was not just a taste for alcohol but a template for how human societies would handle it forever after. Alcohol was simultaneously sacred and profane, celebrated and feared, regulated and indulged. It was nutrition for the poor and luxury for the rich. It was medicine, currency, social glue, and poison, often at the same time and in the same cup. And it was, from very early on, a business — one that already understood that the drink itself was only half of what it was selling. Nine thousand years before the first television commercial, before the first painted tavern sign, before the first printed advertisement, the alcohol trade was already doing the thing it does best: meeting the world exactly as it found it, and selling whatever version of the drink that world was prepared to buy.
"""

CHAPTER_2 = """
When Rome fell, Europe did not stop drinking. If anything it drank more. The thousand years we loosely call the medieval period were, among many other things, an age of ale: thick, dark, sweetish, and everywhere. Clean water was scarce and often dangerous in crowded settlements, and weak ale — brewed from cereal grains and safer than the local water — became a genuine source of daily nutrition for ordinary people. It was breakfast and wages and comfort. It was also, predictably, a source of widespread drunkenness across every level of society, from the manor to the field. Monasteries brewed at enormous scale, both for their own tables and for sale, and in doing so the Benedictine and Trappist orders became, almost by accident, some of the most accomplished brewers in history.

It is hard to overstate how much medieval Europe drank, because the drinking was bound up with simple survival. In crowded towns the water carried disease, and weak ale — boiled in the brewing, and mildly preserved by its own alcohol — was often the safer choice. By some estimates, the average person in parts of Europe around the year 1500 might consume on the order of a hundred liters of alcoholic beverage a year, much of it low-strength ale standing in for water that could not be trusted. To call this a drinking culture is almost to miss the point. Alcohol was infrastructure. It was how people stayed hydrated, fed, and calorically afloat through hard winters and harder centuries, and the institutions that produced it — the monastery, the manor brewery, the town alehouse — were as woven into daily life as the mill or the church. Out of that ubiquity grew both an industry and a problem, and the medieval mind, like the Roman one before it, learned to hold the two ideas at once: that drink was a daily necessity, and that drunkenness was a daily threat.

The medieval town is also where we find the first thing that looks unmistakably like alcohol marketing. In 1389, King Richard II of England decreed that any establishment selling ale had to hang out a sign announcing the fact. The stated purpose was regulatory — to let the official ale-taster know where to inspect — but the effect was competitive. Once every alehouse had to display a sign, the natural next move was to display a better one, and taverns began vying to hang the most distinctive, eye-catching, memorable marker on the street. The idea of forcing drink-sellers to identify themselves was not original to Richard; some historians trace it to a much older Chinese practice, in which villagers would hoist a banner when they had fresh wine to trade. But Richard's law helped seed a culture of flamboyant marketing competition in the alcohol trade that has never really stopped, and the English pub carries the tavern-sign tradition to this day.

[[IMG:1]]

[[IMG:2]]

The sign was not the only piece of marketing one might meet in a medieval town. Town criers roamed the streets of Europe shouting the news — royal decrees, new laws, and, for a fee, advertisements. In France under King Philip Augustus, around the turn of the thirteenth century, a class of so-called "wine criers" was licensed to stand outside taverns, call out the merits of the wine within, and charge the owner for every customer they pulled through the door. Here, eight centuries before the influencer and the paid endorsement, was the basic transaction already fully formed: a trusted voice, a public medium, and a drinks merchant paying to borrow both.

That transaction is worth holding onto, because it is the seed of something this book will return to again and again. The wine crier worked because he was already part of the town's trusted apparatus of news — the man you listened to for decrees and announcements — and the tavern keeper paid precisely to borrow that trust and bend it toward a sale. Strip away the medieval costume and you have the exact logic of the celebrity endorsement, the sponsored post, and the influencer's casual product mention. The genius was never in the message alone; it was in smuggling the message inside a voice the audience had already decided to believe. Seven hundred years later, the alcohol industry would spend fortunes rediscovering what King Philip's wine criers knew by instinct.

Another medieval trademark has outlived nearly everything around it. In the late 1300s a brewery called Den Hoorn — "the Horn" — was operating in the Brabant town of Leuven, its shop sign decorated with an elaborate golden horn. In 1717 a brewer named Sebastian Artois bought the business and renamed it after himself: Stella Artois. The brand is sold around the world today, and its logo still carries a golden horn descended directly from that fourteenth-century sign. It is, by a reasonable measure, the oldest continuously running marketing campaign on earth — an unbroken thread of brand identity more than seven hundred years long.

It is worth dwelling on what the medieval tavern sign actually was, because it was the first mass branding the alcohol trade ever produced. Most people in a medieval town could not read, which meant a written name above a door was useless. The sign had to communicate through image alone — a bear, a horn, a red lion, a green man — a simple, memorable mark that a customer could recognize at a distance and recall the next time thirst struck. This is precisely the function of the modern logo, and the taverns arrived at it by the same pressures that drive branding today: a crowded market, a barely literate public, and the need to be the establishment a passerby remembered. The competition Richard II's law set off was, in effect, the first logo war, fought with paint and carved wood across the streets of medieval England. The horn that still rides on a bottle of Stella Artois is a survivor of that war — proof that a strong enough mark can outlast empires, languages, and the very memory of what it originally meant.

Behind the signs lay an economy larger than we tend to imagine. Brewing was one of the most widespread trades of the medieval world, often carried out by women — the "alewives" who brewed in their homes and sold the surplus — and scaled up by monasteries into something approaching industrial production. The alehouse and the tavern were not merely places to drink; they were the social infrastructure of the town, serving as meeting halls, courrooms, marketplaces, and inns. To brand a drinking house, then, was to brand a hub of communal life, and the brewers who learned to do it well were learning, centuries early, that you do not merely sell a product — you sell a place, an identity, a destination people choose. The lesson would echo down to the brand websites and social feeds of our own age, where the alcohol industry would again work to become not just a product on a shelf but a place people willingly gathered.

[[IMG:3]]

Out of the medieval world came the Renaissance, and with it a renewed moral anxiety about drink. Drawing on Christian philosophy, Renaissance moralists cast excessive drinking as a contributor to social decay, especially against a backdrop of poor conditions and recurrent disease. The earliest temperance arguments took shape here, and they had a particular structure that would echo for centuries: God, the argument went, had given humanity alcohol for moderate pleasure, enjoyment, and even health — but drunkenness was a sin. Use, yes; abuse, no. It is essentially the same line the industry itself would adopt, much later, when it learned to brand its own products as things to be enjoyed "responsibly."

And yet the most consequential development of the era ran in precisely the opposite direction. While moralists preached restraint, the science of the Renaissance was busy making alcohol far stronger. Distillation — heating a fermented liquid to separate and concentrate its alcohol — was not a European invention. The technique had been refined for centuries by Islamic alchemists, largely for perfumes and medicines. When that knowledge migrated into Christian Europe, fifteenth-century German apothecaries adapted these centuries-old methods into a potent new product, which they marketed, tellingly, as a health tonic. They called it gebrant wein — "burnt wine" — the ancestor of brandy.

The chemistry behind the new product was elegant and, once understood, easy to exploit. Alcohol boils at a lower temperature than water — around seventy-eight degrees Celsius, against water's hundred — so when a fermented liquid is heated within that window, the vapor that rises off it is far richer in alcohol than the liquid left behind. Capture that vapor, cool it, and collect it, and you have a spirit several times stronger than any beer or wine. This was not a refinement of fermentation; it was a multiplication of it.

The arrival of distilled spirits was one of those quiet hinges on which the whole history turns, because it changed not only how strong a drink could be but what kind of business alcohol could become. Beer and wine are bulky, perishable, and bound to a place; they spoil, they sour, they are heavy to move and hard to keep. Spirits are none of these things. A barrel of brandy or rum holds far more intoxication per pound than a barrel of ale, survives a sea voyage that would ruin wine, and can be made from almost any sugar or grain a region happens to have. In economic terms, distillation turned a local, fragile good into a durable, concentrated, globally tradeable commodity — the difference between selling fruit and selling something closer to fuel. It is no accident that the spread of spirits tracks so closely with the age of empire and ocean trade. A drink that could cross oceans without spoiling was exactly the drink an expanding, seafaring Europe needed, and the men who financed the ships understood it well.

It changed the human cost of drinking, too. For all of recorded history, getting truly, dangerously drunk had taken effort and volume; weak ale and watered wine set a natural ceiling on how fast a person could descend. Spirits removed the ceiling. A few cheap measures of strong liquor could do what gallons of beer once required, and societies that had managed their relationship with alcohol for centuries suddenly found themselves confronting a far more concentrated form of the same old temptation. The pattern that the rest of this book traces — a powerful new capability arriving faster than the culture's ability to absorb it — had appeared before in the medium of communication. Here it appeared in the medium of the drink itself. And it changed the nature of the product in the marketplace. Brandy carried a reputation for being lethal in excess, and it earned that reputation honestly, but it was also potent, durable, and profitable in a way nothing before it had been.

The pattern is worth pausing on: a sacred body of knowledge from one culture, repurposed by another into a commercial intoxicant and sold as medicine. We will see it again when American distillers build an industry on a crop held sacred by the peoples they displaced.

Brandy carried a reputation for being deadly in excess, but it sold anyway, and in 1500 it found its great promoter. A German physician named Hieronymus Brunschwig published Liber de arte distillandi — a book devoted entirely to the art of distillation — and it circulated widely. Gutenberg's press had been running for barely half a century, and Brunschwig was among the first to grasp what the new technology could do for technical knowledge: spread it, fast and far, beyond the reach of any single workshop or guild. Just as Cato's farming manual had carried the methods of industrial winemaking across the Roman world some sixteen centuries earlier, Brunschwig's distillation manual helped lay the foundations of a global spirits industry. The medium was new — movable type — and once again the drinks trade was among the very first to exploit it.

The printing press deserves its reputation as a hinge of history, and its role in this story is larger than a single distillation manual. Before print, knowledge moved at the speed of the copyist's hand; a technique known to one workshop might take generations to spread, if it spread at all. Print shattered that limit. A method could now be set in type once and struck off in the hundreds, carrying identical instructions to readers a continent apart. For the alcohol trade this was transformative twice over. First, it standardized and spread the methods of production, so that brewing and distilling could be learned from a book rather than an apprenticeship, and the industry could grow faster than any guild could have allowed. Second, and more lastingly, it created the conditions for mass persuasion. A press that could print a manual could print a handbill, a label, a poster, a notice in a newssheet — the same brand message, reproduced endlessly, placed before thousands of strangers who had never met the seller.

This is the deep reason the alcohol industry's fortunes have tracked so closely with communication technology, and it is worth stating plainly here at the technology's birth. A drink is a local, physical thing; it must be made somewhere and carried somewhere to be sold. But a brand is information, and information wants to travel. Every leap in the capacity to copy and transmit information — from the scribe's manuscript to the printing press, and onward to the broadcast tower and the data center — is therefore a leap in the reach of branding, and the alcohol industry, selling a product whose margins easily funded the experiment, was positioned to seize each leap as it came. Brunschwig's little book on distillation was an early sign of a pattern that would define the industry forever: that whoever masters the newest way of spreading information will master the newest way of selling drink.

The consequences were enormous, because distilled spirits were a fundamentally different kind of product from beer or wine. They were stronger, more concentrated, more portable, and far more durable. They could survive long sea voyages without spoiling. They could be made from almost any fermentable material. They were, in a word, scalable — and they arrived just as European ships were beginning to cross the Atlantic in earnest. The spirits industry and the colonial enterprise would grow up side by side, feeding each other. Neither, as the next chapter shows, was innocent.
"""

CHAPTER_3 = """
The story of rum is the story of empire, and it is written in blood. By the seventeenth century, distilled spirits were available across the western world, and a vast new trading system had organized itself around the sugar of the Caribbean. We remember it by a tidy geometric name — the triangular trade — that does very little to convey what it was. Ships carried manufactured goods and rum from Europe and colonial America to the west coast of Africa. There the cargo was exchanged for human beings. Those men, women, and children were chained into the holds and carried across the Atlantic to the sugar islands, where the survivors were sold to plantations and forced to cultivate and process the cane. The cane became molasses; the molasses was shipped north to distilleries — many in New England, some in the islands themselves, such as Barbados's Mount Gay, established in 1703 and still sold worldwide — and there it became rum. The rum was sold back into Europe and Africa, and the triangle turned again.

The phrase "triangular trade" lends the whole arrangement a clean, almost geometric respectability that the reality did not deserve. What it describes is an economy in which human beings were a commodity priced against barrels of spirit, and in which the demand for cheap sugar and cheap rum drove the demand for stolen labor. Rum distilling became one of colonial New England's signature industries, and the molasses that fed it tied the northern colonies, the Caribbean plantations, and the African coast into a single brutal circuit. The arrangement generated enormous wealth and, with it, enormous friction. When the British Crown moved to tax and regulate the molasses and sugar that the rum trade depended on, the resulting resentment fed directly into the grievances that would erupt into the American Revolution. The country that would soon declare all men created equal was, in significant part, financed by an industry organized around the buying and selling of people.

[[IMG:4]]

[[IMG:5]]

Rum was not a minor by-product of this system. It was an engine of it. Profits from rum financed the purchase of more enslaved people, and rum itself served as currency on the African coast, traded directly for human lives. The economic expansion that the sugar-and-spirits trade made possible was staggering, and it helped build the wealth of colonial America — even as it helped build the tensions that would erupt into revolution. Every barrel of colonial rum carried, invisibly, a ledger of suffering.

That inheritance ran straight through the founding of the United States. A striking number of the founders had direct stakes in the alcohol trade. George Washington operated one of the largest distilleries in the country at Mount Vernon, worked by enslaved distillers; it was strikingly profitable and added meaningfully to his fortune. He also understood alcohol as a political instrument. Campaigning for a seat in Virginia's House of Burgesses in 1758, Washington won over voters with a generous flood of free rum, wine, and beer — retail politics in the most literal sense. Thomas Jefferson, never as commercially successful in drink as Washington, kept a vineyard at Monticello, imported European wine with connoisseur's devotion, left a legendary cellar, and led some of the first serious attempts to cultivate fine wine on American soil. The men who wrote the founding documents were, many of them, also men of the alcohol industry. Drinking was not at the margins of early American life; it was woven into its economy and its politics from the start.

Nor was it a quiet presence. The early republic drank on a scale that would alarm a modern physician. Cider sat on the breakfast table; spirits punctuated the workday; taverns served as courthouses, post offices, and polling places, so that civic life and drinking life were often the same life conducted in the same room. Washington's free-flowing campaign rum was not an aberration but an expression of the norm — politics and alcohol were simply expected to travel together. This was the soil in which the American alcohol industry put down its deepest roots, a culture that treated daily, abundant drinking as unremarkable and that handed the trade a captive, thirsty, growing market. Everything the industry would later build — the brands, the slogans, the campaigns — rose from a foundation of consumption that was already, by any later standard, extraordinary.

As the nineteenth century wore on, that foundation supported something genuinely new: the national brand. For most of history, beer and spirits had been local affairs, brewed and sold within a short distance of where they were made, because the products did not travel well and the country had no way to move them at scale. The industrial revolution changed all three constraints at once. The railroad could carry a product across a continent. Mechanical refrigeration and pasteurization let beer survive the journey. And the new printing and advertising industries could make a brand known in a city it had never physically reached. The brewers who grasped this — the names that would dominate American beer for the next century — stopped thinking of themselves as local producers and began building truly national operations, shipping their product coast to coast and advertising it everywhere the trains ran.

This was a profound shift in what the industry was selling. A local brewer competed on the freshness and price of his beer to neighbors who knew him. A national brewer competed on something far more abstract: reputation, recognition, the feeling a name conjured in the mind of a stranger a thousand miles away who had a dozen other beers to choose from. Winning that stranger's loyalty required exactly the tools the era was perfecting — the trademark, the slogan, the ubiquitous advertisement — and the brewers poured money into all of them. By the time the twentieth century arrived, the largest alcohol companies were among the most sophisticated marketers in the entire economy, masters of a craft that had grown, in a few short decades, from the painted tavern sign into a national apparatus of persuasion. They had also grown rich and powerful enough to provoke a reaction — and the reaction, as the next chapter describes, would try to destroy them outright.

The eighteenth and nineteenth centuries — the age of industrial revolution and surging American capitalism — pushed the trade forward on every front. In 1775 the oldest American distillery still operating today, later known as Buffalo Trace, was founded to make whiskey. Early American whiskey was distinctive because it was distilled largely from corn, rather than the barley, rye, and wheat of the Old World. Corn had been domesticated by the indigenous peoples of southern Mexico and was sacred to many Native cultures — which means that American distillers built a signature industry on an appropriated sacred crop, an unsettling rhyme with those German apothecaries who had turned Islamic alchemy into brandy two centuries before. The roots of familiar brands reach back into this era: a Maryland farmer named Robert Samuels began making the whiskey that would eventually become Maker's Mark, and a man named Jacob Beam started selling a Kentucky corn whiskey that would one day bear the name Jim Beam.

These origin stories are worth lingering on, because the modern whiskey industry sells them as relentlessly as it sells the liquor. Walk down any bourbon aisle and you are surrounded by founding dates, family names, and frontier mythology — the rickhouse, the limestone water, the recipe handed down through generations. Much of this is genuine history, and some of it is marketing burnished to a high shine, but the distinction matters less than the function. A whiskey that can trace itself to a named ancestor in a particular Kentucky hollow is selling more than alcohol; it is selling heritage, authenticity, a story of American self-reliance distilled into a bottle. The industry understood early that the drinker was buying a narrative along with the drink, and that the older and more rooted the narrative, the more the bottle was worth. The frontier distiller and the country he helped settle would be marketed together for the next two centuries, until whiskey and a certain idea of America became almost impossible to pull apart.

It is no small irony that this most "American" of products carries within it the same pattern of appropriation that produced brandy across the Atlantic. The corn at the heart of American whiskey was a gift of the very peoples being driven from the land, repackaged into a symbol of the settlers' national identity. The heritage the bourbon brands sell so lovingly is real, but it is only half the story. The other half — the dispossession that cleared the ground and the sacred crop that was borrowed without acknowledgment — does not appear on any label, then or now.

There is a grim symmetry worth naming before we leave the colonial era behind. Twice now in this story, a people has built a signature spirit on something taken from a culture it was busy destroying. German apothecaries had turned the sacred alchemy of the Islamic world into brandy and sold it as medicine. Now American distillers built their signature whiskey on corn — a crop first domesticated by the indigenous peoples of the Americas and held sacred by many Native nations — even as the new republic dispossessed and slaughtered those same peoples. The rum was a product of slavery; the whiskey, increasingly, a product of conquest. The two founding spirits of the United States were thus soaked, from the start, in the suffering of the people on whose backs the country was built. This is not a metaphor the industry would ever choose for itself, but it is the accurate one, and it belongs in any honest account of where American drinking came from.

The same industrial revolution that scaled up distilling also gave birth to modern advertising — and at every step the alcohol industry was on the frontier of the new media. The first commercial broadsides, printed bills often pasted directly onto walls, appeared in the United States around 1835. Direct-mail advertising was firmly established by the 1840s. Newspaper advertising took off in the 1860s as color presses began to circulate. Branded promotional objects — paper fans, burlap sacks stamped with a distiller's name — spread in the 1880s. And the 1890s brought widespread color advertising in magazines. Just as the medieval tavern keeper had pushed the limits of the painted sign, the American drinks industry pushed the limits of each of these new printing technologies as it arrived, learning to build brand recognition and to tell the public, vividly and repeatedly, exactly where and what to drink.

There is a temptation to see these early advertisements as quaint — the broadside pasted to a brick wall, the ornate magazine engraving, the distiller's mail-order card. But each represented the cutting edge of its moment, and each taught the industry the same expanding lesson. A printed image could travel where no tavern sign could. A slogan repeated across a thousand newspapers could lodge a brand in the minds of strangers who had never met a salesman. A promotional fan or a stamped burlap sack could carry a name into a customer's home and keep it there. With every advance in printing, the distance a brand could reach grew longer and the intimacy with which it could speak grew greater — and the alcohol industry, with its deep pockets and its constant hunger for new customers, was forever among the first to test the limits of what the new presses could do.

It was in these same decades that advertising itself became a profession. What had been the haphazard work of shopkeepers and printers grew into a distinct industry, with agencies that specialized in placing notices, writing copy, and buying space — middlemen who sold persuasion as a service. The alcohol industry was among their most reliable clients, because few products needed persuasion more or could pay for it more handsomely. A distiller did not merely announce that whiskey existed; the public already knew. The job was to make a buyer prefer this whiskey, to attach to a particular label some feeling of quality or pleasure or status that the liquid alone could not supply. That is the work the new advertising profession existed to do, and the alcohol industry helped build it, campaign by campaign.

The surviving artifacts of that era still carry the charge of their moment. A broadside from the 1860s crowds the page with type, shouting its claims in the visual language of a barker on a busy street. A painted billboard from the early 1900s looms at the scale of a building, turning a brand into part of the landscape itself. A mail-order card invites a drink to be summoned to the home by post, a private solicitation slipped among the day's letters. A color newspaper advertisement blooms into the page with an extravagance the black-and-white press could never match. Each was the cutting edge of its decade, and each taught the same expanding lesson — that a printed image could travel where no painted sign ever could, and could repeat itself a thousand times over until a brand felt less like a product than like a fact of the world.

[[IMG:6]]

[[IMG:7]]

[[IMG:8]]

[[IMG:9]]

By the turn of the twentieth century, alcohol was braided so tightly into the economic, political, and social fabric of America that pulling it out seemed unthinkable. Millions of jobs depended on it. Whole regional economies were built around it. The federal government leaned on alcohol excise taxes for a substantial share of its revenue. And the industry's advertising machine had spent decades making drinking feel synonymous with celebration, with masculinity, with patriotism, with freedom itself. It looked permanent. It was about to collide with one of the strangest experiments in the nation's history — and, as the next chapter shows, to survive it.
"""


CHAPTER_4 = """
The story of alcohol in America cannot be told honestly without talking about race, because the same trade that enriched the founders was, from the beginning, organized around exclusion as well as exploitation. Through the long era of emancipation and Jim Crow — from the late nineteenth century into the middle of the twentieth — the alcohol business in the United States was heavily segregated. The sale and consumption of alcohol were frequently restricted for Black Americans, both by discriminatory law and by the everyday violence of segregated bars and establishments. And yet, against those odds, some Black entrepreneurs built businesses in the industry anyway. The A. Smith Bowman Distillery, founded in Fairfax County, Virginia in 1934 — just after Prohibition's repeal — by Abram Smith Bowman and his three sons, was one of the very few Black-owned distilleries in the country, producing bourbon, gin, and brandy.

Such businesses faced obstacles that had nothing to do with the quality of their product. Black-owned bars and clubs were targeted for closure by police, officials, and white supremacist groups using nuisance laws, zoning rules, and outright hostile raids — measures justified in the language of "public safety" and "moral order" but aimed squarely at suppressing Black enterprise and limiting Black mobility. The combined force of Jim Crow and a fast-consolidating industry concentrated alcohol wealth and leadership in almost exclusively white hands — an industry whose foundational products, rum and corn whiskey, were themselves the fruits of slavery and of the displacement of Native peoples. The Bowman distillery survived for several generations within the family before white capital, in the form of the New Orleans–based Sazerac Company, eventually absorbed it through a series of mergers and acquisitions late in the twentieth century. The pattern of the trade — wealth flowing upward and toward whiteness — was remarkably durable.

The injustice ran deeper than exclusion from ownership, because it was woven into the very products on which the industry was built. American drinking rested on two foundational spirits, and both were soaked in the suffering of people of color. Rum was the distillate of a Caribbean sugar economy worked by the enslaved. Corn whiskey was distilled from a crop sacred to the Native nations the young republic was dispossessing, on land taken from those same nations. An industry founded on the forced labor of Black people and the displacement of Native people then spent the following century ensuring that the descendants of both would have the hardest possible time sharing in its profits. The story of the A. Smith Bowman Distillery — a rare Black-owned success, sustained across generations against the grain of every obstacle, before being absorbed by white capital — is exceptional precisely because the system was designed to make it impossible.

The tools of exclusion were rarely as honest as a sign reading "whites only," though those existed too. More often they were the quieter instruments of law and zoning: nuisance ordinances that could shutter a Black-owned bar on a neighbor's complaint, licensing rules applied with a heavy hand to Black applicants and a light one to white, hostile inspections and raids dressed up as concern for public order. The language was always about safety and morality; the effect was always the suppression of Black enterprise and the protection of white market share. It is a particular kind of hypocrisy, and a revealing one. An industry that would spend the twentieth century selling rebellion, freedom, and good times to everyone with a dollar was, behind the marketing, as invested as any other in the racial hierarchies of its age — happy to sell the image of liberation while working to deny the real thing.

It was against this backdrop that the most dramatic experiment in American alcohol policy unfolded. The temperance movement had been gathering force since the early nineteenth century, and by its end it had assembled a genuinely strange coalition: religious conservatives who saw drink as sin, women's-rights activists who saw it as the fuel of domestic violence and ruined households, progressive reformers who saw it as a public-health catastrophe, and — it must be said plainly — nativists and outright racists who associated alcohol with the immigrant communities they despised. The movement was contradictory, and it was effective. In 1920, the Eighteenth Amendment to the Constitution prohibited the manufacture, sale, and transportation of intoxicating beverages across the entire United States.

For the reformers, it was the culmination of a moral crusade nearly a century in the making. The temperance argument had deep and tangled roots — it echoed, in its way, the Renaissance preachers who had condemned drunkenness as a sin while allowing drink in moderation — and by the early twentieth century it carried real human grievances. Women, who had little legal recourse against a husband who drank away the household's money or turned violent, had genuine cause to see the saloon as an enemy of the home. Reformers pointed, accurately, to alcohol's role in poverty, crime, and family ruin. The tragedy of Prohibition is that a movement containing so much legitimate concern produced a policy so spectacularly counterproductive, in part because it was yoked to the uglier energies of nativism and the suspicion of immigrant communities for whom drink was simply part of life. The country approached the noble experiment full of hope that law could cure a social ill. It was about to learn how badly that hope could misfire.

Prohibition was sold as "the noble experiment," and as social policy it was a catastrophe. It did not stop Americans from drinking. It simply drove the drinking underground and handed the trade to organized crime. Speakeasies multiplied into the tens of thousands in a single city. Bootlegging syndicates grew rich and violent. The supply of alcohol did not so much vanish as go feral — less safe, less regulated, and vastly more profitable to criminals.

The human costs piled up in ways the reformers had never imagined. With legitimate production outlawed, drinkers turned to whatever they could get, and what they could get was often dangerous: crudely distilled moonshine, industrial alcohol diverted and inexpertly "renatured," spirits cut with substances that could blind or kill. Deaths from tainted alcohol climbed. Corruption spread through police forces and city governments as bootleggers bought the protection they needed. And the federal treasury, which had leaned on alcohol excise taxes for a substantial share of its revenue, simply lost that money — a self-inflicted wound that looked especially foolish once the Depression arrived and the government found itself desperate for funds. Prohibition managed the rare feat of making nearly every problem it claimed to solve measurably worse: more crime, more dangerous drinking, less public money, and a vast new criminal infrastructure that would not disappear when the law was repealed.

What is less often remembered is how the legitimate alcohol giants survived the dry years, because their survival is a master class in marketing under constraint. Anheuser-Busch publicly pivoted to "near beer," a low-alcohol product that technically complied with the law. But the company also sold malt syrup and yeast — the raw materials of beer — and marketed them, with a wink, to the home brewers who knew exactly what they were for. The product was technically a baking ingredient. Everyone understood what most of it actually became.

The advertising of the dry years is a study in saying one thing while meaning another. The surviving images show brewers performing compliance with enormous, almost theatrical earnestness — wholesome near beer presented as family refreshment, wrapped in the same patriotic and homey imagery the brands had always used, as if nothing at all had changed except the alcohol content. A near-beer advertisement of the period is a small masterpiece of double meaning: every cue says "the beer you have always loved," while a fig leaf of legality covers the single missing ingredient. The brands were not really selling near beer. They were keeping a relationship warm, reminding the public of a name and a feeling, holding a place in the culture against the day the real product could return. And alongside the near beer ran the wink of the malt-syrup advertisement, promoting an "ingredient" with a straight face to customers everyone knew were brewing the genuine article in their basements.

[[IMG:10]]

The Joseph Schlitz Brewing Company, one of the largest brewers in the country before Prohibition, ran the same play, keeping a substantial operation in malt extract alive throughout the dry years even as it advertised its own near beer as wholesome family refreshment.

[[IMG:11]]

[[IMG:12]]

The lesson the survivors absorbed was profound. A great brand is not the same thing as the product in the bottle. A brand is a relationship with the public, and a relationship can be maintained through a decade of prohibition by selling something adjacent — a syrup, a low-proof substitute, a memory — until the real product can return. Keep the name in front of people, keep the story alive, and the law itself becomes a temporary inconvenience rather than a fatal blow.

Prohibition also rewrote the culture of drinking in ways that long outlasted the law. The speakeasy — the hidden bar behind the unmarked door, entered with a password — turned drinking into something glamorous and transgressive, a private rebellion against a public rule. For the first time, respectable women drank alongside men in significant numbers in these illicit rooms, because a bar that was already breaking the law had little reason to honor the old conventions about who belonged at it. The cocktail flourished, partly because mixing cheap or rough bootleg liquor with juice and sugar made it more palatable and harder to identify. An entire aesthetic of sophisticated, slightly dangerous drinking was born in those years, and it proved enormously useful to the industry once the drinking was legal again. Prohibition was supposed to make alcohol shameful. In the speakeasy, it made alcohol cool.

And when repeal came, the industry that walked back into the light was not the sprawling, fragmented trade that had gone underground in 1920. The dry years had thinned the field, rewarding the largest and most adaptable companies — the ones with the capital to wait, to diversify into syrups and tonics and near beer, to keep their brands alive through a decade of enforced silence. The industry that emerged was leaner, more consolidated, and far more sophisticated about the management of public perception than the one that had entered. It had survived the most aggressive attack any government had ever mounted against it, and it had learned, in the surviving, a lesson it would never forget: that the product in the bottle is replaceable, but the relationship with the public is everything, and a great brand can outlast even the Constitution of the United States.

Repeal did not simply restore the old order; it built a new one whose architecture still governs American drinking today. The Twenty-First Amendment handed enormous authority to the states, and most of them erected a so-called three-tier system that legally separated the producers of alcohol from its distributors and its retailers, each tier licensed and taxed and forbidden from owning the others. The system was designed to prevent the abuses that had helped fuel the temperance backlash in the first place — chief among them the pre-Prohibition "tied house," the saloon owned outright by a brewery and pushed to sell as much of its product as aggressively as possible. By law, the new order would keep the maker at arm's length from the point of sale.

Yet the deeper legacy of Prohibition was not regulatory but cultural and criminal. The vast bootlegging networks that had grown rich in the dry years did not evaporate when the liquor became legal; the organized-crime syndicates simply redirected their capital and their muscle into other enterprises, leaving a mark on American life that long outlasted the Eighteenth Amendment. And the consolidated, sophisticated industry that emerged from the dry years was a formidable thing — leaner than before, schooled in the management of public opinion, and acutely aware of how close it had come to extinction. It would spend the next several decades making sure it never came that close again, not by making a better product, but by becoming ever more skilled at the one thing Prohibition had proven decisive: the management of how the public felt about it.

When public opinion finally turned, the Twenty-First Amendment repealed Prohibition in December 1933, and the industry came roaring back. But it returned to a permanently altered landscape. The Twenty-First Amendment handed the states broad authority to regulate and tax alcohol, and although the rules varied from place to place, one change held nearly everywhere: advertisers now had to be more truthful. The wild, pre-Prohibition claims about alcohol's health benefits — the tonics, the cure-alls, the medicinal whiskeys — were reined in, and advertisers were now required to be more truthful about what they were selling.

But honesty about a product's ingredients is a thin kind of regulation, and the years after repeal made that clear. People had hoped that a sober legal framework might finally tame alcohol's harms and nudge the country toward responsible drinking. Instead it became undeniable that alcohol was a genuine public-health problem in its own right — a driver of disease, addiction, underage drinking, and, as the automobile spread, of death on the roads. The law could police what a label claimed. It could do almost nothing about the far more powerful business of what a brand could be made to mean, and it was in that gap between literal truth and emotional suggestion that the modern alcohol industry would build its empire. People hoped the new regulatory framework would encourage responsible drinking and limit alcohol's harms. It turned out to be far more complicated than that, because the framework governed what the industry could literally claim about its products — not the vastly more powerful business of what the industry could make those products mean. And in the decades after repeal, the industry would become extraordinarily good at meaning.
"""

CHAPTER_5 = """
If the alcohol industry has a favorite customer, it is the traumatized one. That is an ugly sentence, but the historical record bears it out with brutal consistency, and nowhere more clearly than in the century of war and economic collapse that ran through the first half of the twentieth century.

The First World War sent millions of young men into a landscape of industrialized horror — trenches half-full of mud and corpses, artillery barrages that lasted for days, a constant grinding threat that produced a wave of psychological injury the era called "shell shock" and we now recognize as post-traumatic stress disorder. Of the 4.7 million American soldiers who served, many came home from the trenches profoundly damaged. In the teeth of that trauma, alcohol was an obvious and available refuge, and there is good reason to think it was used heavily. Some accounts suggest officers themselves supplied it, to steady nerves and stir fighting spirit — an echo, across nearly three thousand years, of Agamemnon pouring wine to ready his army for Troy.

None of this was new; only its scale was. The link between soldiers and alcohol is one of the best-documented patterns in human history, running from the wine rations of the Roman legions through the grog of the age of sail to the cigarettes-and-beer of the twentieth-century front. What the First World War added was industrial magnitude. Never before had so many men been exposed to such concentrated horror for so long, and never before had so many returned home carrying invisible wounds that medicine barely understood and society barely acknowledged. The result was a generation primed, by trauma and by habit, to drink — and an industry perfectly positioned to meet them. The pattern that would define the next half-century was set: mass suffering created mass need, and the alcohol trade was always there to supply it, draped in the language of comfort and courage.

[[IMG:13]]

The veterans of that war returned to an America that was entering Prohibition, which meant that any soldier who wished to keep drinking had to do so through illegal channels. Advertisers across the consumer economy rushed to capture the veteran market, frequently wrapping their products in nationalism and the imagery of the war. Anheuser-Busch built legal marketing campaigns for its near beer aimed squarely at veterans, using patriotic appeals to a sense of duty and sacrifice. The campaigns did something subtle and lasting: they kept beer culturally omnipresent even under Prohibition, and they cemented beer-drinking as part of the iconography of the American veteran.

The irony was sharp. These were men who had been told they fought to defend the American way of life, and they came home to find that the law had just declared one of that life's oldest pleasures a crime. To keep drinking, a veteran now had to break the law, and many did. Meanwhile the brands that could not legally sell him a full-strength beer kept his loyalty warm with patriotic near-beer campaigns, banking the relationship for the day repeal would come. It was the same lesson the breweries had learned about malt syrup, applied to human sentiment: keep the bond alive through the dry years, wrap it in the flag, and the customer will be waiting when the taps run freely again.

[[IMG:14]]

Then came the Great Depression, and with it a different kind of mass suffering — and a different kind of opportunity. The economic collapse that began in 1929 threw millions of Americans into unemployment, foreclosure, and a shattering loss of identity, and historians have documented the predictable rise in drinking as people tried to cope. The industry's response was not sympathy but calibration. Beer and spirits advertising of the 1930s was saturated with images of rugged, self-reliant masculinity — cowboys, frontiersmen, men of action and command. The strategy, historians argue, was aimed precisely at men whose sense of manhood had been gutted by their inability to provide, offering alcohol as a way to feel powerful and certain again.

It is worth being precise about what this means, because it is the second of the industry's enduring habits caught in the act. These advertisements did not simply sell beer to men who happened to be suffering. They were calibrated to the suffering itself — designed to find the specific wound of Depression-era masculinity, the humiliation of the breadwinner who could no longer provide, and to offer the product as a salve for exactly that wound. The cowboy and the frontiersman were not random images of rugged appeal; they were fantasies of competence and command sold to men who felt they had lost both. This is the difference between advertising to a vulnerable population and exploiting its vulnerability, and the alcohol industry, then as in every era, knew precisely which one it was doing.

[[IMG:15]]

[[IMG:16]]

Women were targeted too, but along a different axis. Wine advertising presented drinking as a marker of sophistication and refinement, leaning on ideals of beauty and aesthetic taste, and aimed at women who were commanding a growing share of household spending. From this Depression-era split emerged a gendered template that has proven astonishingly durable: beer and hard liquor coded masculine, wine coded feminine — a division that still shapes how alcohol is sold in the twenty-first century.

Walk down the aisle of any modern store and the template is still there, almost unchanged. The beer is sold with sports and trucks and friendship among men; the wine is sold with candlelight and self-care and the language of indulgence aimed at women; the hard seltzer and the sweetened spirit chase younger drinkers along the same old gendered grooves. These categories feel natural, even obvious, which is exactly the measure of the marketing's success. A genuinely natural division would not have needed ninety years of relentless advertising to establish and maintain it. What feels like common sense about who drinks what is, in large part, a Depression-era sales strategy that simply never stopped running — a portrayal so thoroughly absorbed that we mistake it for a fact about ourselves.

[[IMG:17]]

Then the Second World War handed the industry its greatest opening since repeal. More than sixty million people died in the war and the genocide and famine surrounding it, and fear and anxiety ran high across the entire globe. Demand for alcohol surged past even pre-Prohibition levels as soldiers, civilians, and war workers drank to manage their dread. Alcohol also became a major source of government revenue through wartime taxes. The war cut both ways for the trade — grain, sugar, and fuel were rationed, and the government restricted the use of food crops for distilling — but marketing never paused. Producers urged Americans toward frugality, wrapped their brands in the flag, and suggested that buying the product was itself a way to support the war effort.

The wartime situation forced a delicate balancing act, and the industry performed it with practiced skill. On one side, the resources of brewing and distilling — grain, sugar, fuel, even the industrial alcohol the war demanded for munitions and synthetic rubber — were suddenly precious, and the government restricted their use for drink. Production capacity rose and fell with the shifting needs of the war machine. On the other side, demand had never been higher, as soldiers, civilians, and the millions who flooded into war work all reached for relief from fear and exhaustion. Caught between scarce supply and surging demand, the industry leaned harder than ever on the one thing it never lacked: the message. If a brand could not always promise abundance, it could promise meaning — it could attach its name to sacrifice, to thrift, to the boys overseas, to the very idea of the nation at war.

So the advertisements of the period sold not just whiskey or beer but membership in the war effort. They told the drinker that frugality was patriotic and that the dollars not spent should go to war bonds; they wrapped their labels in flags and eagles and the iconography of duty; they implied that to choose this brand was, somehow, to do one's part. It was a remarkable feat of repositioning — turning a moment of rationing and restraint into an opportunity to deepen the public's emotional bond with the product. And it worked, because it always works: when a brand attaches itself to something a frightened people already hold sacred, it borrows that sanctity at no cost to itself. The whiskey advertisement urging you to buy war bonds was performing the same trick as the wine crier eight centuries before, simply on a national stage and in a moment of maximum vulnerability.

[[IMG:18]]

[[IMG:19]]

Over sixteen million Americans served, and many came home carrying the same psychological wounds as their fathers' generation. The drinking culture of that war became part of its mythology: jubilant troops celebrating with abundant French wine after the Normandy landings, U.S. infantry drawing daily beer rations in the Pacific, British sailors lining up for the rum ration that the Royal Navy would maintain into the 1970s.

The military's long entanglement with alcohol is worth pausing on, because it reveals how deeply the bond between drinking and soldiering runs, and how useful that bond has been to the trade. The British naval rum ration — the daily "tot" — had been an institution for centuries, a fixture of shipboard life so entrenched that its eventual abolition in 1970 was mourned as the end of an era. Armies across history issued drink to steady nerves before battle and to celebrate after it, to dull the cold and the fear and the boredom that fill the vast stretches of a soldier's life between moments of terror. Officers understood it as a tool of morale and control; soldiers understood it as one of the few reliable comforts available in unbearable circumstances. Either way, the young men who survived came home having learned, in the most intense conditions imaginable, to reach for a drink when life became hard to bear.

This was a lesson the industry was delighted to reinforce, and the imagery of the world wars handed it a gift it would draw on for generations. Photographs of GIs toasting their survival in liberated Paris, of infantrymen sharing a beer ration in a Pacific clearing, of sailors lined up for their tot, became part of the visual mythology of victory itself — alcohol fused, in the national memory, with courage, camaraderie, and freedom hard-won. Advertisers reached for that fusion immediately and never really let go. The beer in the veteran's hand was not just a beer; it was a symbol of the brotherhood of the front and the sweetness of coming home alive. That association, forged in the most consequential events of the century, would prove among the most durable the industry ever made, and it primed a whole generation of returning servicemen — many of them already carrying the invisible wounds of combat — to drink, and to keep drinking, in numbers that public-health researchers would be documenting for decades.

[[IMG:20]]

[[IMG:21]]

[[IMG:22]]

It is well documented that the veterans of the Second World War, like those of the First, suffered disproportionately high rates of alcoholism. War after war, the same machinery turned: trauma created need, availability met it, and marketing wrapped the whole thing in honor and belonging. And when peace finally came, it brought a long economic boom and a flood of leisure that pushed per-capita consumption steadily upward through the 1950s and 1960s and into the 1970s.

The postwar surge was not simply a matter of prosperity, though prosperity certainly fueled it. It was the cumulative aftershock of three world-historical traumas stacked one upon another within a single lifetime: the First World War, the Great Depression, and the Second World War. A generation had come of age inside a near-continuous emergency, and the alcohol industry had been there at every stage — supplying the trenches, exploiting the breadlines, wrapping itself in the flag of two world wars. By the time the shooting stopped for good, drinking had been thoroughly braided into the American idea of relief, reward, and return to normal life. The veteran drank to quiet what he had seen. The suburb drank to celebrate what it had built. The advertising of the era met both impulses with a flood of catchy, aspirational campaigns that normalized heavy drinking as the natural accompaniment to the good life everyone had supposedly fought for.

What made this moment so consequential for the rest of our story is that it set the stage on which the next medium would perform. A prosperous, suburban, leisure-rich nation, primed by decades of hardship to see alcohol as comfort and celebration alike, was settling down each evening in front of a brand-new appliance that beamed sound and image directly into the living room. The audience was ready. The associations were already in place. All that remained was for the most powerful advertising medium yet invented to arrive and supercharge them — and it was, at that very moment, flickering to life in millions of American homes. A newly prosperous, newly suburban America was primed to drink — and a new medium was arriving in its living rooms, one that would prove more powerful at selling alcohol than anything that had come before.
"""

CHAPTER_6 = """
In the early 1950s a new appliance settled into the American living room and rearranged the culture around it. Television combined the visual punch of print with the emotional intimacy of radio and delivered both directly into the home, night after night, to an audience that sat still and paid attention. The advertising industry exploded to meet it, in a burst of creativity and spending so intense that historians still call the period the Golden Age of Advertising — the world later mythologized, half a century on, in the figure of the fictional ad man Don Draper.

The postwar economy was booming and hungry for goods, and the agencies that had honed their craft on wartime propaganda turned their full attention to selling consumer products as the furniture of a modern, aspirational life. Television commercials showed happy families, sleek cars, gleaming kitchens. And they leaned on a new weapon: the slogan, the jingle, the line so catchy it lodged in the memory and would not leave. "When you say Budweiser, you've said it all." The goal was no longer merely to inform people that a product existed. It was to fuse the product to a feeling and play that fusion into millions of minds at once.

Television was uniquely suited to that work, and the alcohol industry understood why almost immediately. Earlier media had each offered one channel to the senses: the broadside and the magazine were silent and still; radio was sound without image. Television fused them. It could show a frosted glass catching the light and a crowd of friends laughing around it, and it could do so with a voice and a jingle layered on top, all delivered into the living room at the hour when families gathered and their guard was down. And it reached almost everyone. Within a single decade, the television set went from a novelty to a fixture in nearly every American household, the new hearth around which domestic life arranged itself. For an industry whose entire art was the association of a product with a feeling, no medium had ever offered so direct a path from the screen to the psyche.

Beer brands seized that path with particular zeal, and the lifestyle commercial became their signature form. The point of such a spot was never really the beer; the beer appeared for only a few seconds. The point was the world around it — the cookout, the ballgame, the camaraderie, the reward at the end of an honest day's labor — a world the viewer was invited to enter by way of the can. The slogans did the binding work, compressing that whole world into a phrase catchy enough to survive the trip from the screen into memory and out again at the moment of purchase. "When you say Budweiser, you've said it all" does not describe a beer. It describes a feeling of completeness, of having arrived, and it staples that feeling to a label. Repeated nightly across a nation of new television sets, such associations did not just sell beer. They taught a whole postwar generation what beer was supposed to mean.

Into this golden age the alcohol industry walked with a peculiar, self-imposed limp — one that turned out to be a brilliant long-term strategy. In 1948 the distilled-spirits trade adopted a voluntary code that kept hard liquor off television and radio entirely. No law required it; the industry simply chose it. The decision was framed as corporate responsibility, and for decades it bought the spirits business an invaluable reputation for restraint — a supposed track record of policing itself that it would invoke, again and again, to fend off government regulators in the decades to come.

With the spirits houses standing voluntarily aside, the television age became, overwhelmingly, the age of beer. The brewers poured their fortunes into the new medium and learned its grammar faster than almost any other industry, and the results became part of the American soundtrack. The slogan and the jingle were their chosen weapons, because a tune is even stickier than a sentence; it installs itself in the memory and replays unbidden, carrying the brand along with it. A campaign from the 1950s promised that "where there's life, there's Bud," stapling a beer to the very idea of vitality and good living. Decades later, the same brand would conquer the culture again with nothing more than a group of friends drawling a single syllable at one another over the phone — a catchphrase so infectious it escaped the commercial entirely and entered everyday speech, which is the holy grail of advertising: a pitch that the audience carries and repeats for free.

[[IMG:23]]

Look at how little such advertisements actually concern the product. The beer is barely present; what fills the screen is a feeling — leisure, friendship, belonging, the easy warmth of being among your people at the end of the day. That was the genius and the danger of television. It did not argue that one beer tasted better than another, a claim a viewer might evaluate and resist. It simply showed, night after night, a desirable world with a particular beer sitting quietly at its center, until the brand and the warmth grew indistinguishable in the mind. Repeated across a nation of new television sets, in the relaxed and receptive hours of the evening, these gentle, story-shaped pitches did work that no broadside or billboard ever could. And because the spirits houses had voluntarily ceded the medium, it was beer that did this teaching — beer that owned the evenings and the ball games and the jingles, beer that became the default American drink in the most powerful advertising medium ever built.

The voluntary spirits ban is worth dwelling on, because it reveals how industry self-regulation actually functions. In a self-regulated system, the industry writes its own rules and faces no real penalty for breaking them. And the spirits code did almost nothing to reduce the volume of alcohol advertising on television. It simply ceded the entire airwave to beer, which rushed in to fill the vacuum with a torrent of commercials. The total quantity of televised alcohol marketing did not fall. Only its labeling changed. This is the recurring magic trick of self-regulation: it generates the public appearance of caution while leaving the underlying business — the saturation of the culture with drink — entirely intact. Decades of research would later confirm that even with such voluntary codes in place, the harms they claimed to prevent, including underage drinking, remained stubbornly common.

The Golden Age of alcohol advertising coincided, not by accident, with a rising national reckoning about alcohol abuse. America was living through the long aftershock of three traumas stacked on top of one another — the First World War, the Great Depression, and the Second World War — and a quieter counter-current was building against the noise of the commercials. In the postwar years, scholarly efforts such as the Yale Center of Alcohol Studies and grassroots movements such as Alcoholics Anonymous advanced a disease model of alcoholism: the idea that addiction arose from biological, psychological, and social mechanisms and could be treated, rather than being merely a moral failing to be scolded away. The federal government began to take the problem seriously. Under the Nixon administration, the Comprehensive Alcohol Abuse and Alcoholism Prevention, Treatment, and Rehabilitation Act of 1970 created the National Institute on Alcohol Abuse and Alcoholism, cementing alcoholism as a national concern worthy of systematic research and treatment.

This marked a genuine shift in how the country thought about drink. For most of the long history told in this book, the response to alcohol's harms had focused on the supply side — on prohibition, regulation, taxation, the control of who could sell what to whom. The disease model pointed somewhere new: toward the demand side, toward research, education, and treatment for the drinker as a person with a treatable condition rather than a moral failing to be punished or a market to be served. It was a more humane and more accurate way of understanding addiction, and it slowly reshaped medicine, policy, and public attitudes.

The new understanding grew from several directions at once. Researchers at institutions such as the Yale Center of Alcohol Studies began to investigate alcoholism with scientific rigor, treating it as a condition to be understood rather than a vice to be condemned. The fellowship of Alcoholics Anonymous, founded in the 1930s, carried a parallel message out of the laboratories and into church basements across the country, insisting that the alcoholic was a sick person who could recover, not a sinner who simply lacked willpower. Together these currents marked a genuine break with the moralizing of the temperance era — a recognition that drinking sat at the intersection of biology, psychology, and social circumstance, and that punishment and prohibition had never been equal to it.

But here the deeper irony of the postwar decades comes into focus. The very years in which America built its institutions for understanding and treating alcohol problems were the years in which the industry achieved its most complete saturation of the culture. As doctors and researchers worked to frame heavy drinking as a treatable illness, advertisers worked, with far larger budgets, to frame it as the natural reward of a life well lived. The science and the salesmanship were pulling in opposite directions, and the salesmanship had the money, the airwaves, and a head start of several thousand years. It is a tension that runs straight through to the present: a society that funds research into alcohol's harms while bathing itself, day and night, in messages designed to make those harms feel like fun. The public-health understanding of alcohol has never lacked for evidence. What it has always lacked is anything close to the industry's reach.

But here lies one of the quiet tragedies of the period, and a pattern worth naming. Even as the public-health understanding of alcohol deepened, the industry's cultural entrenchment deepened faster. The very decades in which America was building institutions to study and treat alcoholism were the decades in which alcohol advertising achieved its most total saturation of the culture. The science was catching up to the harm; the marketing was racing further ahead. And the industry, reading the room, could see what the future might hold. When cigarette advertising was driven off the airwaves in 1971, the alcohol industry watched a fellow vice industry get hauled before the regulators and stripped of its broadcast access — and it began, methodically, to build forms of advertising that no broadcast ban could ever touch. If regulators might someday close the front door, the industry would learn to enter through the story itself.

That same period produced a telling contrast. In 1971, cigarette advertising was banned from American radio and television. Alcohol advertising was not. The discrepancy captured the cultural common sense of the moment: alcohol, however harmful, was still regarded as acceptable in moderation in a way tobacco no longer was. But the industry seems to have read the writing on the wall. Sensing that regulators might eventually come for alcohol the way they had come for tobacco, the trade began developing new advertising techniques — methods designed to thread alcohol so deeply and so quietly into the fabric of entertainment and everyday life that no broadcast ban could ever reach them. If the front door of overt advertising might someday close, the industry would simply learn to come in through the walls.
"""

CHAPTER_7 = """
The most effective advertising is the kind that does not look like advertising at all. As the twentieth century wore on, the alcohol industry mastered three techniques built on exactly that principle — product placement, sponsorship, and celebrity endorsement — each one designed to dissolve the line between the commercial and the cultural until the audience could no longer tell where the entertainment ended and the sales pitch began.

Product placement is the art of slipping a branded product into a film or television show so that it rides along inside the story. The technique had simmered for years, but it became a phenomenon after the 1982 film E.T., which famously used a trail of Reese's Pieces at a pivotal moment and sent sales of the candy soaring. The lesson was not lost on the alcohol industry, which had in fact been practicing the craft for decades. The most storied example is James Bond. When researchers systematically analyzed six decades of Bond films, they counted 109 distinct alcohol product-placement events — an average of four and a half per film — featuring brands like Smirnoff and Heineken. Bond's famous vodka martini was, among other things, one of the longest-running advertising campaigns in cinema history.

What makes product placement so powerful is precisely that it switches off the viewer's defenses. People watching a commercial know they are being sold to, and they brace accordingly — a skepticism that decades of advertising have trained into us. But a drink in the hand of a beloved character arrives with no such warning label. It is part of the story, and we have agreed, by sitting down to watch, to believe in the story. The brand slips past the gate disguised as a narrative detail, and the glamour, competence, or warmth of the character transfers quietly onto the product. When Bond orders his vodka with that famous instruction, the audience is not weighing a purchase; it is admiring a man, and the admiration sticks to the bottle on its way past.

There was a second, colder genius to the technique, and the timing was no accident. As the previous chapter described, the alcohol industry could see regulators tightening their grip on broadcast advertising. Product placement offered an elegant escape. A brand woven into the fabric of a film or a sitcom was not a broadcast advertisement in any sense a regulator could easily define, and so it sailed clean past rules built for an older kind of selling. The industry had found a way to advertise inside the very entertainment that the advertising rules were meant to protect — and it spread the practice widely, from Bond's martinis to the cans of a particular beer that seemed to materialize in every other film and television series of the era, present in the story but rarely noticed as a pitch.

[[IMG:24]]

Bond was only the most glamorous case. Budweiser placed its products across an enormous range of films and television series, from Top Gun to Friends, until a can of a particular beer in a character's hand became a quiet, constant background hum of branding — present in the story but rarely noticed as a pitch.

Sponsorship works on a closely related principle: a company pays to attach its name to something people already love — a team, a festival, a performer — in the hope that the audience's affection for the thing will rub off onto the brand. If you are having the time of your life at a stadium and an alcohol brand's name is everywhere around you, your good feelings and that brand begin, subtly, to merge. From the 1970s onward, alcohol brands chased these associations relentlessly. Miller Lite's long sponsorship of the Dallas Cowboys, beginning in 1991, is among the most recognizable in American sports, binding a beer to one of the country's most beloved franchises and, through it, to entire generations of fans.

[[IMG:25]]

The genius of such deals is their reach across time. A child who grows up watching a team wrapped in a particular beer's branding inherits that association and may carry it for life. And the consequences are measurable: research has linked alcohol's sponsorship of sports to higher drinking levels among young people and to hazardous drinking among adults. The good feeling is the product; the beer is just what the good feeling is attached to.

From the 1970s onward, the alcohol industry pursued these associations across every arena where people gathered to feel good together — stadiums and arenas, music festivals and concert tours, cultural celebrations of every kind. The logic was always the same. A person who is happy, excited, and surrounded by others having a wonderful time is in a uniquely receptive state, and a brand present in that moment gets folded into the memory of the joy. Repeat the pairing across a season, a career of fandom, a lifetime of summers, and the brand stops being a sponsor and becomes part of the experience itself — inseparable, in the fan's mind, from the team or the music or the festival that they love.

The consequences are not merely sentimental; they are measurable, and they fall hardest on the young. Research has consistently linked alcohol's sponsorship of sport to higher levels of drinking among young people and to hazardous drinking among adults, precisely because the technique works the way it is designed to: it builds an emotional bond early and reinforces it for years. A child who spends a decade of Sundays watching a beloved team play in a stadium wrapped in a single beer's branding does not merely learn that the beer exists. The child learns that the beer belongs to something they love, and may carry that lesson, unexamined, for the rest of their life. The sponsorship deal is, in effect, an investment in a generation's affections — bought cheaply now, and collected on for decades.

Sport was only the beginning. From the 1970s onward, alcohol brands wove themselves into music festivals, concert tours, and cultural celebrations of every description, chasing the same prize wherever people gathered to feel joy together. A beer's name on a summer festival, a spirit's logo behind a concert stage, a brand's tent at a street celebration — each purchase rented a slice of someone's happiest memory. And the brands learned to court the tastemakers of each scene, the musicians and artists and local heroes whose approval signaled belonging, so that the drink arrived not as an intrusion on the culture but as a native part of it. The strategy reached its natural endpoint in the festival built and named by a brand itself, an entire event existing to be a branded good time — the logical conclusion of a logic that began when a medieval tavern keeper realized that the right sign could turn a building into a destination.

What ties all of these techniques together — placement, sponsorship, endorsement — is that none of them looks like an advertisement, and that is the entire point. By the end of the twentieth century the alcohol industry had learned to stop interrupting the culture and start inhabiting it: to live inside the films and the games and the music and the famous faces that people sought out for pleasure, rather than buying the gaps in between. It was a quieter kind of selling, and a far more powerful one, and it left the audience with no obvious moment at which to raise its guard. The pitch was no longer a thing that happened to you. It was woven into the things you loved.

Seen from the vantage of the present, these late-century techniques look like a dress rehearsal. Product placement taught the industry to hide inside a story; sponsorship taught it to buy its way into joy; celebrity endorsement taught it to borrow a trusted face. Each dissolved a little more of the boundary between advertising and ordinary life, and each trained the public to accept a world in which the selling was always present but never announced. What none of them could do, for all their power, was reach the individual directly. The film played to everyone in the theater; the stadium banner greeted every fan alike; the celebrity smiled at the whole audience at once. The techniques were intimate in feeling but mass in fact.

That was the last barrier, and it was about to fall. The same decades that perfected the invisible advertisement also produced the technologies that would make it personal — first the website and the email, then the social feed — until the industry could take everything it had learned about hiding inside the things people love and aim it, at last, at one person at a time. The celebrity endorsement would become the influencer's post. The product placement would become the sponsored content sliding through the feed. The sponsorship would become the brand that talks to you like a friend. Everything the twentieth century taught the industry about inhabiting the culture, the twenty-first would let it whisper directly into your ear. To understand how, we have to follow the alcohol industry online.

Celebrity endorsement completes the trio, lending a brand a famous person's charisma, glamour, and trust. Alcohol brands have leaned on celebrity since the early twentieth century, but the practice accelerated sharply from the 1970s on, and a few partnerships show its raw power. When the rapper and entrepreneur Sean "Diddy" Combs took on Ciroc vodka in the mid-2000s, the brand's sales rose meteorically, and it penetrated young Black American consumers — a demographic vodka marketers had long struggled to reach. When the actor George Clooney co-founded the tequila brand Casamigos in 2013, the brand and the man became almost a single entity; Clooney's persona and lifestyle were the product, and the alcohol came along for the ride.

The Ciroc story repays a closer look, because it shows celebrity endorsement working as cultural alchemy rather than mere advertising. Vodka had long struggled to reach young Black consumers, and no amount of conventional marketing had cracked the problem. What cracked it was authenticity, or the appearance of it: when Sean Combs took the brand on not as a paid spokesman reading lines but as a partner with a real stake, weaving the vodka into the imagery and aspiration of hip-hop culture, the brand stopped feeling like an outsider's pitch and started feeling like an insider's choice. The endorsement worked precisely because it did not look like an endorsement. It looked like belonging. That is the lever celebrity has always pulled, from the wine criers of medieval France to the matinee idols of the mid-century: it borrows a trusted, admired voice and lets the product ride in on the trust.

[[IMG:26]]

And the trust is bankable, sometimes to an almost absurd degree. The proof arrived a few years after Casamigos launched, when it sold to the drinks giant Diageo in a deal valued at roughly a billion dollars — an astonishing sum for a brand only a few years old, and a precise market valuation of a fused celebrity-and-bottle identity. What Diageo bought was not a distillery or a recipe; those are cheap. It bought a feeling that millions of people already had about a famous man, transferred onto a product, and it paid a billion dollars for the transfer. Decades of research confirm what that price tag implies. Product placement, sponsorship, and celebrity endorsement are all genuinely effective at building awareness and driving consumption, and their effect is strongest among the young, who are still assembling their identities out of the idols and images around them. The techniques work because they smuggle the sales pitch inside the things we love and admire, where our defenses do not think to look.

By the close of the twentieth century, alcohol had achieved something beyond mere ubiquity. It was no longer simply advertised to Americans. It was woven into their stories, their teams, their idols — ambient, constant, and almost invisible. And then the ground shifted again, because a new kind of technology was about to arrive, one that would let the industry not merely surround its audience but speak to each member of it by name.
"""


CHAPTER_8 = """
In the mid-1990s a new medium began to command the public's attention, and the alcohol industry, true to a habit by then nine thousand years old, was among the first to see its potential. The earliest internet advertising was crude — static banner ads pinned to web pages, a digital echo of the painted billboard, mostly text and low-resolution images. But the principle was familiar, and the pioneers were familiar too. Budweiser was among the first major alcohol brands to buy banner space on popular sites, planting its flag on the new frontier exactly as the trade had planted broadsides on city walls in the 1830s.

The early web was a strange and exhilarating place for an advertiser, and the alcohol industry approached it with the same opportunism it had brought to every medium before. No one yet knew the rules. The technology was crude, the audiences were small but growing explosively, and the conventions that would later govern online advertising had not been invented. Into this open frontier the brands rushed to experiment — buying banners, registering domain names, building the first clumsy websites — not because anyone had proven these things worked, but because the industry's long experience had taught it a simple lesson: get there first. The medium that looks like a toy today is the dominant force tomorrow, and the brand that learned its grammar early would own it when it mattered.

What the early internet promised, even in its primitive form, was a quality no previous medium could offer: interactivity. A billboard could only be seen; a commercial could only be watched. But a website could be explored, clicked through, lingered in. The audience was no longer purely passive, and that changed the nature of the relationship. An advertiser could now invite the customer into an experience and watch, in increasing detail, how the customer responded — which pages held attention, which links were followed, which offers were taken. Each interaction left a trace, and the traces could be gathered and studied. The crude banner ad of the 1990s was the thin edge of a very large wedge, the first hint that the new medium would not merely show advertising to people but learn from them in return — and that the learning would eventually become the most valuable thing about it.

The static banner was only a beginning. As bandwidth and browsers improved, the industry moved toward richer forms, and the most important of these was the brand's own website — a permanent digital home, open at all hours, that doubled as advertisement and shrine. Jack Daniel's launched its site in the mid-1990s and never looked back, filling it with the lore of its whiskey-making, cocktail recipes, and folksy historical detail, all of it quietly telling a single compelling story with the brand at its center. A website was something no billboard or television spot could be: a place a customer chose to visit, and lingered in, on purpose.

This was a genuine inversion of the old relationship between seller and sold-to. For nine thousand years, advertising had been something done to people — a sign thrust before the eye, a commercial that interrupted the program, a mailer that arrived unbidden. The audience was a target, and the advertiser's central problem was capturing attention that was always trying to wander off. The brand website flipped the arrangement. Now the customer came to the advertiser, voluntarily, and stayed a while, and the brand's task was no longer to interrupt but to reward the visit — with recipes, with lore, with the carefully tended story of the product. A visitor who chooses to spend twenty minutes browsing a distillery's history is in a state no billboard could ever produce: relaxed, curious, and already half in love. The site was not an advertisement that a person endured. It was an experience a person sought out, which made it a far more powerful instrument of persuasion than anything that had to shout for attention first.

And it pointed toward the future in another way. A website let a brand behave like a publisher — generating a steady stream of content designed to draw people in and keep them coming back. The alcohol companies were learning to be media outlets in their own right, producing the entertainment rather than merely buying space inside someone else's. It was a small step from there to the logic of the social platform, where a brand would post and comment and joke and engage exactly like a person, and the line between a company talking to you and a friend talking to you would finally dissolve altogether.

[[IMG:27]]

Two further techniques pushed personalization further than any earlier medium had managed. The first was programmatic advertising — paying to place dynamic, interactive ads on the specific sites a target demographic was known to frequent, with each ad linking straight back to the product. Unlike a banner shouting at everyone, programmatic advertising aimed, and it linked the moment of desire directly to the means of purchase. The second was the email campaign. As email became a primary channel of everyday communication, alcohol companies began collecting addresses — through their own sites, through online contests and promotions — and using them to send messages tailored to the individual recipient: their name, their location, their tastes. For the first time in the long history of the trade, a brand could speak to a customer not as part of a crowd but as a named individual, in a private channel the customer checked many times a day.

The engine beneath all of this was data. Every visit to a website, every contest entered, every address surrendered for a coupon, added a line to a growing dossier, and the dossiers made a new kind of advertising possible. The old media had always been blunt instruments. A billboard could be placed on a busy road, a commercial slotted into a popular program, a magazine ad aimed at a known readership — but in every case the advertiser was guessing at a crowd and hoping the right people were in it. Digital advertising replaced the guess with a record. A brand could now know, with increasing precision, who you were: your age, your location, the sites you frequented, the things you had bought, the hours you were awake. And it could use that knowledge to decide which message to show you, and when.

For an industry whose oldest skill was matching the pitch to the prey — the cowboy ad for the humiliated breadwinner, the refined wine ad for the aspirant housewife — this was a kind of homecoming. The Depression-era marketer had to intuit the vulnerabilities of a demographic and broadcast a single guess at all of them. The digital marketer could segment that demographic into thousands of individuals and tailor a different appeal to each. The instinct was ancient; only the precision was new, and the precision changed everything.

That shift — from broadcasting at a mass to whispering to an individual — was the quiet turning point in the whole history of alcohol marketing. For nine thousand years the industry had been limited by the nature of its media. The tavern sign reached whoever walked past. The newspaper reached whoever bought it. The television commercial reached whoever was watching, all of them receiving the same message at the same moment. Digital media broke that constraint. Now the message could bend to fit the receiver. The industry could not only surround its audience; it could divide that audience into ever-smaller segments and address each one in its own language. The reach of alcohol advertising expanded exponentially, and its aim grew correspondingly precise.

It is worth pausing to appreciate how complete a reversal this represented. For nine thousand years the fundamental limitation of alcohol marketing had been the medium's indifference to the individual. The tavern sign greeted the saint and the drunkard with the same painted face. The newspaper carried the same advertisement to every reader. Even television, the most powerful persuasion machine yet built, spoke to the whole audience at once and trusted the law of averages to do the rest. The drinker had always been, to the advertiser, a member of a mass. Digital media dissolved the mass into its individual parts and let the brand address each one by name, in private, in a channel that person carried with them and consulted compulsively. The reach grew while the aim narrowed — a combination no earlier medium had ever managed, and a foretaste of what would arrive when these tools were fused with a platform built not for documents or email but for human relationships themselves.

And yet, for all its power, the early digital era still kept a recognizable boundary between the advertiser and the advertised-to. A banner was a banner; a brand's website was obviously the brand's; an email from a distillery was clearly a message from a company that wanted your money. You knew, more or less, when you were being sold to. The next medium would erase even that. It would take the industry's oldest trick — the wine crier's borrowed voice, the placement that hides inside a story — and perfect it, by building a place where commercial messages could be made to look exactly like the words of a friend.
"""

CHAPTER_9 = """
Social media changed everything again, and it did so by collapsing the last distinction the industry had left to overcome: the line between an advertisement and a personal message from someone you trust.

Social-media marketing uses networks like Facebook, Instagram, and the platform once called Twitter to deliver tailored promotions straight into a user's feed. The ads take many forms — still images, videos, scrolling carousels — but the most insidious of them is the sponsored post engineered to blend seamlessly into the ordinary content around it, so that a commercial message sits in the feed looking just like a photo from a cousin or a friend. A close cousin of this technique is influencer marketing, which is essentially celebrity endorsement democratized: people with large followings, who may be famous nowhere outside their platforms, are paid to feature a product to an audience that regards them as peers rather than pitchmen. The whole apparatus is built to do one thing supremely well — to segment an audience by age, interest, behavior, and mood, and then to reach each segment with a message designed to feel native to it.

Two features of this medium make it unlike anything that came before, and both work in the industry's favor. The first is the depth of its targeting. A social platform does not merely know your age and your town, as a website's cookie might; it knows your friends, your interests, the things that make you linger and the things that make you scroll past, and from these it infers your moods and your moments. It can distinguish the user celebrating a promotion from the one nursing a breakup, and it can serve each a different message at the instant they are most receptive. The second feature is the disguise. On a social feed, a paid message does not arrive in a box marked "advertisement," set apart from the real content. It arrives in the stream itself, shaped to look exactly like the posts of friends and family that surround it — same format, same tone, same intimate register. The most sophisticated marketing in human history is engineered, above all, to not look like marketing at all.

The alcohol industry has been working this medium effectively since the late 2000s, and its conduct reveals how little the underlying logic has changed since the wine criers of medieval France. Consider the case of Smirnoff's Facebook strategy, exposed in an analysis of internal industry documents. The approach was deliberately designed so that the brand's posts would read as though a friend, not a corporation, were speaking — casual, warm, personal — and it actively cultivated user comments to amplify that effect. It worked. But the same documents revealed something darker: Smirnoff insiders were aware that nearly three-quarters of the brand's Facebook followers were under the legal drinking age.

[[IMG:28]]

That detail deserves to be sat with, because it is the entire problem in miniature. A brand engineered its marketing to look like friendship, deployed it on a platform saturated with minors, and understood — in its own internal accounting — that most of the people it was reaching were too young to legally buy what it sold. This is the structural failure of self-regulation, the same failure we met with the spirits industry's voluntary television code, now transplanted to a medium built for precisely the kind of boundary-blurring the industry had always craved. Decades of research have linked social-media alcohol marketing to the initiation, use, and abuse of alcohol, with the strongest effects, predictably, among young people.

To understand why social media is so potent a vehicle for all of this, it helps to see what these platforms actually sell. Their product is not the service they offer to users; it is the users' attention, packaged and sold to advertisers. Everything about the design — the endless scroll, the unpredictable rewards of likes and notifications, the feed tuned to show whatever keeps a person engaged — is engineered to capture and hold attention, because attention is the inventory the business exists to sell. An alcohol brand buying access to that attention is therefore buying into a machine optimized, down to the smallest detail, to keep people looking. The medium does the hard work of holding the audience; the advertiser need only supply the message.

Influencer marketing fits this machine perfectly, and it represents the wine crier's ancient trick brought to its logical extreme. The medieval crier borrowed the trust of a town's official voice; the modern influencer rents out the trust of a friend. Followers feel they know these figures — they have watched their daily lives, absorbed their tastes, come to regard them as intimates rather than broadcasters — and that felt intimacy is precisely what makes a product recommendation land. When an influencer holds up a particular drink, it does not register as an advertisement at all. It registers as a tip from someone you trust, which is the most persuasive form of marketing ever devised, and the hardest to defend against. The brand has not bought a billboard. It has bought a friendship, or a convincing imitation of one, and pointed it at whatever it is selling.

The deeper danger is the dissolution of the very category of "advertisement." When a college student posts a photo holding a particular craft beer, is that an advertisement? When an influencer mentions a tequila in a video, is that a paid endorsement or a genuine recommendation? When a brand's account posts a meme that thousands of users cheerfully share, who exactly is doing the advertising? The lines do not blur by accident. They are blurred on purpose, because a message that does not look like marketing slips past the defenses that people have learned to raise against marketing. The blurring is the product, and the blurring is profitable.

The danger is sharpest where the audience is youngest. Young people live on these platforms, and the same features that make social media so effective at selling — the intimacy, the disguise, the relentless personalization — make it nearly impossible to keep alcohol marketing away from those too young to drink. The Smirnoff case was not an aberration but a preview. When a platform is built to maximize engagement, and a brand is built to maximize the appeal of its product, their incentives align against the very protections that are supposed to shield minors. Age gates ask a birthdate and accept any answer. Content that is shared by users travels far beyond whatever audience a brand officially "targeted." And a marketing message engineered to look like a friend's post does not announce itself to a sixteen-year-old any more than to anyone else. The structural failure here is not a single company's bad behavior; it is a medium whose fundamental design rewards exactly the boundary-blurring the alcohol industry has wanted since the days of the wine crier — now operating at planetary scale, in the pockets of children.

What makes the Smirnoff case so damning is not that a brand reached underage users — on a platform saturated with minors, some of that is nearly inevitable — but that the analysis of internal industry documents showed the company understood what it was doing. The strategy to make corporate posts read like a friend's was deliberate. The cultivation of user comments to amplify the brand's reach was deliberate. And the awareness that a large majority of the brand's followers were below the legal drinking age was, the documents indicate, simply known — a fact in the business's possession, not a surprise it might have missed. This is the difference between an unfortunate side effect and a tolerated cost of doing business, and the documents suggest the latter.

It would be comforting to treat this as a single firm's lapse, but the history in this book argues otherwise. From the malt syrup sold with a wink to home brewers, to the spirits code that policed appearances while ceding the airwaves to beer, to the wartime advertisement that draped a bottle in the flag, the industry has shown again and again that its public posture of responsibility and its private conduct are managed as separate things. The social feed did not corrupt an otherwise scrupulous industry. It handed a very old industry a very powerful new tool, perfectly shaped to the boundary-blurring it had always practiced, and removed most of the friction that had previously limited it. The result is the most efficient machine for selling alcohol ever built, pointed at an audience that includes, by the industry's own knowledge, a great many people too young to legally drink.

There is one more feature of this medium that the rest of this book has been building toward. Social platforms are not merely channels; they are amplifiers that respond to crisis. The alcohol industry has a documented history of leaning hardest on people in moments of fear and upheaval — the trenches of the First World War, the despair of the Great Depression, the dread of the Second. Social media takes that ancient instinct and arms it with unprecedented reach and precision, delivering tailored messages to individuals in real time, calibrated to their mood and their moment. So it should not surprise us — though it should alarm us — what happened when the precise targeting of social media met the heightened vulnerability of a frightened, isolated population during the largest public-health crisis in a century. The industry's three oldest habits were about to converge, all at once, on a captive audience locked inside its own homes.
"""

CHAPTER_10 = """
Step back from the long chronicle of the preceding chapters and three patterns stand out from the noise — three habits so consistent across nine thousand years that they amount to a description of the industry's character.

The first is that the alcohol industry's success has always been bound to advances in communication technology. This is the thread that has run through every chapter of this book. From the hand-written manuals of Cato and Brunschwig to the painted signs of medieval taverns; from the broadsides, billboards, and color magazines of the industrial age to the television commercials of the postwar boom; from the brand websites and targeted emails of the early internet to the social feeds of the present day — the industry has been among the earliest and most aggressive adopters of every new medium humanity has invented. The consistency is so total that it stops looking like coincidence and starts looking like the defining trait of the enterprise.

The second pattern is darker. The alcohol industry has a long, documented history of reaching for the vulnerable. Roman commanders used wine to motivate and control the legions. Medieval peasants were a captive market, drinking for nutrition and escape in a cycle that deepened their poverty. The American industry was built on the profound injustice of slavery and the dispossession of Native peoples. And in more recent times the trade has preyed, again and again, on the traumatized and the desperate — war veterans, the unemployed, the grieving, the frightened. Whether by intention or by reflex, the industry has profited from human vulnerability with remarkable reliability.

The third pattern concerns endurance. The portrayals the alcohol industry creates do not fade when the campaign ends; they harden into culture and outlive everyone involved. The association of drink with soldiers and courage runs unbroken from the Iliad to the modern recruiting-age beer commercial. The use of alcohol as a salve for hardship runs from the Depression-era cowboy ad to the present. The gendered split of masculine liquor and feminine wine, forged in the 1930s, still shapes the shelves of every store. James Bond's martini and Clooney's tequila taught millions to read drinking as a marker of sophistication and success. These portrayals do not merely sell a product. They build a world, and then they leave us living in it.

It is worth saying clearly, as the historian of any such argument must, that this is not the only story one could tell. History, as the scholar John Lewis Gaddis has written, is less a list of facts than an act of mapmaking: the historian, like the cartographer, must choose which features to include and which to leave out, because no map can show everything. The three patterns mapped here are a selection, not the whole landscape. But they are a useful map, and they were drawn for a reason — because in the third decade of the twenty-first century, all three converged at once.

In late 2019 a virus began to spread, and within months the world shut its doors. The COVID-19 pandemic produced exactly the conditions under which the industry's three habits reinforce one another. It drove people online in unprecedented numbers, as work, school, family, and leisure all migrated onto screens and social-media use surged — a vast expansion of the communication technology the industry knew how to exploit. It inflicted mass psychological trauma — fear, grief, isolation, the desperate search for solace by any available means — the precise vulnerability the industry has always sought out. And it filled the media with powerful narratives, among them the industry's own: alcohol delivered to the door, the virtual happy hour, drink reframed as a reasonable way to cope with lockdown and to stay connected while staying apart. The oldest sell had met the newest medium under the worst possible conditions, and it is this convergence that my doctoral research set out to measure.

Consider how precisely the pandemic assembled the conditions for the industry's three habits to reinforce one another. The technology habit was satisfied by a historic surge online: with workplaces, schools, and social life forced onto screens, social-media use climbed to unprecedented levels, and platforms like Twitter became the central nervous system of a frightened public, carrying everything from government guidance to the chatter of neighbors. The vulnerability habit was satisfied by trauma on a scale not seen in generations — grief, fear, financial ruin, and an isolation that gnawed at mental health month after month. Researchers who studied how people coped found them reaching in every direction for relief: "reaching up" toward spiritual practice, "reaching down" toward nature and the garden, "reaching in" toward whatever inner resilience they could muster, and "reaching around" toward one another through the only doors left open, which were digital. And the portrayal habit was satisfied by a flood of narratives in which alcohol was cast as a reasonable companion to lockdown — the delivered case of wine, the virtual happy hour, the meme about day-drinking through the apocalypse — all of it framed as coping, as connection, as a way to endure. Every lever the industry had spent nine thousand years learning to pull was, in the spring of 2020, sitting right there within reach.

The findings are sobering. To study what alcohol brands actually said during the pandemic, the research analyzed nearly six thousand tweets from major alcohol brands alongside more than 486,000 tweets from ordinary users, spanning January 2019 to July 2021 — before, during, and after the first lockdowns. Because no human team could hand-sort hundreds of thousands of messages into themes, the work required a new computational method, one that pairs statistical clustering with the interpretive power of modern language models to discover the themes hiding in the text and then sort every message into them. That method surfaced five primary marketing themes, and three of them rose sharply when the lockdowns came down. Messaging around alcohol delivery and isolation drinking climbed by roughly ninety percent. Promotions framed as support for struggling restaurants — corporate social responsibility as a sales vehicle — rose by nearly one hundred ninety percent. Social-media promotions rose by half again. And the themes the industry leaned into during lockdown carried measurably stronger emotional language, consistent with appeals aimed at heightened psychological vulnerability.

A word about how one even measures such a thing, because the method matters to the trustworthiness of the finding. Hundreds of thousands of tweets are far too many for any team of researchers to read and sort by hand, and the themes that run through them are not labeled in advance; they have to be discovered. The research therefore leaned on a new approach that pairs the pattern-finding power of statistics with the interpretive power of modern language models. In broad strokes, every message is first converted into a mathematical representation of its meaning, so that messages saying similar things land near one another; software then finds the natural clusters in that space; and a language model reads representative messages from each cluster, names the theme it finds there, and goes on to sort the whole corpus accordingly. The technique lets a researcher hear what a half-million messages are saying without pretending to read each one — and it is itself a small example of this book's central theme, a brand-new communication technology turned, this time, to the work of holding the alcohol industry to account.

What the numbers describe, in human terms, is an industry that read the moment and leaned into it. As people sat frightened and alone in their homes, the messaging that surged was the messaging that met them exactly there: drink delivered to your door, so you need not go out; drink as a way to support your shuttered neighborhood restaurant, so that consumption became a civic virtue; drink promoted relentlessly across the very platforms where the isolated were spending their days. And the language grew more emotional precisely as the audience grew more raw. None of this required a conspiracy. It required only an industry doing what it has always done, with tools sharper than any it has ever held.

A second study asked the question that matters most: who was leading whom? Were the brands merely echoing what the public already felt, or were they driving the conversation? Using a statistical model designed to test influence running in both directions over time, the research found a single, unmistakable, one-way street. Increased exposure to brand tweets promoting alcohol delivery and isolation drinking predicted later increases in similar content from ordinary users — but the reverse pathway did not hold. The public did not lead the industry. The industry led the public. Across nine thousand years, that may be the oldest pattern of all: the seller speaks first, and the culture follows.

It is worth being clear about why that finding matters so much, because it cuts against the industry's favorite defense. Whenever the alcohol business is challenged, it reaches for the language of reflection: it claims merely to mirror a culture that already loves to drink, to meet a demand it did not create, to give people what they were going to want anyway. The cross-lagged analysis tests that claim directly, by asking whether the industry's messaging followed the public's behavior or led it — and the answer, in the data of the pandemic, was that the industry led. The surge in brand messaging about delivery and isolation drinking came first; the matching surge in ordinary people's content came after. The mirror, it turns out, was projecting. And if that is true in the well-measured case of a single pandemic, the long history assembled in this book suggests it has been true, in less measurable ways, for a very long time.

This is the quiet revelation beneath the whole story. We tend to imagine that our tastes, our rituals, our sense of when a drink is called for, arise naturally from within us. But a great deal of what feels like our own free relationship with alcohol was taught to us — by the Depression-era ad that gendered the drink, by the wartime photograph that fused it with freedom, by the television commercial that made it the reward for a hard day, by the feed that dressed it as a friend's suggestion. None of this means a person cannot freely choose to drink, or to abstain. It means only that the choice has been the target of the most sustained, sophisticated, and well-funded campaign of persuasion in human history, and that recognizing the campaign is the precondition for the choice being truly one's own.

Patterns this durable will not be wished away, but they can be governed, and the history assembled in this book points toward where the leverage lies. Because the industry's power has always traveled through communication media, the platforms that carry alcohol marketing today are the natural place to act — through real content moderation, genuine age verification rather than the click-to-confirm fiction, and transparency about the algorithms that decide who sees what. And because the industry has so reliably exploited the vulnerable, accountability cannot remain voluntary, as it has been since the spirits code of 1948. That means mandatory disclosure of marketing practices, real liability for foreseeable harms, enforceable standards, and codes that specifically govern how the industry may behave during crises — the very moments, as this whole history shows, when it has always reached hardest.

The case for making these measures mandatory rather than voluntary is written across the preceding chapters in the industry's own conduct. The spirits trade's celebrated decision to keep itself off television in 1948 did not reduce the volume of alcohol on the airwaves; it merely handed the airwaves to beer and bought the industry a reputation for restraint that it spent decades cashing in against regulators. The social-media codes of the present era have fared no better, presiding over a platform in which a brand could engineer its posts to look like friendship while knowing that most of its followers were underage. The lesson is consistent and old: an industry permitted to write and enforce its own rules will write rules that look responsible and enforce them only as far as is convenient. There is a precedent for doing otherwise. When the country finally decided that the tobacco industry could no longer be trusted to police its own marketing, it acted — it pulled the advertising, mandated the disclosures, and imposed real consequences. Alcohol has never faced a reckoning of that seriousness, in part because the industry has been so much better at making its product feel like a normal, even cherished, part of life. But the harms are real, the history is clear, and the tools the industry now wields are more powerful than any regulator has yet reckoned with.

What the pandemic revealed, in the end, was not a new villain but an old pattern operating with new efficiency. There was no need to imagine the alcohol industry hatching a scheme to exploit a global catastrophe. The industry simply did what nine thousand years had trained it to do — attach itself to the dominant medium, reach for people at their most vulnerable, and build portrayals that would outlast the crisis — and it did so with the most precise targeting tools ever invented, aimed at a population that had nowhere to go but online and nothing to feel but afraid. The studies that anchor this chapter did not uncover a conspiracy. They documented a reflex, measured for once with the rigor the subject deserves, and the measurement confirmed what the long history would have predicted: when the world broke, the industry leaned in, and the public followed where it led.

That is precisely why the response cannot be left to the industry's goodwill, and why it must be built for the next crisis rather than the last one. There will be another emergency — another disaster, another upheaval, another frightened and isolated population reachable through whatever medium has by then become the air we breathe. The history in this book all but guarantees that the alcohol industry will be there when it comes, fluent in the new technology and practiced at the old exploitation. The question this book leaves on the table is whether, the next time, anyone will be ready for it: whether the platforms will be governed, the marketing made transparent, the accountability made real and mandatory before the crisis rather than debated uselessly after it. We have the history. We have, now, the evidence. What remains is the will.

None of this denies anyone a drink. It is a personal choice, and a free one. But the choice is only genuinely free when it is made with open eyes — when we can see the nine-thousand-year-old machinery of persuasion for what it is, and recognize the oldest sell each time it arrives wearing the costume of the newest medium. Understanding how we got here is the first step toward deciding where we go next. That decision belongs to all of us.
"""

AFTERWORD = """
I began this research as a doctoral student at The University of Texas at Austin, surrounded by a drinking culture so pervasive that it had become invisible. Austin loves its bars and its breweries, its live music and its bottomless weekend brunches, and I took part in that culture the way most people do — without thinking very hard about where it came from or who profits from it.

Writing this history changed that, and not because the facts were shocking. Most of them are hiding in plain sight, scattered across archives and advertisements and academic papers that almost no one reads end to end. What changed me was seeing the facts laid out in sequence, because in sequence they reveal a pattern that is difficult to unsee. The alcohol industry is not simply a business that makes a product many people enjoy. It is one of the most adaptive, strategically patient, and historically consequential commercial enterprises in human civilization. It has outlasted prohibition, regulation, and reform movements that would have destroyed lesser industries, and it has done so by mastering, in turn, every medium through which human beings have ever spoken to one another.

I want to be careful about what that recognition does and does not imply. It is not a call to abstinence, and it is not an indictment of the millions of people who drink without harm, in good company, as humans have done since the first beer fermented in a clay pot at Jiahu. Alcohol is woven into our celebrations and our griefs, our religions and our friendships, and there is something almost beautiful in how old and how human that is. The argument of this book is narrower and, I think, harder to dismiss: that an enterprise this powerful, this practiced, and this willing to find us at our weakest deserves to be seen clearly, named honestly, and held to account — and that the seeing is something each of us can do, the next time a message about a drink arrives wearing the friendly face of the newest medium.

I do not presume to tell anyone whether to drink, or how much. That is not the purpose of this book. But I have come to believe that a decision made in ignorance is not fully a free decision, and that the long history of alcohol marketing has worked, deliberately, to keep us comfortable and uncurious. My hope is that these pages have supplied a little of the context that genuine choice requires.

The industry will keep adapting. There will be a medium after the feed, and the oldest sell will be waiting inside it, wearing whatever costume the new technology provides. The only real question is whether the rest of us will be paying attention. I wrote this book in the hope that more of us will.
"""

SOURCES = """
This book is a work of popular history adapted from the historical chapters of my doctoral dissertation, Alcohol Marketing on Social Media During the COVID-19 Pandemic: Historical Perspectives, Modern Evidence, and Future Regulation (The University of Texas at Austin, 2025). For readers who want the full scholarly apparatus — the complete citations behind every fact, figure, and claim in these pages, along with the methods and statistical detail of the two studies discussed in the final chapter — the dissertation is the authoritative source and is available, with all of its references, in the project repository.

In the interest of readability I have kept the narrative free of inline citations. That choice should not obscure my debts. Several existing cultural histories of alcohol were indispensable in assembling this account, and any reader who enjoys this book should seek them out: Tom Standage's A History of the World in 6 Glasses, Iain Gately's Drink: A Cultural History of Alcohol, and Rod Phillips's Alcohol: A History. The history of advertising and media that runs through these chapters draws on a wide scholarly literature on marketing, public health, and communication, all of it cited in full in the dissertation.

The historical advertisements, photographs, and other figures reproduced throughout this book are drawn from public archives and library collections — among them Duke University Libraries, the Smithsonian Institution, the Yale Center for British Art, the National WWI Museum and Memorial, and the National WWII Museum — and are credited individually in their captions. They appear here for historical and educational commentary, and rights to the images remain with their respective holders.
"""

CHRONOLOGY = [
    ("c. 7000 BCE", "The earliest known intentionally fermented beverage is brewed at Jiahu, in present-day northern China."),
    ("c. 1750 BCE", "The Code of Hammurabi regulates Babylon's taverns, the price of beer, and the honesty of the measure."),
    ("c. 160 BCE", "Cato's De Agri Cultura sets out the methods of industrial winemaking and circulates among Roman landowners."),
    ("1389 CE", "King Richard II requires English alehouses to hang signs, igniting the first competition in branding."),
    ("1500", "Hieronymus Brunschwig's printed manual on distillation spreads the making of spirits across Europe."),
    ("1717", "The Leuven brewery Den Hoorn is renamed Stella Artois; its medieval golden horn survives to this day."),
    ("1775", "The distillery later known as Buffalo Trace is founded to make American corn whiskey."),
    ("1830s–1890s", "Broadsides, billboards, direct mail, color newspapers, and magazine advertising arrive in turn — the alcohol industry an early adopter of each."),
    ("1920", "The Eighteenth Amendment begins national Prohibition; brewers pivot to near beer and malt syrup."),
    ("1933", "The Twenty-First Amendment repeals Prohibition; the three-tier system reshapes the legal trade."),
    ("1948", "The distilled-spirits industry adopts a voluntary ban on radio and television advertising, ceding the airwaves to beer."),
    ("1962", "James Bond orders an on-screen vodka in Dr. No, beginning six decades of product placement."),
    ("1970–1971", "The NIAAA is established; cigarette advertising is banned from broadcast — alcohol is left exempt."),
    ("1991", "Miller Lite begins its long sponsorship of the Dallas Cowboys."),
    ("mid-1990s", "Alcohol brands launch the first websites and banner advertisements."),
    ("late 2000s", "The industry embraces social-media marketing; internal documents later reveal awareness of underage reach."),
    ("2013", "George Clooney co-founds Casamigos, sold a few years later to Diageo for roughly a billion dollars."),
    ("2020", "COVID-19 lockdowns; brand messaging about alcohol delivery and isolation drinking surges, and the public follows."),
]


def build_chronology(doc):
    add_paragraph(doc,
                  "A glance across nine thousand years, tracing the alcohol industry "
                  "from one communication medium to the next.",
                  font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=18, line_spacing=1.2)
    for label, text in CHRONOLOGY:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        r1 = p.add_run(f"{label}  ")
        set_run_font(r1, size=11, bold=True, color=ACCENT)
        r2 = p.add_run(text)
        set_run_font(r2, size=11)


def add_section_title(doc, title):
    """A page-breaking title for the unnumbered sections (intro, afterword, etc.)."""
    doc.add_page_break()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(60)
    t.paragraph_format.space_after = Pt(28)
    trun = t.add_run(title)
    set_run_font(trun, size=20, bold=True)


CHAPTERS = [
    ("Chapter One",   "The First Pour",                  CHAPTER_1),
    ("Chapter Two",   "Signs, Spirits, and the Press",   CHAPTER_2),
    ("Chapter Three", "Blood, Sugar, and Rum",           CHAPTER_3),
    ("Chapter Four",  "The Noble Experiment",            CHAPTER_4),
    ("Chapter Five",  "A Soldier's Drink",               CHAPTER_5),
    ("Chapter Six",   "The Golden Age and the Screen",   CHAPTER_6),
    ("Chapter Seven", "The Art of the Invisible Ad",     CHAPTER_7),
    ("Chapter Eight", "Logging On",                      CHAPTER_8),
    ("Chapter Nine",  "The Feed",                        CHAPTER_9),
    ("Chapter Ten",   "The Isolation Economy",           CHAPTER_10),
]


def main():
    doc = Document()

    # Default body font: a classic book serif (Word substitutes if unavailable).
    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.3
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Garamond")
    rFonts.set(qn("w:hAnsi"), "Garamond")

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # ---- FRONT MATTER ----
    build_title_page(doc)
    build_copyright_page(doc)
    build_epigraph_page(doc)
    contents = [(None, "Introduction")]
    contents += [(k, t) for (k, t, _) in CHAPTERS]
    contents += [(None, "Afterword"),
                 (None, "A Chronology of Alcohol and the Media"),
                 (None, "A Note on Sources")]
    build_contents_page(doc, contents)

    # ---- INTRODUCTION ----
    add_section_title(doc, "Introduction")
    add_body(doc, INTRODUCTION)

    # ---- CHAPTERS ----
    for kicker, title, body in CHAPTERS:
        add_chapter_title(doc, kicker, title)
        add_body(doc, body)

    # ---- BACK MATTER ----
    add_section_title(doc, "Afterword")
    add_body(doc, AFTERWORD)
    add_section_title(doc, "A Chronology of Alcohol and the Media")
    build_chronology(doc)
    add_section_title(doc, "A Note on Sources")
    add_body(doc, SOURCES)

    # ---- SAVE ----
    output_dir = os.path.join(HERE, "deliverables")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "the_oldest_sell.docx")
    doc.save(output_path)
    print(f"Book saved to: {output_path}")

    total = len(INTRODUCTION.split()) + len(AFTERWORD.split()) + len(SOURCES.split())
    for _, _, body in CHAPTERS:
        total += len(body.split())
    n_imgs = sum(1 for body in
                 [INTRODUCTION, AFTERWORD, SOURCES] + [b for _, _, b in CHAPTERS]
                 for line in body.split("\n\n") if line.strip().startswith("[[IMG:"))
    print(f"Total word count: {total}")
    print(f"Illustrations embedded: {n_imgs}")


if __name__ == "__main__":
    main()
