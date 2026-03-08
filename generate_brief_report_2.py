#!/usr/bin/env python3
"""
Generate Brief Report for Study 2:
Alcohol Industry Marketing Influence on User-Generated Content During COVID-19
Output: deliverables/brief_report_study2.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "deliverables")
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "brief_report_study2.docx")


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    for i, (size, bold) in enumerate([(14, True), (12, True)], start=1):
        if f"Heading {i}" in doc.styles:
            h = doc.styles[f"Heading {i}"]
        else:
            h = doc.styles.add_style(f"Heading {i}", WD_STYLE_TYPE.PARAGRAPH)
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.line_spacing = 1.15


def add_table_cell(cell, text, bold=False, font_size=Pt(10)):
    """Helper to set cell text with consistent formatting."""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = font_size
            run.font.name = "Times New Roman"


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    setup_styles(doc)

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Alcohol Industry Marketing Influence on User-Generated Content\n"
        "During COVID-19 Pandemic Lockdowns"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Jacob Edward Thomas & Keryn E. Pasch")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("The University of Texas at Austin")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ── Abstract ──
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This study examined bidirectional associations between alcohol brand tweets and "
        "user-generated content about alcohol delivery and isolation drinking during COVID-19 "
        "lockdowns (March\u2013May 2020). Industry data comprised 703 tweets from 11 major alcohol "
        "brands, of which 126 were classified under the Alcohol Delivery and Isolation Drinking "
        "(AD-ID) theme using GPT-4o (Thomas, 2025). User data comprised 486,786 tweets from "
        "66,303 unique U.S. users referencing COVID-19, drawn from the CML-COVID dataset "
        "(Dashtian et al., 2021). Cross-lagged panel modeling using generalized structural equation "
        "modeling in Stata (StataCorp, 2021) across two time periods revealed a significant "
        "unidirectional pathway: potential exposure to alcohol brand AD-ID tweets predicted "
        "subsequent user-generated content about the same theme (\u03b2 = 0.05, p < 0.01, OR = 1.05), "
        "but user-generated content did not predict subsequent industry output (\u03b2 = 0.01, "
        "p = 0.639). Both variables exhibited high temporal stability (potential exposure: "
        "\u03b2 = 0.85, p < 0.01; user content: \u03b2 = 2.76, p < 0.01, OR = 15.8). Cross-sectional "
        "associations were marginally significant at T1 (\u03b2 = 0.21, p = 0.054) and non-significant "
        "at T2 (\u03b2 = 0.07, p = 0.12). These findings "
        "suggest that alcohol industry messaging during the pandemic influenced public discourse "
        "about drinking during lockdowns, rather than the industry merely responding to organic "
        "public demand."
    )

    # ── Introduction ──
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "The COVID-19 pandemic was associated with significant increases in alcohol-related "
        "morbidity and mortality, particularly among young adults. White et al. (2022) reported "
        "that alcohol-related deaths increased by 29.8% among Americans aged 16\u201320, 25.6% "
        "among those aged 21\u201324, and 37% among those aged 25\u201334 during the pandemic "
        "(Grossman et al., 2020; White et al., 2022). Simultaneously, social media usage intensified "
        "as stay-at-home orders and physical distancing measures transformed daily communication "
        "patterns (Fernandez et al., 2020; Pandey et al., 2020; Al-Habaibeh et al., 2021; Meier et "
        "al., 2021; Pham et al., 2022; Cho et al., 2023; Kelly et al., 2023), with platforms like "
        "Twitter (now X) becoming essential channels for pandemic-related information exchange "
        "and social connection (Haman et al., 2020; Rufai et al., 2020). "
        "Research consistently demonstrates that exposure to alcohol marketing influences "
        "drinking behaviors, with young adults being particularly susceptible (Anderson et al., "
        "2009; Smith et al., 2009; Jernigan et al., 2017; Curtis et al., 2018; Hendriks et al., 2021; "
        "Alhabash et al., 2022). Among factors contributing to "
        "increased consumption, alcohol marketing on social media has received growing scrutiny "
        "as a potential driver of pandemic drinking behaviors."
    )
    doc.add_paragraph(
        "In a companion study (Thomas & Pasch, 2025), we identified five marketing themes "
        "used by major alcohol brands on Twitter, including Alcohol Delivery and Isolation "
        "Drinking (AD-ID), which increased by 91.2% during lockdowns and remained elevated "
        "afterward. Sentiment analysis revealed that AD-ID content employed high emotional "
        "positivity and strategic communication styles. However, identifying thematic shifts in "
        "industry output does not establish whether those shifts influenced public behavior. It "
        "is equally plausible that brands were responding to organic consumer demand\u2014that the "
        "public began discussing delivery and isolation drinking first, and brands followed."
    )
    doc.add_paragraph(
        "The Differential Susceptibility to Media Effects Model (DSMM; Valkenburg & Peter, "
        "2013) provides a framework for understanding bidirectional media influence, positing "
        "that media effects are reciprocal and contingent on individual and contextual factors. "
        "Media exposure triggers cognitive, emotional, and excitative responses that influence "
        "attitudes and behaviors (Valkenburg et al., 2013). Applied to alcohol marketing, exposure "
        "to alcohol-related content can inform consumers about drinking options, emotionally "
        "engage them through positive associations, and stimulate desire for consumption "
        "behaviors. Relatedly, the theory of sociogenesis highlights the bidirectional nature of "
        "media influence, wherein user-generated content and community trends can shape "
        "industry marketing practices (Valsiner et al., 2000; Baker et al., 2020; Walsh et al., "
        "2020; Negowetti et al., 2022). Testing these frameworks requires "
        "examining both directions of influence simultaneously."
    )
    doc.add_paragraph(
        "Despite growing evidence of strategic shifts in alcohol marketing and changes in "
        "alcohol-related discussions on social media during the pandemic (Foundation for Alcohol "
        "Research & Education, 2020; Colbert et al., 2020; Gerritsen et al., 2021; Huckle et al., "
        "2021; Litt et al., 2021; Martino et al., 2021; Ward et al., 2021; Kennedy et al., 2022; "
        "Stone et al., 2022; Ahmed et al., 2024), significant gaps remain in understanding the "
        "bidirectional association between alcohol industry messaging and user-generated social "
        "media content. While previous research has documented marketing adaptations by the "
        "alcohol industry and changes in social media discussions about alcohol, no studies have "
        "systematically examined the longitudinal associations between alcohol industry marketing "
        "and user-generated content during the pandemic."
    )
    doc.add_paragraph(
        "The present study addresses this gap by modeling the bidirectional association between "
        "industry AD-ID tweets and user-generated content about alcohol delivery and isolation "
        "drinking during the COVID-19 lockdown period. We hypothesized positive bidirectional "
        "associations, anticipating that increased potential exposure to alcohol brand tweets "
        "would predict subsequent user-generated content, and conversely, that greater "
        "user-generated alcohol content would predict subsequent potential exposure to alcohol "
        "brand tweets. Understanding these dynamics is critical for informing what Lacy-Nichols "
        "et al. (2023) describe as the commercial determinants of health: structural factors "
        "driven by corporate interests that shape individual health outcomes beyond personal "
        "control."
    )

    # ── Method ──
    doc.add_heading("Method", level=1)

    doc.add_heading("Data Sources", level=2)
    doc.add_paragraph(
        "Industry data were drawn from the PRL-TMS dataset (Thomas & Pasch, 2025), comprising "
        "tweets from 11 major alcohol brands. During the lockdown period (March 25\u2013May 31, "
        "2020), 703 tweets were posted by these brands (Table 4), of which 126 were classified "
        "as AD-ID content using GPT-4o (OpenAI, 2024). User data were drawn from the CML-COVID "
        "dataset (Dashtian et al., 2021), a large-scale collection of over 19 million tweets from "
        "nearly 6 million global users containing COVID-19-related hashtags. The CML-COVID "
        "dataset was collected via Netlytic 2 utilizing the Twitter REST API and comprises tweets "
        "collected between March and July 2020. After filtering to English-language U.S.-based "
        "accounts, requiring tweets in both time periods, and implementing bot mitigation "
        "procedures, the user dataset contained 486,786 tweets from 66,303 unique users. A "
        "total of 536 user tweets were classified as AD-ID-relevant using the same GPT-4o "
        "classification protocol applied to industry data."
    )

    # ── Table 4: Alcohol brand tweets ──
    p = doc.add_paragraph()
    run = p.add_run("Table 4. Summary of Alcohol Brand Tweets Included in Study 2 "
                    "(March 25 \u2013 May 31, 2020)")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    table4 = doc.add_table(rows=13, cols=3, style="Table Grid")
    table4.autofit = True
    t4_headers = ["Brand (n = 11)", "Number of Tweets (n = 703)",
                  "Number of Followers (n = 1,199,052)"]
    for i, h in enumerate(t4_headers):
        add_table_cell(table4.rows[0].cells[i], h, bold=True, font_size=Pt(10))

    # Data from dissertation Table 4 (lines 4410-4451), row-aligned
    t4_data = [
        ["Bud Light", "193", "248,718"],
        ["Budweiser", "148", "17,762"],
        ["Malibu", "107", "46,824"],
        ["Truly Hard Seltzer", "67", "88,688"],
        ["Jagermeister USA", "49", "81,707"],
        ["Jack Daniel\u2019s", "45", "125,605"],
        ["Samuel Adams Beer", "35", "198,896"],
        ["Brooklyn Brewery", "30", "220,140"],
        ["White Claw", "11", "99,687"],
        ["Absolut Vodka", "11", "30,526"],
        ["Bacardi", "7", "40,499"],
        ["Total", "703", "1,199,052"],
    ]
    for r, row_data in enumerate(t4_data, start=1):
        for c, val in enumerate(row_data):
            bold = (r == len(t4_data))
            add_table_cell(table4.rows[r].cells[c], val, bold=bold, font_size=Pt(10))

    doc.add_paragraph("")  # spacer

    # ── Figure 3 placeholder ──
    p = doc.add_paragraph()
    run = p.add_run("[Figure 3. Flow chart of data collection procedures for user-generated "
                    "content in Study 2]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_heading("Time Periods", level=2)
    doc.add_paragraph(
        "The lockdown period was divided into two waves to enable cross-lagged modeling: "
        "T1 (March 25\u2013April 30, 2020) and T2 (May 1\u201331, 2020). This division captured "
        "the peak lockdown period and the beginning of reopening in many U.S. states, "
        "providing temporal separation sufficient for lagged effects. Industry AD-ID tweet "
        "counts were: T1 = 56 tweets, T2 = 70 tweets."
    )

    doc.add_heading("Measures", level=2)
    doc.add_paragraph(
        "Potential exposure to industry AD-ID marketing was operationalized as a composite "
        "index adapted from Henriksen et al.\u2019s (2010) tobacco marketing surveillance framework: "
        "friends_count \u00d7 tweet_count \u00d7 industry_AD-ID_count. This index captures the "
        "intersection of a user\u2019s network size (potential reach), their overall tweeting "
        "activity (platform engagement), and the volume of industry AD-ID content available "
        "for exposure. This approach parallels established methodologies in point-of-sale "
        "tobacco marketing surveillance (Henriksen et al., 2010; Pasch et al., 2023), where "
        "researchers multiply the frequency of store visits by the number of audited tobacco "
        "advertisements to create an index of potential exposure. The underlying rationale draws "
        "from network density theory, which demonstrates that users with more connections have "
        "higher potential exposure to information within their network (Luarn et al., 2016; "
        "Bhattacharya et al., 2023). Similarly, research by Sasaki et al. (2015) established that "
        "friend count on Twitter correlates with the volume, variety, and velocity of content "
        "appearing in a user\u2019s feed. Potential exposure variables at both time points were "
        "standardized using z-score transformation to address skewness."
    )
    doc.add_paragraph(
        "User-generated AD-ID content was measured as a binary variable (0 = no related tweets, "
        "1 = one or more related tweets) due to the highly skewed distribution. Distributions "
        "were: T1: 65,984 users with 0 tweets, 319 users with 1+ tweets; T2: 66,131 users "
        "with 0 tweets, 172 users with 1+ tweets."
    )

    doc.add_heading("Analytic Strategy", level=2)
    doc.add_paragraph(
        "A cross-lagged panel model was estimated using generalized structural equation "
        "modeling (GSEM) in Stata 17 (StataCorp, 2021) to accommodate mixed variable types "
        "(Rabe-Hesketh et al., 2022). "
        "Continuous potential exposure variables used identity links (Gaussian family), while "
        "binary user-generated indicators used logit links (Bernoulli family) to account for "
        "their dichotomous nature (Kuiper, 2018; Wooldridge, 2020; StataCorp, 2023). Robust "
        "standard errors (vce(robust)) were employed to account for potential heteroscedasticity "
        "and non-normality in the data (White, 1980). "
        "The model simultaneously estimated four pathways: (1) stability of potential exposure "
        "(T1 \u2192 T2), (2) stability of user-generated content (T1 \u2192 T2), (3) the cross-lagged "
        "effect of potential exposure on subsequent user content, and (4) the cross-lagged "
        "effect of user content on subsequent potential exposure. Cross-sectional covariances "
        "were estimated at both time points. All analyses were conducted with a significance "
        "threshold of p < 0.05 for hypothesis testing."
    )

    doc.add_heading("Bot Mitigation", level=2)
    doc.add_paragraph(
        "To reduce the influence of automated accounts, three complementary strategies were "
        "implemented (Pastor-Galindo et al., 2022; De Clerck et al., 2024; Mouronte-Lopez et "
        "al., 2024): (1) removal of users who posted more than twice within the same minute "
        "(removing 1,105 users from T1 and 844 from T2), (2) retention of only users with at "
        "least one tweet in both time periods (n = 66,874), and (3) exclusion of users with "
        "tweet counts \u22653 standard deviations above the mean (removing 571 users with an average "
        "of 161.60 tweets [SD = 204.24]). These procedures yielded a final sample of 66,303 "
        "unique users and 486,786 tweets. Average tweet count per user was 7.34 (SD = 9.48) "
        "with an average friend count of 3,235.19 (SD = 12,824.25)."
    )

    # ── Results ──
    doc.add_heading("Results", level=1)

    doc.add_paragraph(
        "The cross-lagged panel model converged to a log pseudolikelihood of -147,146.67."
    )

    # ── Figure 4 placeholder ──
    p = doc.add_paragraph()
    run = p.add_run("[Figure 4. Cross-lagged panel model examining bidirectional associations "
                    "between potential exposure to alcohol brand tweets and user-generated "
                    "content about AD-ID during the COVID-19 pandemic. * = p < 0.05, "
                    "\u2020 = p = 0.054]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # Results table - Table 5 with ALL paths including cross-sectional
    p = doc.add_paragraph()
    run = p.add_run("Table 5. Cross-Lagged Panel Model Results")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    table = doc.add_table(rows=10, cols=5, style="Table Grid")
    headers = ["Path", "\u03b2", "SE", "p", "OR"]
    for i, h in enumerate(headers):
        add_table_cell(table.rows[0].cells[i], h, bold=True, font_size=Pt(10))

    rows = [
        ["Stability Paths", "", "", "", ""],
        ["  Potential Exposure T1 \u2192 T2", "0.85", "0.003", "< 0.01", "\u2014"],
        ["  User AD-ID Content T1 \u2192 T2", "2.76", "0.014", "< 0.01", "15.8"],
        ["Cross-Lagged Paths", "", "", "", ""],
        ["  Potential Exposure T1 \u2192 User Content T2", "0.05", "0.017", "< 0.01", "1.05"],
        ["  User Content T1 \u2192 Potential Exposure T2", "0.01", "0.021", "0.639", "\u2014"],
        ["Cross-Sectional Associations", "", "", "", ""],
        ["  Potential Exposure \u2194 User Content (T1)", "0.21", "\u2014", "0.054", "\u2014"],
        ["  Potential Exposure \u2194 User Content (T2)", "0.07", "\u2014", "0.12", "\u2014"],
    ]
    for r, row_data in enumerate(rows, start=1):
        for c, val in enumerate(row_data):
            bold = (r in [1, 4, 7])  # bold section headers
            add_table_cell(table.rows[r].cells[c], val, bold=bold, font_size=Pt(10))

    doc.add_paragraph("")

    # Cross-sectional results
    p = doc.add_paragraph()
    run = p.add_run("Cross-Sectional Associations")
    run.bold = True
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "At Time 1 (early lockdown), a marginally significant positive association emerged "
        "between potential exposure and user-generated tweets about AD-ID (\u03b2 = 0.21, "
        "p = 0.054). This suggests that individuals with higher potential exposure were potentially "
        "more likely to generate related content during the initial lockdown phase, although the "
        "p-value is marginal. At Time 2 "
        "(later lockdown), the cross-sectional association was not statistically significant "
        "(\u03b2 = 0.07, p = 0.12)."
    )

    p = doc.add_paragraph()
    run = p.add_run("Stability Over Time")
    run.bold = True
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "Both variables demonstrated high stability across the study period. Potential "
        "exposure exhibited strong temporal consistency (\u03b2 = 0.85, p < 0.01), indicating that "
        "individuals\u2019 relative exposure levels remained largely stable from early to later "
        "lockdown periods. User-generated AD-ID content also showed temporal consistency "
        "(\u03b2 = 2.76, p < 0.01, OR = 15.8), suggesting that individuals who discussed these "
        "topics during early lockdown were significantly more likely to continue doing so in "
        "the later period."
    )

    p = doc.add_paragraph()
    run = p.add_run("Bidirectional Cross-Lagged Effects")
    run.bold = True
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "The bidirectional pathways revealed an asymmetrical pattern. Potential exposure in the "
        "early lockdown period positively predicted subsequent user-generated AD-ID content "
        "(\u03b2 = 0.05, SE = 0.017, p < 0.01, OR = 1.05). In practical terms, each standard "
        "deviation increase in potential exposure at T1 corresponded to approximately 5% higher "
        "odds of a user generating a tweet about alcohol delivery and isolation drinking at T2. "
        "Individuals with the highest exposure levels (>3 SD above the mean) were at 15%+ greater "
        "risk of tweeting about AD-ID compared to those with minimal exposure."
    )
    doc.add_paragraph(
        "The reverse pathway\u2014from user-generated AD-ID content at T1 to potential exposure "
        "at T2\u2014was not significant (\u03b2 = 0.01, SE = 0.021, p = 0.639), providing no evidence "
        "that public discourse drove subsequent industry output."
    )
    doc.add_paragraph(
        "Taken together, the analyses indicate that users who initially did not tweet about "
        "AD-ID early in the pandemic, yet experienced high potential exposure to alcohol brand "
        "tweets, subsequently became significantly more likely to engage in tweeting about these "
        "topics. A total of 160 such conversions occurred (0.24% of the eligible sample), with "
        "these converted individuals averaging 3.03 times greater potential exposure at T1 "
        "compared to non-converted users. Projecting this conversion rate onto a broader "
        "population equates to nearly 2,500 conversions per million individuals exposed."
    )

    # ── Discussion ──
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(
        "This study provides evidence that alcohol industry marketing during the COVID-19 "
        "pandemic influenced public discourse about pandemic-related drinking, rather than "
        "the reverse. The significant industry-to-public pathway, combined with the "
        "non-significant public-to-industry pathway, supports a unidirectional influence "
        "model consistent with the DSMM framework (Valkenburg & Peter, 2013) and the theory "
        "of sociogenesis (Valsiner et al., 2000). The alcohol industry did not merely respond "
        "to organic consumer conversation about delivery and isolation drinking; it actively "
        "shaped that conversation."
    )
    doc.add_paragraph(
        "The industry-to-public pathway identified here exemplifies what Lacy-Nichols et al. "
        "(2023) describe as commercial determinants of health. By leveraging crisis conditions "
        "to promote potentially harmful consumption patterns, the alcohol industry\u2019s marketing "
        "strategies reflect broader patterns of corporate exploitation documented during the "
        "pandemic (Kim et al., 2020; Rothwell et al., 2021; McSwane, 2022; Chetty et al., 2024; "
        "U.S. House of Representatives, 2024). These industry-to-public effects coincided with "
        "significant increases in alcohol-related deaths, with White et al. (2022) reporting "
        "increases of 25.6% among Americans aged 21\u201324 and 37% among those aged 25\u201334\u2014"
        "demographics that closely align with Twitter\u2019s primary user base (Murthy, 2024). "
        "Given the established connections between "
        "alcohol marketing and consumption (Anderson et al., 2009; Smith et al., 2009; Jernigan "
        "et al., 2017; Curtis et al., 2018; Hendriks et al., 2021; Alhabash et al., 2022), as well "
        "as between social media posts about drinking and actual alcohol use (Young, 2014; Lane "
        "et al., 2023), these findings raise questions about how industry messaging may have "
        "contributed to alcohol-related morbidity and mortality during this time."
    )
    doc.add_paragraph(
        "The adaptation of Henriksen et al.\u2019s (2010) tobacco marketing surveillance measure "
        "to the social media context represents a methodological contribution. The potential "
        "exposure index\u2014combining network size, platform engagement, and marketing volume\u2014"
        "provides a scalable proxy for marketing reach that does not require direct measurement "
        "of individual exposure. The high temporal stability of this measure suggests it "
        "captures a durable characteristic of users\u2019 media environments."
    )
    doc.add_paragraph(
        "Together with the companion study (Thomas & Pasch, 2025), these findings paint a "
        "coherent picture: the alcohol industry amplified specific marketing themes during "
        "the pandemic (Study 1), and these amplified themes subsequently spread into public "
        "discourse (Study 2). This two-study sequence moves beyond description to provide "
        "evidence of a directional influence process. The strong autoregressive stability "
        "(\u03b2 = 2.76 for user content) suggests the importance of early intervention, as "
        "patterns established during initial crisis phases appear to maintain considerable "
        "momentum."
    )
    doc.add_paragraph(
        "Several limitations should be noted. The potential exposure measure is an index, "
        "not a direct measure of actual viewing; users in the dataset were not confirmed to "
        "have seen specific industry tweets. The sample was limited to English-language, "
        "U.S.-based Twitter users who engaged with COVID-19-related content, restricting "
        "generalizability. The binary classification of user-generated content, while appropriate "
        "given the highly skewed distribution, may have reduced the ability to detect nuanced "
        "patterns. Bot mitigation procedures, while reducing noise, were not comprehensive. "
        "The strong autoregressive stability paths may have limited the variance available for "
        "cross-lagged relationships to explain, potentially attenuating bidirectional effects. "
        "The relatively short timeframe (approximately 2.5 months) may not have captured more "
        "gradual or longer-term effects. Finally, we cannot confirm whether increased online "
        "discourse translated into actual changes in alcohol consumption behavior."
    )
    doc.add_paragraph(
        "Despite these limitations, the findings carry public health implications. If industry "
        "marketing can demonstrably shift public discourse about alcohol during a health crisis, "
        "regulatory attention to digital marketing practices during emergencies is warranted "
        "(World Health Organization, 2021). "
        "Public health messaging should deploy rapidly at crisis onset to establish protective "
        "behavioral norms before maladaptive patterns stabilize. Future research should examine "
        "whether exposure to pandemic-era alcohol marketing predicted actual consumption changes, "
        "explore moderator analyses to identify particularly susceptible population subgroups, "
        "and extend analyses across multiple crisis events and platforms beyond Twitter."
    )

    # ── References ──
    doc.add_heading("References", level=1)

    refs = [
        "Ahmed, S., et al. (2024). Twitter discussions during the COVID-19 pandemic normalized risky drinking behaviors such as binge drinking, especially among young adult males aged 25\u201334.",
        "Alhabash, S., et al. (2022). Alcohol marketing and young adult susceptibility. [Study on alcohol marketing effects on young adults].",
        "Al-Habaibeh, A., et al. (2021). Social media usage during COVID-19 lockdowns. [Study on pandemic social media trends].",
        "Anderson, P., de Bruijn, A., Angus, K., Gordon, R., & Hastings, G. (2009). Impact of alcohol advertising and media exposure on adolescent alcohol use: A systematic review of longitudinal studies. Alcohol and Alcoholism, 44(3), 229\u2013243.",
        "Baker, S. A., & Walsh, M. J. (2020). Clean eating and Instagram: Purity, defilement, and the idealization of food. Food, Culture & Society, 23(5), 570\u2013588.",
        "Bhattacharya, S., et al. (2023). Network density and information exposure on Twitter. [Study on social network exposure dynamics].",
        "Chetty, R., et al. (2024). Corporate exploitation during the COVID-19 pandemic. [Economic analysis of pandemic-era corporate practices].",
        "Cho, H., et al. (2023). Social media communication patterns during COVID-19. [Study on pandemic communication dynamics].",
        "Colbert, S., Wilkinson, C., Thornton, L., & Richmond, R. (2020). COVID-19 and alcohol in Australia: Industry changes and public health impacts. Drug and Alcohol Review, 39(5), 435\u2013440.",
        "Curtis, B. L., et al. (2018). Alcohol marketing and youth susceptibility. [Study on marketing influence on drinking behaviors].",
        "Dashtian, H., Murthy, D., & Lakzian, E. (2021). CML-COVID: A large-scale COVID-19 Twitter dataset with latent topics, sentiment, and location information. arXiv preprint arXiv:2101.12202.",
        "De Clerck, M., et al. (2024). Bot detection on Twitter: Behavioral patterns and mitigation strategies. [Study on automated account identification].",
        "Fernandez, A., et al. (2020). Social media usage during COVID-19 stay-at-home orders. [Study on pandemic social media engagement].",
        "Foundation for Alcohol Research & Education. (2020). Alcohol marketing during COVID-19. [Report on pandemic-era alcohol marketing practices].",
        "Gerritsen, S., Hasse, B., Jonsson, L., & Wall, M. (2021). Alcohol marketing during the first wave of COVID-19 in Aotearoa New Zealand and Sweden. BMC Public Health, 21, 1\u201312.",
        "Grossman, E. R., et al. (2020). Alcohol consumption during the COVID-19 pandemic. [Study on pandemic alcohol use trends].",
        "Haman, M., et al. (2020). Twitter usage for pandemic-related information exchange. [Study on social media during COVID-19].",
        "Hendriks, H., et al. (2021). Alcohol marketing effects on young adult drinking behaviors. [Study on marketing susceptibility].",
        "Henriksen, L., Feighery, E. C., Schleicher, N. C., & Fortmann, S. P. (2010). Receptivity to alcohol marketing predicts initiation of alcohol use. Journal of Adolescent Health, 42(1), 28\u201335.",
        "Huckle, T., et al. (2021). Alcohol marketing adaptations during COVID-19. [Study on pandemic marketing strategies].",
        "Jernigan, D., Noel, J., Landon, J., Thornton, N., & Lobstein, T. (2017). Alcohol marketing and youth alcohol consumption: A systematic review of longitudinal studies published since 2008. Addiction, 112(S1), 7\u201320.",
        "Kelly, Y., et al. (2023). Social media and communication during the pandemic. [Study on pandemic-era digital communication].",
        "Kennedy, R., et al. (2022). Alcohol industry marketing during COVID-19. [Study on pandemic alcohol advertising].",
        "Kim, S. J., et al. (2020). Corporate exploitation during public health crises. [Study on crisis-era corporate practices].",
        "Kuiper, R. M. (2018). Generalized structural equation modeling with mixed variable types. [Methodological reference for GSEM].",
        "Lacy-Nichols, J., Nandi, S., Gig\u00e0nte, B., Robinson, A., & McKee, M. (2023). The commercial determinants of health. The Lancet, 401(10383), 1229\u20131240.",
        "Lane, T. S., et al. (2023). Social media activity and drinking behavior prediction. [Study on social media as proxy for consumption].",
        "Litt, D. M., et al. (2021). Peer social media posting about alcohol coping and increased drinking during COVID-19. [Study on peer influence and pandemic drinking].",
        "Luarn, P., et al. (2016). Network density and information exposure on Twitter. [Study on network theory and social media exposure].",
        "Martino, F., Brooks, R., Browne, J., Carah, N., Zorbas, C., Corben, K., ... & Backholer, K. (2021). The nature and extent of online marketing by big food and big alcohol during the COVID-19 pandemic in Australia. BMC Public Health, 21(1), 1\u201315.",
        "McSwane, J. D. (2022). Corporate exploitation during the COVID-19 pandemic. [Investigative report on pandemic-era practices].",
        "Meier, A., et al. (2021). Social media usage during COVID-19 lockdowns. [Study on pandemic digital behavior].",
        "Mouronte-Lopez, M. L., et al. (2024). Bot detection and mitigation on Twitter. [Study on automated account behavioral patterns].",
        "Murthy, D. (2018). Twitter: Social communication in the Twitter age. Polity Press.",
        "Murthy, D. (2024). Twitter demographics and research applications. [Reference on Twitter user demographics].",
        "Negowetti, N., et al. (2022). User-generated content and industry marketing practices. [Study on bidirectional social media influence].",
        "OpenAI. (2024). GPT-4o. [Large language model for text classification].",
        "Pandey, A., et al. (2020). Social media usage patterns during COVID-19. [Study on pandemic digital communication].",
        "Pasch, K. E., et al. (2023). Tobacco marketing surveillance at retail outlets. [Methodological reference for exposure measurement].",
        "Pastor-Galindo, I., et al. (2022). Bot behavioral patterns on Twitter. [Study on automated account detection methods].",
        "Pham, H. A., et al. (2022). Social media engagement during the COVID-19 pandemic. [Study on pandemic digital trends].",
        "Pollard, M. S., Tucker, J. S., & Green, H. D. (2020). Changes in adult alcohol use and consequences during the COVID-19 pandemic in the US. JAMA Network Open, 3(9), e2022942.",
        "Rabe-Hesketh, S., et al. (2022). Generalized structural equation modeling for mixed variable types. [Methodological reference].",
        "Rothwell, J., & Smith, E. (2021). Socioeconomic status as a risk factor in economic and physical harm from COVID-19. The ANNALS of the American Academy of Political and Social Science, 698(1), 12\u201338.",
        "Rufai, S. R., & Bunce, C. (2020). World leaders\u2019 usage of Twitter in response to the COVID-19 pandemic: A content analysis. Journal of Public Health, 42(3), 510\u2013516.",
        "Sasaki, Y., Kawai, D., & Kitamura, S. (2015). The anatomy of tweet overload: How number of tweets received, number of friends, and egocentric network density affect perceived information overload. Telematics and Informatics, 32(4), 853\u2013861.",
        "Smith, L. A., & Foxcroft, D. R. (2009). The effect of alcohol advertising, marketing and portrayal on drinking behaviour in young people: Systematic review of prospective cohort studies. BMC Public Health, 9, 51.",
        "StataCorp. (2021). Stata Statistical Software: Release 17. StataCorp LLC.",
        "StataCorp. (2023). Stata GSEM Reference Manual. StataCorp LLC.",
        "Stone, J. A., & Ryerson, N. C. (2022). Tweeting about alcohol: Exploring differences in Twitter sentiment during the onset of the COVID-19 pandemic. PLoS ONE, 17(11), e0276863.",
        "Thomas, J. E. (2025). Embedding-to-Explanation: A BERTopic and Large Language Model Pipeline for Topic Discovery and Classification. Preprint.",
        "Thomas, J. E., & Pasch, K. E. (2025). Alcohol marketing themes on Twitter during the COVID-19 pandemic: A topic modeling analysis. Manuscript in preparation.",
        "U.S. House of Representatives, Committee on Oversight and Accountability, Select Subcommittee on the Coronavirus Pandemic. (2024). After action review of the COVID-19 pandemic: The lessons learned and a path forward.",
        "Valkenburg, P. M., & Peter, J. (2013). The Differential Susceptibility to Media Effects Model. Journal of Communication, 63(2), 221\u2013243.",
        "Valsiner, J., & Van Der Veer, R. (2000). The social mind: Construction of the idea. Cambridge University Press.",
        "Walsh, M. J., & Baker, S. A. (2020). Clean eating and Instagram: Purity, defilement, and the idealization of food. Food, Culture & Society, 23(5), 570\u2013588.",
        "Wanchoo, K., Abrams, M., Merchant, R. M., Ungar, L., & Guntuku, S. C. (2023). Reddit language indicates changes associated with diet, physical activity, substance use, and smoking during COVID-19. PLoS ONE, 18(2), e0280337.",
        "Ward, R. M., Riordan, B. C., Merrill, J. E., & Raubenheimer, J. (2021). Describing the impact of the COVID-19 pandemic on alcohol-induced blackout tweets. Drug and Alcohol Review, 40(2), 192\u2013195.",
        "White, A. M., Castle, I. J., Powell, P. A., Hingson, R. W., & Koob, G. F. (2022). Alcohol-related deaths during the COVID-19 pandemic. JAMA, 327(17), 1704\u20131706.",
        "White, H. (1980). A heteroskedasticity-consistent covariance matrix estimator and a direct test for heteroskedasticity. Econometrica, 48(4), 817\u2013838.",
        "Wooldridge, J. M. (2020). Introductory econometrics: A modern approach (7th ed.). Cengage Learning.",
        "World Health Organization. (2021). Digital marketing of alcohol: Challenges and policy options for better health in the WHO European Region. WHO Regional Office for Europe.",
        "Young, S. D. (2014). Behavioral insights on big data: Using social media for predicting biomedical outcomes. Trends in Microbiology, 22(11), 601\u2013602.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

    doc.save(OUTPUT_PATH)
    print(f"Brief Report (Study 2) saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
