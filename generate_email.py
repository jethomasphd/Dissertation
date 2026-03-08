#!/usr/bin/env python3
"""
Generate email draft to Dr. Keryn Pasch regarding publication strategy.
Output: deliverables/email_to_pasch.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "deliverables")
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "email_to_pasch.docx")


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # ── Email Headers ──
    headers = [
        ("From:", "Jacob (Jake) Thomas"),
        ("To:", "Dr. Keryn Pasch"),
        ("Subject:", "My Plan to Publish \u2014 An Atypical but Strategic Approach"),
    ]
    for label, value in headers:
        p = doc.add_paragraph()
        run_label = p.add_run(label + " ")
        run_label.bold = True
        p.add_run(value)

    # Separator
    p = doc.add_paragraph()
    run = p.add_run("\u2014" * 40)
    run.font.color.rgb = None

    # ── Body ──
    doc.add_paragraph(
        "Hi Keryn,"
    )

    doc.add_paragraph(
        "I wanted to lay out my plan for publishing from the dissertation, because I know "
        "it\u2019s going to look a little different from the standard approach. I\u2019ve thought about "
        "this a lot, and I think the atypical route is actually the smarter one given where "
        "things stand. I\u2019d love your feedback."
    )

    doc.add_paragraph(
        "Here\u2019s the honest reality: the data are from 2019\u20132021, the dissertation took seven "
        "years, and the field moves fast. If I go the traditional route\u2014submit two full-length "
        "journal articles, wait 6\u201312 months per round of review, revise, resubmit\u2014I\u2019m looking "
        "at another one to two years before anything is actually out there. By then, the "
        "pandemic alcohol marketing literature will have moved on, and the window for this "
        "work to make an impact will have narrowed considerably."
    )

    doc.add_paragraph(
        "So instead, I\u2019m planning to dissect the dissertation into multiple outputs that "
        "can be released quickly, each building on the others:"
    )

    # Numbered list
    items = [
        (
            "A popular history book based on Part A. ",
            "I\u2019ve already written a full historical account of the alcohol industry\u2019s "
            "relationship with marketing regulation. There is genuinely nothing like it out "
            "there\u2014no single book tells this story in an accessible way. This is a unique "
            "contribution that no journal article could capture, and it\u2019s essentially ready to go."
        ),
        (
            "A methodology preprint and open-source tool (TopicFlow). ",
            "The BERTopic + LLM pipeline I developed is novel and useful beyond this specific "
            "project. I can publish the preprint immediately to a preprint server, establishing "
            "priority on the methodology. The software is already packaged and documented. This "
            "gives the method a citable home right away."
        ),
        (
            "Two brief research reports, one for each study. ",
            "These are shorter, more focused than full journal articles, and faster to review. "
            "Study 1 covers the topic model and prevalence analysis; Study 2 covers the "
            "cross-lagged model linking industry tweets to user-generated content. Both cite "
            "the preprint for methodology, so they stay concise. Brief reports are still "
            "peer-reviewed\u2014they\u2019re not lesser publications, just more efficient ones."
        ),
        (
            "An HTML portfolio piece. ",
            "A visual, interactive summary showing how all the pieces fit together. This is "
            "partly for future employers and collaborators, but it also demonstrates that I can "
            "think strategically about research dissemination, not just write papers."
        ),
    ]
    for bold_part, rest in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(bold_part)
        run.bold = True
        p.add_run(rest)

    doc.add_paragraph(
        "Here\u2019s why I think this works better than the conventional path. Speed is the biggest "
        "factor\u2014the data are aging, and this strategy gets findings into the world in months "
        "rather than years. The methodology is genuinely novel and deserves standalone "
        "publication rather than being buried in a methods section. The book is something no "
        "journal could accommodate. And critically, each piece cites and builds on the others, "
        "creating a coherent body of work rather than isolated fragments. The preprint "
        "establishes priority on the computational approach, the brief reports provide the "
        "empirical evidence, the book provides historical context, and the portfolio ties it "
        "all together."
    )

    doc.add_paragraph(
        "I know this isn\u2019t the path you probably envisioned when we started. But I genuinely "
        "believe it\u2019s the most strategic way to maximize the impact of work that I\u2019m proud of "
        "and that I think matters. And honestly, I wouldn\u2019t be in a position to pull this off "
        "without the years of mentorship and patience you\u2019ve given me. I am deeply grateful for "
        "that\u2014more than I\u2019ve probably said."
    )

    doc.add_paragraph(
        "Would love to talk through this whenever you have time. No rush."
    )

    # Sign-off
    doc.add_paragraph("")
    doc.add_paragraph("Jake")

    doc.save(OUTPUT_PATH)
    print(f"Email draft saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
