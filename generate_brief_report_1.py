#!/usr/bin/env python3
"""
Generate Brief Report for Study 1:
Alcohol Marketing Themes on Twitter During the COVID-19 Pandemic
Output: deliverables/brief_report_study1.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "deliverables")
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "brief_report_study1.docx")


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    # Heading style
    for i, (size, bold) in enumerate([(14, True), (12, True)], start=1):
        if f"Heading {i}" in doc.styles:
            h = doc.styles[f"Heading {i}"]
        else:
            h = doc.styles.add_style(f"Heading {i}", WD_STYLE_TYPE.PARAGRAPH)
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.line_spacing = 1.15


def add_paragraph(doc, text, style="Normal", bold=False, italic=False, alignment=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if alignment is not None:
        p.alignment = alignment
    return p


def add_table_cell(cell, text, bold=False, font_size=Pt(9)):
    """Helper to set cell text with consistent formatting."""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = font_size
            run.font.name = "Times New Roman"


def build_document():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    setup_styles(doc)

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Alcohol Marketing Themes on Twitter During the COVID-19 Pandemic:\n"
        "A Topic Modeling Analysis"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Jacob Edward Thomas & Keryn E. Pasch")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("The University of Texas at Austin")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ── Abstract ──
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Alcohol brands rapidly adapted their social media marketing strategies during the "
        "COVID-19 pandemic, yet few studies have systematically tracked thematic shifts across "
        "pandemic phases. This study analyzed 5,872 tweets posted by 11 major alcohol brands "
        "between January 2019 and July 2021 using a novel computational pipeline combining "
        "BERTopic modeling with hyperparameter optimization and GPT-4o classification (Thomas, "
        "2025). Five marketing themes were identified: Alcohol Delivery and Isolation Drinking "
        "(AD-ID), Restaurant Support, Social Media Promotions, Sports, and Gaming. Proportional "
        "z-tests with Bonferroni correction revealed that three themes showed statistically "
        "significant increases during COVID-19 lockdowns (March\u2013May 2020): AD-ID (+91.2%), "
        "Restaurant Support (+189.5%), and Social Media Promotions (+150.0%). Critically, AD-ID "
        "and Social Media Promotions remained significantly elevated in the post-lockdown period "
        "(June 2020\u2013July 2021), suggesting durable shifts in alcohol marketing strategy rather "
        "than temporary responses to crisis conditions. Sentiment analysis using LIWC (Boyd et al., "
        "2022) revealed that the three pandemic-responsive themes employed stronger emotional "
        "appeals, with high positive tone (80.16%\u201398.49%) and minimal negative affect, raising "
        "concerns about potential exploitation of pandemic-related psychological vulnerability. "
        "These findings extend prior qualitative work by providing longitudinal, quantitative "
        "evidence of how the alcohol industry leveraged a public health crisis to normalize home "
        "delivery and isolation drinking."
    )

    # ── Introduction ──
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "The COVID-19 pandemic fundamentally altered patterns of daily life beginning in "
        "March 2020, when governments worldwide implemented lockdowns, business closures, and "
        "social distancing mandates. These disruptions created conditions associated with "
        "increased alcohol consumption: social isolation, economic stress, disrupted routines, "
        "and heightened anxiety (Pollard et al., 2020). Concurrently, alcohol-related deaths "
        "increased substantially, with deaths among Americans aged 16\u201320 rising by 29.8%, those "
        "aged 21\u201324 by 25.6%, and those aged 25\u201334 by 37% (White et al., 2022). Social media "
        "usage surged as people turned to digital platforms for connection, entertainment, and "
        "information (Wiederhold, 2020). This confluence of increased alcohol risk factors and "
        "heightened digital engagement created a unique opportunity for alcohol marketers."
    )
    doc.add_paragraph(
        "The alcohol industry has a well-documented history of adapting its marketing strategies "
        "to capitalize on cultural moments and public health events (Jernigan & Ross, 2020). "
        "Research consistently demonstrates that exposure to alcohol marketing influences drinking "
        "behaviors, with young adults being particularly susceptible (Anderson et al., 2009; Smith "
        "et al., 2009; Jernigan et al., 2017). Social media platforms are particularly amenable to "
        "such adaptation because they allow brands to shift messaging rapidly and at low cost "
        "(Barry et al., 2016). Early research during the pandemic identified several marketing "
        "themes adopted by alcohol brands, including home delivery promotions, isolation drinking "
        "normalization, and corporate social responsibility messaging (Martino et al., 2021; "
        "Gerritsen et al., 2021; Colbert et al., 2020). However, these studies were largely "
        "qualitative, focused on narrow time windows, and did not provide systematic comparisons "
        "of theme prevalence across pre-pandemic, lockdown, and post-lockdown periods."
    )
    doc.add_paragraph(
        "Moreover, existing computational approaches to analyzing large-scale social media "
        "marketing data have been limited by methodological constraints. Traditional topic "
        "models such as Latent Dirichlet Allocation (LDA) struggle with short texts (Qiang "
        "et al., 2020), and manual content analysis is impractical at scale. Recent advances "
        "in transformer-based topic modeling\u2014particularly BERTopic (Grootendorst, 2022)\u2014offer "
        "promising alternatives but require careful hyperparameter optimization and human "
        "validation to produce interpretable results."
    )
    doc.add_paragraph(
        "The present study addresses these gaps by applying a novel computational pipeline, "
        "Embedding-to-Explanation (Thomas, 2025), to a longitudinal dataset of alcohol brand "
        "tweets spanning pre-pandemic, lockdown, and post-lockdown periods. The pipeline "
        "combines BERTopic modeling with systematic hyperparameter optimization, large language "
        "model (LLM) topic naming, and GPT-4o classification, informed by the Themes and Topics "
        "framework (Gillies et al., 2022). We pursued two aims: (1) identify the principal "
        "marketing themes used by major alcohol brands on Twitter, and (2) compare the prevalence "
        "of these themes across pandemic phases to determine whether observed shifts were "
        "temporary or durable. Additionally, we explored the cognitive, emotional, and excitative "
        "composition of theme content during the peri-pandemic period through the lens of the "
        "Differential Susceptibility to Media Effects Model (DSMM; Valkenburg & Peter, 2013)."
    )

    # ── Method ──
    doc.add_heading("Method", level=1)

    doc.add_heading("Sample", level=2)
    doc.add_paragraph(
        "The dataset comprised 5,872 original tweets posted by 11 major alcohol companies "
        "between January 2019 and July 2021. These companies were drawn from the PRL-TMS "
        "(Publicly Retrievable Longitudinal Twitter Marketing Sample), a purpose-built dataset "
        "capturing the complete tweet histories of prominent alcohol brands with active U.S. "
        "marketing presences. Retweets and replies were excluded to focus on original marketing "
        "content. The sample spanned three analytically defined periods: pre-pandemic (January "
        "2019\u2013January 2020), peri-lockdown (March\u2013May 2020), and post-lockdown (June 2020\u2013"
        "July 2021). Table 1 summarizes the tweet distribution across the 11 companies."
    )

    # ── Table 1: Company Tweet Summary ──
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Table 1. Summary of Tweets Included in Study 1")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    table1 = doc.add_table(rows=13, cols=3, style="Table Grid")
    table1.autofit = True
    t1_headers = ["Company (n = 11)", "Number of Tweets (n = 5,872)",
                  "Number of Followers (n = 1,199,052)"]
    for i, h in enumerate(t1_headers):
        add_table_cell(table1.rows[0].cells[i], h, bold=True, font_size=Pt(9))

    t1_data = [
        ["Bud Light", "1,142", "248,718"],
        ["Truly Hard Seltzer", "755", "17,762"],
        ["Malibu", "695", "46,824"],
        ["J\u00e4germeister", "653", "88,688"],
        ["Samuel Adams Beer", "594", "81,707"],
        ["Brooklyn Brewery", "592", "125,605"],
        ["Jack Daniel\u2019s", "535", "198,896"],
        ["Budweiser", "409", "220,140"],
        ["Bacardi", "233", "99,687"],
        ["Absolut Vodka", "147", "30,526"],
        ["White Claw Hard Seltzer", "117", "40,499"],
        ["Total", "5,872", "1,199,052"],
    ]
    for r, row_data in enumerate(t1_data, start=1):
        for c, val in enumerate(row_data):
            bold = (r == len(t1_data))  # bold the total row
            add_table_cell(table1.rows[r].cells[c], val, bold=bold, font_size=Pt(9))

    doc.add_paragraph("")  # spacer

    # ── Figure 1 placeholder ──
    p = doc.add_paragraph()
    run = p.add_run("[Figure 1. Month-over-month volume of tweets included in Study 1, "
                    "COVID-19 lockdown period highlighted (n = 5,872)]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_heading("Theme Identification", level=2)
    doc.add_paragraph(
        "Marketing themes were identified using the Embedding-to-Explanation pipeline described "
        "in Thomas (2025) and available as open-source software. The procedure involved three "
        "stages, informed by the Themes and Topics framework (Gillies et al., 2022). "
        "First, BERTopic (Grootendorst, 2022) was applied with systematic hyperparameter "
        "optimization: 50 random iterations were conducted for each of 24 candidate topic "
        "solutions (ranging from 2 to 25 topics), yielding 1,200 fitted models. The best-performing "
        "model was selected based on c_v coherence, a validated measure of topic interpretability "
        "(R\u00f6der et al., 2015). Second, LLM-based democratic topic naming was performed by "
        "submitting representative documents from each topic to GPT-4o across 5,000 iterations, "
        "with final labels determined by majority vote. Third, human thematic synthesis was "
        "conducted by two researchers who reviewed BERTopic output, sample tweets, and LLM "
        "interpretations through the lens of current alcohol marketing literature (Martino et al., "
        "2021; Gerritsen et al., 2021). The optimal model contained 7 granular topics "
        "(coherence = 0.768), which were synthesized into 5 interpretable themes: Alcohol "
        "Delivery and Isolation Drinking (AD-ID), Restaurant Support, Social Media Promotions, "
        "Sports, and Gaming."
    )

    # ── Figure 2 placeholder ──
    p = doc.add_paragraph()
    run = p.add_run("[Figure 2. Topic model performance: coherence scores across 1,200 "
                    "candidate models]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── Table 2: Theme Structure ──
    p = doc.add_paragraph()
    run = p.add_run("Table 2. Theme Structure and Representative Content")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    table2 = doc.add_table(rows=8, cols=4, style="Table Grid")
    table2.autofit = True
    t2_headers = ["Theme", "Constituent Topics\n(LLM Interpretation)",
                  "Representative Words", "Example Tweets"]
    for i, h in enumerate(t2_headers):
        add_table_cell(table2.rows[0].cells[i], h, bold=True, font_size=Pt(8))

    t2_data = [
        [
            "Alcohol Delivery\nand Isolation\nDrinking\n(AD-ID)",
            "Pandemic Alcohol\nDelivery and\nEnjoyment",
            "sunshine, drizly, bit, day,\norder, happy, part,\noriginal, drink, get",
            "Pro tip: Put your beer in a\ncoffee mug and no one knows\nyou\u2019re drinking on the video\nconference.\n\nBring the beachside vibe\n*inside* with Malibu delivered\nto your door in 60 minutes or\nless. Get $5 off your first\norder, courtesy of Drizly,\nwith code \u2018Sunshine20\u2019!"
        ],
        [
            "Restaurant\nSupport",
            "COVID-19 Restaurant\nSupport and Donation\nInitiatives",
            "restaurant, donate, support,\nworker, help, fund,\nrestaurantstrong, strong,\ncovid, donation",
            "ATTENTION TWITTER: The\npeople who serve us every day\nare depending on our support.\nRestaurants and bars, let us\nknow that you\u2019re\n#OpenForTakeout, so we can\nspread the word.\n\nYeah, we miss bars too. But\nwe can still support restaurant\nand bar workers today. Order\ntakeout and we\u2019ll buy your\nfirst round when they reopen."
        ],
        [
            "Social Media\nPromotions",
            "Virtual Bar Tours &\nLive Instagram Events",
            "live, instagram, tune,\ntonight, ig, go, dive, tour,\nbar, edition",
            "Going live now! doing a little\nthing with @budlight tonight\non IG and facebook live. come\nhang 6pm pst / 9pm est\n#divebartour"
        ],
        [
            "",
            "Pandemic Social\nEngagement &\nPromotions",
            "whassup, say, tag, win,\nmerch, friend, chance,\nneed, collection, week",
            "Win fresh Bud Whassup\ncrewnecks for your happy hour\ncrew! RETWEET + FOLLOW\nfor the chance to outfit your\ncrew."
        ],
        [
            "Sports",
            "Virtual NFL Draft\nParty with Celebrity\nHosts",
            "draft, boo, commish, nfl,\nyoutube, drafterparty,\nrobgronkowski, guest,\nhost, camillekostek",
            "TONIGHT!!!! @nfl @budlight"
        ],
        [
            "",
            "Home Sports\nChallenge Campaigns",
            "challenge, team, house,\nsport, league, submit,\nhome, athlete, chance,\nvideo",
            "Want to win a @trulyseltzer\nputting mat? Show us your\n#dailyninesetup. Photo, video,\nputting, chipping, whatever.\nShow us."
        ],
        [
            "Gaming",
            "Charity Gaming\nEvents by Alcohol\nBrands",
            "royale, twitch, tv, charity,\nmoney, tonight, dkm,\ntherealgavinlux, snellzilla,\nfinal",
            "Excited to announce the Bud\nLight Seltzer Charity Royale.\n24 pro athletes compete in\nCall of Duty: Warzone to raise\nmoney for awesome local\ncauses. Live right now on\nhttp://Twitch.tv/BudLight.\nTune in!"
        ],
    ]
    for r, row_data in enumerate(t2_data, start=1):
        for c, val in enumerate(row_data):
            add_table_cell(table2.rows[r].cells[c], val, font_size=Pt(8))

    doc.add_paragraph("")  # spacer

    doc.add_heading("Classification", level=2)
    doc.add_paragraph(
        "All 5,872 tweets were classified into the five themes using GPT-4o with structured "
        "prompting. Classification validity was established by comparing GPT-4o labels against "
        "human coding of a random subsample of 250 tweets, yielding agreement exceeding 85%. "
        "Tweets that did not clearly map to any theme were categorized as unclassified and "
        "excluded from prevalence analyses."
    )

    doc.add_heading("Prevalence Analysis", level=2)
    doc.add_paragraph(
        "Theme prevalence was calculated as the proportion of classified tweets assigned to "
        "each theme within each time period. Proportional z-tests were used to compare "
        "prevalence across the three periods (pre-pandemic, peri-lockdown, post-lockdown), "
        "with Bonferroni correction applied to account for multiple comparisons (5 themes \u00d7 "
        "3 comparisons = 15 tests; corrected \u03b1 = 0.0033). Analyses were conducted in Stata 17 "
        "(StataCorp, 2021)."
    )

    doc.add_heading("Sentiment Analysis", level=2)
    doc.add_paragraph(
        "To describe the communication mechanisms used in these themes during the peri-pandemic "
        "period as framed by the DSMM (Valkenburg & Peter, 2013), the entire collection of tweets "
        "for each theme was analyzed using LIWC (Chung et al., 2012; Boyd et al., 2022). LIWC "
        "functions as a linguistic magnifying glass, helping researchers understand the frequency "
        "of specific word types to gain insights into emotional states, social concerns, and "
        "thinking styles reflected in written content."
    )

    # ── Results ──
    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "The BERTopic pipeline identified an optimal 7-topic model (c_v coherence = 0.768) from "
        "1,200 candidate models, which was synthesized into five marketing themes (Table 2). "
        "Table 3 presents the prevalence of each theme across the three pandemic periods."
    )

    # ── Table 3: Prevalence ──
    p = doc.add_paragraph()
    run = p.add_run("Table 3. Theme Prevalence Across Periods")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    table3 = doc.add_table(rows=6, cols=7, style="Table Grid")
    table3.autofit = True
    headers = ["Theme", "Pre (%)", "Peri (%)", "Post (%)",
               "Pre\u2192Peri z (p)", "Peri\u2192Post z (p)", "Pre\u2192Post z (p)"]
    for i, h in enumerate(headers):
        add_table_cell(table3.rows[0].cells[i], h, bold=True, font_size=Pt(9))

    data = [
        ["AD-ID", "9.1", "17.4", "11.9",
         "6.03 (<0.001)*", "-4.13 (<0.001)*", "2.98 (0.0029)*"],
        ["Restaurant\nSupport", "1.9", "5.5", "1.7",
         "4.89 (<0.001)*", "-6.18 (<0.001)*", "-0.50 (0.62)"],
        ["Social Media\nPromotions", "6.6", "16.5", "10.6",
         "7.80 (<0.001)*", "-4.62 (<0.001)*", "4.59 (<0.001)*"],
        ["Sports", "8.2", "11.4", "8.9",
         "2.59 (0.01)", "-2.17 (0.03)", "0.83 (0.41)"],
        ["Gaming", "1.1", "2.0", "1.6",
         "1.80 (0.07)", "-0.79 (0.43)", "1.40 (0.16)"],
    ]
    for r, row_data in enumerate(data, start=1):
        for c, val in enumerate(row_data):
            add_table_cell(table3.rows[r].cells[c], val, font_size=Pt(9))

    p = doc.add_paragraph()
    run = p.add_run("Note. * = statistically significant at Bonferroni-corrected \u03b1 = 0.0033.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    doc.add_paragraph("")  # spacer

    doc.add_paragraph(
        "Three themes demonstrated statistically significant increases from the pre-pandemic "
        "to the peri-lockdown period. AD-ID prevalence rose from 9.1% to 17.4%, an increase "
        "of 91.2% (z = 6.03, p < 0.001). Restaurant Support increased from 1.9% to 5.5%, "
        "a 189.5% increase (z = 4.89, p < 0.001). Social Media Promotions rose from 6.6% to "
        "16.5%, a 150.0% increase (z = 7.80, p < 0.001). These increases were substantively "
        "large, reflecting a rapid reorientation of marketing emphasis during lockdowns."
    )
    doc.add_paragraph(
        "Critically, not all increases were transient. AD-ID prevalence remained significantly "
        "elevated in the post-lockdown period (11.9%) relative to pre-pandemic levels "
        "(z = 2.98, p = 0.0029), though it significantly declined from the peri-lockdown peak "
        "(z = -4.13, p < 0.001). Social Media Promotions similarly remained elevated "
        "post-lockdown (10.6%) relative to pre-pandemic levels (z = 4.59, p < 0.001), with a "
        "significant decline from the peri-lockdown peak (z = -4.62, p < 0.001). In contrast, "
        "Restaurant Support returned to near-baseline levels in the post-lockdown period (1.7%; "
        "peri\u2192post z = -6.18, p < 0.001), with no significant difference between pre- and "
        "post-pandemic levels (z = -0.50, p = 0.62), suggesting this was a temporary crisis "
        "response. Sports and Gaming themes showed no significant changes across any period "
        "at the Bonferroni-corrected threshold."
    )

    # ── Sentiment Analysis Results ──
    doc.add_heading("Sentiment Analysis", level=2)
    doc.add_paragraph(
        "Linguistic markers derived from LIWC provided a descriptive analysis of the five themes "
        "during the peri-pandemic response period in terms of their cognitive, emotional, and "
        "excitative characteristics."
    )
    doc.add_paragraph(
        "The AD-ID theme exhibited high analytical thinking (85.63%) with moderate clout "
        "(76.92%) and relatively low authenticity (36.78%), suggesting a guarded, strategic "
        "communication style. The notably higher tone score (80.16%) indicated predominantly "
        "positive emotional valence, with higher use of perceptual processes (9.64%) and spatial "
        "references (5.75%), and minimal negative emotions (0.22%) or conflict-related language "
        "(0%). The presence of food-related terms (4.89%) and physical references (5.84%) "
        "suggests content focused on sensory or experiential descriptions."
    )
    doc.add_paragraph(
        "The Restaurant Support theme presented the highest clout score (94.38%) and high "
        "analytical thinking (89.68%), coupled with the highest tone score (98.49%) across all "
        "themes. It showed relatively lower authenticity (9.14%), suggesting formal or carefully "
        "constructed content, with stronger social references (17.04%) and higher affiliation "
        "markers (7.77%). The complete absence of negation words (0%) and negative emotions "
        "(0%) reinforces its overwhelmingly positive orientation."
    )
    doc.add_paragraph(
        "The Social Media Promotions theme maintained relatively high analytical thinking "
        "(88.89%) and clout (93.8%), with moderate authenticity (32.37%). It showed the "
        "strongest perception-related content (12.04%) and motion references (2.73%), "
        "suggesting dynamic content, with high positive emotional tone (88.52%) and minimal "
        "negative emotions (0.20%)."
    )
    doc.add_paragraph(
        "The Sports theme showed the second-highest analytical thinking score (92.32%) with "
        "relatively higher clout (87.37%), but a lower tone score (67.22%). The Gaming theme "
        "exhibited the highest analytical thinking (98.37%) but lower clout (71.90%) and tone "
        "(51.75%), with the highest use of big words (26.99%) and numbers (10.02%), suggesting "
        "technical or specialized content."
    )
    doc.add_paragraph(
        "Through the DSMM lens (Valkenburg & Peter, 2013), the consistently high analytical "
        "scores (85.63%\u201398.37%) suggest content that elicits substantial cognitive engagement. "
        "The emotional processing dimension showed a gradient, with AD-ID, Restaurant Support, "
        "and Social Media Promotions displaying high positive emotional tone (80.16%\u201398.49%) "
        "and minimal negative affect\u2014suggesting emotionally engaging and accessible content\u2014"
        "while Sports and Gaming shifted toward more neutral, technically-oriented content "
        "(67.22% and 51.75%, respectively). Regarding excitative processes, Restaurant Support\u2019s "
        "high social and affiliation markers (17.04% and 7.77%) may facilitate stronger state "
        "excitation through social engagement, while Social Media Promotions\u2019 elevated perceptual "
        "(12.04%) and motion references (2.73%) suggest different levels of arousal potential."
    )

    # ── Discussion ──
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(
        "This study provides the first longitudinal, quantitative analysis of alcohol marketing "
        "theme prevalence on social media across pre-pandemic, lockdown, and post-lockdown "
        "periods. The findings support the hypothesis that alcohol brands rapidly amplified "
        "specific marketing themes during COVID-19 lockdowns and, importantly, demonstrate that "
        "some of these shifts persisted beyond the acute crisis period."
    )
    doc.add_paragraph(
        "The durability of the AD-ID theme is particularly concerning from a public health "
        "perspective. Alcohol delivery and isolation drinking messaging nearly doubled during "
        "lockdowns (from 9.1% to 17.4%) and remained significantly elevated at 11.9% even after "
        "restrictions eased. This pattern suggests that the pandemic may have accelerated a "
        "structural shift in how the alcohol industry markets its products\u2014normalizing home "
        "delivery and solitary consumption as routine rather than crisis-specific behaviors. This "
        "finding aligns with and extends qualitative observations by Martino et al. (2021) and "
        "Gerritsen et al. (2021), who identified similar themes but could not assess their "
        "trajectory over time. The persistence of these themes is especially concerning given the "
        "established association between alcohol marketing exposure and increased alcohol use "
        "(Anderson et al., 2009; Smith et al., 2009; Jernigan et al., 2017) and the concurrent "
        "increase in alcohol-related deaths during this period (White et al., 2022)."
    )
    doc.add_paragraph(
        "The transient nature of Restaurant Support messaging is consistent with a corporate "
        "social responsibility (CSR) interpretation: brands promoted restaurant partnerships "
        "during acute closures but reverted to other strategies once the immediate crisis passed. "
        "While supporting struggling restaurants was pro-social, CSR also serves alcohol "
        "companies via reputation management and direct advancement of commercial interests "
        "(Mialon et al., 2018; Yoon et al., 2013). By positioning themselves as champions of "
        "restaurant recovery while simultaneously promoting home delivery and isolation drinking, "
        "alcohol companies created a dual narrative that both legitimized their pandemic response "
        "through apparent social responsibility and expanded direct-to-consumer channels."
    )
    doc.add_paragraph(
        "Social Media Promotions, while declining from their lockdown peak, remained "
        "significantly elevated (from 6.6% pre-pandemic to 10.6% post-pandemic, a 60.6% "
        "increase), suggesting that brands learned the effectiveness of digital engagement "
        "strategies and continued to invest in them. The intersection of intensified alcohol "
        "marketing and heightened social media use is especially concerning given the established "
        "association between social media use and increased risks for mental health problems "
        "and problematic substance use behaviors."
    )
    doc.add_paragraph(
        "Methodologically, this study demonstrates the utility of the Embedding-to-Explanation "
        "pipeline (Thomas, 2025) for analyzing marketing content at scale. The combination of "
        "BERTopic with systematic hyperparameter optimization, LLM-based naming, and validated "
        "GPT-4o classification provides a reproducible framework that addresses key limitations "
        "of both manual content analysis and traditional topic models. The pipeline\u2019s open-source "
        "availability facilitates replication and extension to other marketing domains."
    )
    doc.add_paragraph(
        "Several limitations warrant acknowledgment. The sample included only 11 alcohol "
        "companies on a single platform (Twitter), limiting generalizability to broader industry "
        "practices or other social media environments. Twitter\u2019s demographics (typically younger, "
        "urban, and technologically adept) may influence the types of themes identified. The "
        "observational design cannot establish a causal link between marketing exposure and "
        "alcohol consumption behaviors. Classification, while validated against human labels, "
        "relied on LLM inference and may contain systematic biases. The sentiment analysis "
        "was exploratory and descriptive, providing aggregate rankings rather than statistical "
        "tests. Future research should link marketing theme exposure to behavioral outcomes "
        "using prospective designs and extend analyses to platforms such as Instagram and "
        "TikTok, where alcohol marketing is increasingly concentrated."
    )

    # ── References ──
    doc.add_heading("References", level=1)

    refs = [
        "Anderson, P., de Bruijn, A., Angus, K., Gordon, R., & Hastings, G. (2009). Impact of alcohol advertising and media exposure on adolescent alcohol use: A systematic review of longitudinal studies. Alcohol and Alcoholism, 44(3), 229\u2013243.",
        "Barry, A. E., Bates, A. M., Olusanya, O., Vinal, C. E., Martin, E., Peoples, J. E., ... & Montano, J. R. (2016). Alcohol marketing on Twitter and Instagram: Evidence of directly advertising to youth/adolescents. Alcohol and Alcoholism, 51(4), 487\u2013492.",
        "Boyd, R. L., Ashokkumar, A., Seraj, S., & Pennebaker, J. W. (2022). The Development and Psychometric Properties of LIWC-22. University of Texas at Austin.",
        "Chung, C. K., & Pennebaker, J. W. (2012). Linguistic inquiry and word count (LIWC): Pronounced \u201cLuke\u201d... and other useful facts. In P. M. McCarthy & C. Boonthum-Denecke (Eds.), Applied Natural Language Processing (pp. 206\u2013229). IGI Global.",
        "Colbert, S., Wilkinson, C., Thornton, L., & Richmond, R. (2020). COVID-19 and alcohol in Australia: Industry changes and public health impacts. Drug and Alcohol Review, 39(5), 435\u2013440.",
        "Gerritsen, S., Hasse, B., Jonsson, L., & Wall, M. (2021). Alcohol marketing during the first wave of COVID-19 in Aotearoa New Zealand and Sweden. BMC Public Health, 21, 1\u201312.",
        "Gillies, M., Murthy, D., Robinson, L., & McCracken, H. (2022). Themes and topics: An integrated framework for qualitative and computational approaches to text analysis. Journal of Communication, 72(5), 550\u2013570.",
        "Grootendorst, M. (2022). BERTopic: Neural topic modeling with a class-based TF-IDF procedure. arXiv preprint arXiv:2203.05794.",
        "Jernigan, D. H., & Ross, C. S. (2020). The alcohol marketing landscape: Alcohol industry size, structure, strategies, and public health responses. Journal of Studies on Alcohol and Drugs, Supplement 19, 13\u201325.",
        "Jernigan, D., Noel, J., Landon, J., Thornton, N., & Lobstein, T. (2017). Alcohol marketing and youth alcohol consumption: A systematic review of longitudinal studies published since 2008. Addiction, 112(S1), 7\u201320.",
        "Lacy-Nichols, J., Nandi, S., Gig\u00e0nte, B., Robinson, A., & McKee, M. (2023). The commercial determinants of health. The Lancet, 401(10383), 1229\u20131240.",
        "Martino, F., Brooks, R., Browne, J., Carah, N., Zorbas, C., Corben, K., ... & Backholer, K. (2021). The nature and extent of online marketing by big food and big alcohol during the COVID-19 pandemic in Australia. BMC Public Health, 21(1), 1\u201315.",
        "Mialon, M., & McCambridge, J. (2018). Alcohol industry corporate social responsibility initiatives and harmful drinking: A systematic review. European Journal of Public Health, 28(4), 664\u2013672.",
        "Pollard, M. S., Tucker, J. S., & Green, H. D. (2020). Changes in adult alcohol use and consequences during the COVID-19 pandemic in the US. JAMA Network Open, 3(9), e2022942.",
        "Qiang, J., Qian, Z., Li, Y., Yuan, Y., & Wu, X. (2020). Short text topic modeling techniques, applications, and performance: A survey. IEEE Transactions on Knowledge and Data Engineering, 34(3), 1427\u20131445.",
        "R\u00f6der, M., Both, A., & Hinneburg, A. (2015). Exploring the space of topic coherence measures. Proceedings of the Eighth ACM International Conference on Web Search and Data Mining, 399\u2013408.",
        "Smith, L. A., & Foxcroft, D. R. (2009). The effect of alcohol advertising, marketing and portrayal on drinking behaviour in young people: Systematic review of prospective cohort studies. BMC Public Health, 9, 51.",
        "StataCorp. (2021). Stata Statistical Software: Release 17. StataCorp LLC.",
        "Thomas, J. E. (2025). Embedding-to-Explanation: A BERTopic and Large Language Model Pipeline for Topic Discovery and Classification. Preprint.",
        "Valkenburg, P. M., & Peter, J. (2013). The Differential Susceptibility to Media Effects Model. Journal of Communication, 63(2), 221\u2013243.",
        "White, A. M., Castle, I. J., Powell, P. A., Hingson, R. W., & Koob, G. F. (2022). Alcohol-related deaths during the COVID-19 pandemic. JAMA, 327(17), 1704\u20131706.",
        "Wiederhold, B. K. (2020). Using social media to our advantage: Alleviating anxiety during a pandemic. Cyberpsychology, Behavior, and Social Networking, 23(4), 197\u2013198.",
        "Yoon, Y., G\u00fcrhan-Canli, Z., & Schwarz, N. (2006). The effect of corporate social responsibility (CSR) activities on companies with bad reputations. Journal of Consumer Psychology, 16(4), 377\u2013390.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

    doc.save(OUTPUT_PATH)
    print(f"Brief Report (Study 1) saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
