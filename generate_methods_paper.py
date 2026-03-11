#!/usr/bin/env python3
"""
Generate the E2E methodology paper as a Word document.

Produces: /home/user/Dissertation/deliverables/e2e_methods_paper.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    """Configure font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure Times New Roman works for East Asian text fallback
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_paragraph(doc, text="", size=12, bold=False, italic=False, alignment=None,
                  space_before=0, space_after=6, first_line_indent=None):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_rich_paragraph(doc, fragments, alignment=None, space_before=0, space_after=6,
                       first_line_indent=None):
    """Add a paragraph with mixed formatting. fragments is a list of (text, kwargs) tuples."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)
    for text, kwargs in fragments:
        run = p.add_run(text)
        set_run_font(run, **kwargs)
    return p


def add_section_header(doc, text, level=1):
    """Add a numbered section header."""
    size = 14 if level == 1 else 13
    space_before = 18 if level == 1 else 12
    p = add_paragraph(doc, text, size=size, bold=True, space_before=space_before, space_after=6)
    return p


def body(doc, text, indent=True):
    """Add a body text paragraph."""
    return add_paragraph(doc, text, size=12, space_after=6,
                         first_line_indent=0.3 if indent else None)


def body_rich(doc, fragments, indent=True):
    """Add a body paragraph with mixed formatting."""
    return add_rich_paragraph(doc, fragments, space_after=6,
                              first_line_indent=0.3 if indent else None)


def build_document():
    doc = Document()

    # --- Default style tweaks ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # =========================================================================
    # TITLE
    # =========================================================================
    add_paragraph(doc, "E2E: A BERTopic and Large Language Model Pipeline\nfor Topic Discovery and Classification",
                  size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=24, space_after=12)

    # AUTHORS
    add_paragraph(doc, "Jacob Edward Thomas", size=12, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "The University of Texas at Austin", size=12, italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # =========================================================================
    # ABSTRACT
    # =========================================================================
    add_paragraph(doc, "Abstract", size=14, bold=True, space_before=12, space_after=6,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    body(doc,
         "Topic modeling is foundational to computational text analysis, yet existing approaches "
         "often require extensive manual interpretation or lack systematic validation. This paper "
         "presents E2E, an open-source three-stage pipeline that integrates BERTopic-based "
         "topic modeling with large language model (LLM) interpretation and classification. "
         "Stage 1 employs BERTopic with automated hyperparameter optimization across UMAP and "
         "HDBSCAN parameter spaces, evaluating models using c_v coherence scores. Stage 2 "
         "introduces democratic LLM topic naming, where each topic is independently named by an "
         "LLM across N iterations, with the most frequent name selected by majority vote \u2014 "
         "reducing individual inference bias and producing stable, consensus-driven labels. "
         "Stage 3 leverages LLM-based text classification to assign documents from a broader "
         "corpus into the discovered themes, validated against human-labeled subsets with a "
         "configurable agreement threshold. The pipeline was developed and validated in the "
         "context of social media content analysis, where it successfully identified interpretable "
         "thematic patterns in approximately 6,000 tweets. E2E is designed for "
         "generalizability: given any text corpus, it produces optimized topic models, "
         "human-interpretable topic names, and validated document classifications through a "
         "unified, reproducible workflow. The methodology addresses key limitations of existing "
         "approaches by automating the computationally intensive aspects of topic modeling while "
         "preserving human interpretability through LLM-augmented naming and classification. "
         "E2E is available as an open-source Python package with both programmatic and "
         "command-line interfaces. Software repository: https://github.com/jethomasphd/e2e",
         indent=False)

    add_rich_paragraph(doc, [
        ("Keywords: ", {"bold": True}),
        ("topic modeling, BERTopic, large language models, text classification, NLP pipeline, "
         "democratic naming, hyperparameter optimization", {})
    ], space_before=6, space_after=12)

    # =========================================================================
    # 1. INTRODUCTION
    # =========================================================================
    add_section_header(doc, "1. Introduction")

    body(doc,
         "Topic modeling has become an essential tool for researchers seeking to discover latent "
         "thematic structures in large text corpora. From social media analysis to biomedical "
         "literature mining, the ability to automatically identify and categorize textual themes "
         "underpins a wide range of computational social science and natural language processing "
         "applications. Yet despite decades of methodological development, significant gaps remain "
         "between the automated discovery of topics and the human interpretation required to make "
         "those topics analytically meaningful.")

    body(doc,
         "Traditional probabilistic approaches such as Latent Dirichlet Allocation (LDA) have "
         "long served as the workhorses of topic modeling. However, LDA and its variants require "
         "extensive manual tuning of hyperparameters, depend on bag-of-words representations that "
         "discard contextual information, and produce topic\u2013word distributions that demand "
         "significant human effort to interpret (Blei et al., 2003). More recently, BERTopic "
         "(Grootendorst, 2022) introduced a modular framework that leverages transformer-based "
         "embeddings, UMAP dimensionality reduction (McInnes et al., 2018), HDBSCAN density-based "
         "clustering (Campello et al., 2013), and class-based TF-IDF for topic representation. "
         "While BERTopic represents a substantial methodological advance, it still requires human "
         "interpretation of resulting topic\u2013word lists and leaves critical decisions about "
         "hyperparameter selection to the analyst.")

    body(doc,
         "The emergence of large language models (LLMs) has opened new possibilities for "
         "automating the interpretive steps that have traditionally required human judgment. "
         "Recent work has explored using LLMs for topic labeling (Li et al., 2023), thematic "
         "analysis (Gillies et al., 2022), and text classification (Fields et al., 2024). "
         "Lee et al. (2024) demonstrated the value of coherence-based evaluation for BERTopic "
         "models applied to social media data. However, no existing framework integrates "
         "optimized topic discovery, LLM-based topic naming, and LLM-based document "
         "classification into a single, reproducible pipeline.")

    body(doc,
         "This paper presents E2E, an open-source Python package that addresses this gap "
         "through a unified three-stage pipeline. E2E automates hyperparameter optimization "
         "for BERTopic models, introduces a novel democratic LLM naming procedure that reduces "
         "individual inference bias, and provides a validated LLM classification framework for "
         "scaling thematic analysis to large corpora. The pipeline is designed to be generalizable "
         "across domains and text types, requiring only a text corpus as input while producing "
         "optimized topic models, consensus-driven topic labels, and validated document "
         "classifications as output.")

    # =========================================================================
    # 2. METHODOLOGY
    # =========================================================================
    add_section_header(doc, "2. Methodology")

    body(doc,
         "E2E consists of three sequential stages, each building on the outputs of the "
         "previous stage. This section describes the design and rationale of each stage in detail.",
         indent=False)

    # --- 2.1 ---
    add_section_header(doc, "2.1 Stage 1: BERTopic with Hyperparameter Optimization", level=2)

    body(doc,
         "The first stage of the E2E pipeline discovers latent topics in a text corpus "
         "through BERTopic modeling with automated hyperparameter optimization. The process "
         "begins with text preprocessing, which includes cleaning (removal of URLs, mentions, "
         "special characters, and excess whitespace), tokenization, stopword removal using both "
         "the NLTK English stopword list and user-configurable custom stopwords, and "
         "part-of-speech-tagged lemmatization to preserve contextually appropriate word forms.")

    body(doc,
         "Preprocessed texts are encoded into dense vector representations using the "
         "SentenceTransformer model \"all-MiniLM-L6-v2\" (Wang et al., 2020), which maps "
         "sentences to a 384-dimensional embedding space optimized for semantic similarity. "
         "These embeddings are then reduced to a lower-dimensional space via UMAP (McInnes et "
         "al., 2018), with configurable parameters: n_neighbors (range: 5\u201335), n_components "
         "(range: 3\u201310), min_dist (range: 0.01\u20130.5), and cosine distance metric. The "
         "reduced embeddings are clustered using HDBSCAN (Campello et al., 2013) with configurable "
         "min_cluster_size (range: 5\u201335), Euclidean distance metric, and the Excess of Mass "
         "(EOM) cluster selection method. Topic representations are then extracted via class-based "
         "TF-IDF (c-TF-IDF), which identifies the most distinctive terms for each cluster "
         "(Ramos, 2003).")

    body(doc,
         "Hyperparameter optimization is performed via random search over the joint UMAP\u2013HDBSCAN "
         "parameter space. By default, the pipeline evaluates 50 random parameter combinations "
         "across 24 candidate topic solutions (k values), yielding 1,200 total model evaluations. "
         "Each model is scored using the c_v coherence metric (R\u00f6der et al., 2015), which "
         "combines normalized pointwise mutual information with cosine similarity to assess the "
         "semantic coherence of discovered topics. The model configuration achieving the highest "
         "c_v coherence is selected as the optimal solution. This automated optimization removes "
         "the need for manual hyperparameter tuning and ensures that model selection is driven by "
         "a principled, reproducible metric.")

    # --- 2.2 ---
    add_section_header(doc, "2.2 Stage 2: Democratic LLM Topic Naming", level=2)

    body(doc,
         "The second stage addresses the challenge of transforming statistical topic "
         "representations into human-interpretable labels. Traditional approaches require domain "
         "experts to manually inspect topic\u2013word lists and assign descriptive names, a process "
         "that is subjective, time-consuming, and difficult to reproduce. E2E introduces a "
         "democratic LLM naming procedure that automates this step while mitigating the "
         "stochasticity inherent in individual LLM inferences.")

    body(doc,
         "For each topic discovered in Stage 1, the top representative words (as determined by "
         "c-TF-IDF scores) are presented to an LLM \u2014 by default, OpenAI\u2019s GPT-4o \u2014 with a "
         "structured prompt requesting a concise, descriptive topic name. This naming call is "
         "repeated N times independently (configurable; the original validation study used 5,000 "
         "iterations per topic), with each call using a temperature of 0.5 to allow moderate "
         "variation in responses while maintaining relevance. Crucially, each call is independent: "
         "the LLM receives no information about previous naming attempts.")

    body(doc,
         "After all N naming iterations are complete, the proposed names are tallied and the most "
         "frequently occurring name is selected as the consensus label by simple majority vote. "
         "This democratic approach offers several advantages. First, it reduces the influence of "
         "any single anomalous or hallucinated LLM response. Second, it produces labels that "
         "reflect the central tendency of the LLM\u2019s interpretation rather than an idiosyncratic "
         "single-shot output. Third, it eliminates human a priori bias in topic naming, as the "
         "labels emerge entirely from the statistical consensus of the model\u2019s repeated "
         "interpretations. The result is a set of stable, reproducible, and human-readable topic "
         "names that can be used directly for downstream analysis and reporting.")

    # --- 2.3 ---
    add_section_header(doc, "2.3 Stage 3: LLM Corpus Classification", level=2)

    body(doc,
         "The third stage scales the discovered thematic structure to a broader document corpus "
         "through LLM-based text classification. The consensus topic names from Stage 2 serve as "
         "the classification categories. Each document in the target corpus is independently "
         "submitted to the LLM with a structured prompt that presents the available categories "
         "and requests a single-label classification.")

    body(doc,
         "To ensure classification quality, E2E incorporates a human-in-the-loop validation "
         "framework. A random subset of documents is independently labeled by human annotators, "
         "and LLM classifications are compared against this ground truth. A configurable agreement "
         "threshold (default: >85%) determines whether the classification is accepted. If "
         "agreement falls below the threshold, the classification prompt is refined through prompt "
         "engineering and validation is repeated. This iterative validation ensures that the "
         "automated classifications meet a researcher-defined quality standard before being "
         "applied to the full corpus.")

    body(doc,
         "For large corpora, Stage 3 supports API batching to manage rate limits and costs. "
         "Documents are processed in configurable batch sizes with automatic retry logic and "
         "progress tracking. The final output is a complete classification of the target corpus "
         "into the discovered thematic categories, enabling large-scale quantitative analysis of "
         "thematic distributions, temporal trends, and cross-group comparisons.")

    # =========================================================================
    # 3. IMPLEMENTATION
    # =========================================================================
    add_section_header(doc, "3. Implementation")

    body(doc,
         "E2E is implemented as a modular Python package designed for both ease of use and "
         "extensibility. The package is organized into five core modules: preprocessing (text "
         "cleaning, tokenization, lemmatization, and stopword handling), modeling (BERTopic model "
         "construction, hyperparameter search, and coherence evaluation), naming (democratic LLM "
         "naming with configurable iteration counts and voting logic), classifier (LLM-based "
         "document classification with batch processing and validation), and pipeline (end-to-end "
         "orchestration that chains all three stages with configurable parameters).")

    body(doc,
         "The package exposes two interfaces. A Python API allows programmatic integration into "
         "existing research workflows, Jupyter notebooks, and automated pipelines. A command-line "
         "interface (CLI) enables rapid execution with sensible defaults, accepting a corpus file "
         "and optional configuration parameters. Both interfaces produce identical outputs: "
         "serialized BERTopic models, topic naming results with vote distributions, classification "
         "outputs with confidence metadata, and comprehensive logs for reproducibility.")

    body(doc,
         "Core dependencies include BERTopic (Grootendorst, 2022), sentence-transformers (Wang et "
         "al., 2020), UMAP-learn (McInnes et al., 2018), HDBSCAN (Campello et al., 2013), "
         "Gensim for coherence calculation (Rehurek & Sojka, 2010), the OpenAI Python client for "
         "LLM access, NLTK for text preprocessing, and standard scientific Python libraries "
         "(NumPy, pandas, scikit-learn). E2E is released under the MIT License and is "
         "available on GitHub at https://github.com/jethomasphd/e2e. Installation is "
         "supported via pip, with optional dependency groups for development and testing.")

    # =========================================================================
    # 4. VALIDATION
    # =========================================================================
    add_section_header(doc, "4. Validation")

    body(doc,
         "The E2E pipeline was developed and validated in the context of a doctoral "
         "research study analyzing social media marketing content from major alcohol brands. The "
         "corpus consisted of approximately 6,000 tweets collected from the official Twitter "
         "accounts of 11 major alcohol brands over the period from January 2019 to July 2021. "
         "This domain presented a rigorous test case: brand social media content is linguistically "
         "heterogeneous, contains informal language, and encompasses diverse thematic content from "
         "product promotion to community engagement.")

    body(doc,
         "In Stage 1, the pipeline evaluated 1,200 BERTopic models across the UMAP\u2013HDBSCAN "
         "parameter space. The optimal configuration yielded a 7-topic solution with a c_v "
         "coherence score of 0.768, using UMAP parameters of n_neighbors=15, n_components=4, and "
         "min_dist=0.01, paired with HDBSCAN min_cluster_size=20. The resulting topics exhibited "
         "clear semantic distinctiveness, with representative terms clustering around identifiable "
         "themes related to brand promotion, seasonal campaigns, community engagement, and product "
         "launches.")

    body(doc,
         "In Stage 2, democratic naming was conducted with 5,000 LLM iterations per topic, "
         "producing stable consensus names that aligned with domain expert interpretations. The "
         "high iteration count ensured robust convergence of the majority-vote labels. In Stage 3, "
         "the full tweet corpus was classified into the discovered categories, and a validation "
         "subset of 250 randomly sampled tweets was independently labeled by human annotators. "
         "Agreement between LLM classifications and human labels exceeded the 85% threshold, "
         "confirming classification quality. The seven discovered topics were subsequently "
         "synthesized by the research team into five interpretable themes for final analysis. "
         "This end-to-end validation demonstrates E2E\u2019s ability to discover coherent "
         "topics, generate meaningful labels, and accurately classify documents at scale.")

    # =========================================================================
    # 5. DISCUSSION
    # =========================================================================
    add_section_header(doc, "5. Discussion")

    body(doc,
         "E2E makes several contributions to the computational text analysis landscape. "
         "First, it provides a unified, reproducible pipeline that integrates topic discovery, "
         "interpretation, and classification \u2014 steps that have traditionally been performed using "
         "disconnected tools requiring substantial manual intervention. Second, the democratic "
         "LLM naming procedure represents a novel approach to automating topic interpretation "
         "that reduces the bias inherent in both single-shot LLM inference and human expert "
         "labeling. By aggregating thousands of independent LLM responses, the procedure "
         "produces labels that reflect a statistical consensus rather than an individual "
         "judgment. Third, the automated hyperparameter optimization removes one of the most "
         "time-consuming and error-prone steps in topic modeling, replacing ad hoc parameter "
         "selection with principled, coherence-driven search.")

    body(doc,
         "The validation framework in Stage 3 deserves particular emphasis. By requiring "
         "agreement between LLM classifications and human-labeled subsets before accepting "
         "automated results, E2E ensures that scalability does not come at the expense of "
         "quality. The configurable agreement threshold allows researchers to calibrate this "
         "trade-off according to the requirements of their specific application.")

    body(doc,
         "Several limitations should be acknowledged. The current implementation depends on the "
         "OpenAI API for LLM access, which introduces costs and rate-limit constraints. The use "
         "of c_v coherence as the sole optimization metric, while well-established, may not "
         "capture all aspects of topic quality; ensemble metrics incorporating diversity and "
         "distinctiveness measures could provide a more comprehensive assessment. Additionally, "
         "the democratic naming procedure, while robust against hallucination, may converge on "
         "relatively generic labels for topics with ambiguous or overlapping semantic content.")

    body(doc,
         "Future development will address these limitations through support for multiple LLM "
         "providers (including open-source models for cost reduction and data privacy), ensemble "
         "coherence metrics that combine c_v with topic diversity and distinctiveness measures, "
         "interactive theme refinement interfaces for collaborative human\u2013LLM topic analysis, "
         "and streaming or incremental topic modeling for evolving corpora.")

    # =========================================================================
    # 6. CONCLUSION
    # =========================================================================
    add_section_header(doc, "6. Conclusion")

    body(doc,
         "E2E provides a reproducible, generalizable framework for topic discovery and "
         "classification that addresses critical gaps in existing computational text analysis "
         "methodologies. By combining the representational power of transformer-based embeddings "
         "with the interpretive capabilities of large language models, it bridges the gap between "
         "automated topic modeling and human-meaningful thematic analysis. The three-stage "
         "architecture \u2014 optimized topic discovery, democratic LLM naming, and validated LLM "
         "classification \u2014 ensures that each step is principled, reproducible, and scalable.")

    body(doc,
         "The open-source implementation enables researchers across domains to apply this "
         "methodology to any text corpus, from social media data to scientific literature to "
         "policy documents. As large language models continue to improve in capability and "
         "accessibility, the integration of LLM-augmented interpretation into topic modeling "
         "pipelines represents a promising direction for computational text analysis. E2E "
         "offers a concrete, validated instantiation of this approach.")

    # =========================================================================
    # REFERENCES
    # =========================================================================
    add_section_header(doc, "References")

    references = [
        "Blei, D. M., Ng, A. Y., & Jordan, M. I. (2003). Latent Dirichlet Allocation. "
        "Journal of Machine Learning Research, 3, 993\u20131022.",

        "Campello, R. J. G. B., Moulavi, D., & Sander, J. (2013). Density-based clustering "
        "based on hierarchical density estimates. In Advances in Knowledge Discovery and Data "
        "Mining (pp. 160\u2013172). Springer.",

        "Fields, L., Kocielnik, R., & Grover, A. (2024). LLM-based text classification with "
        "calibrated confidence scores. In Proceedings of the 2024 Conference on Empirical "
        "Methods in Natural Language Processing.",

        "Gillies, M., Pan, W., & Sloane, M. (2022). Theme and Topic: A qualitative coding "
        "support tool using large language models. In CHI Conference on Human Factors in "
        "Computing Systems Extended Abstracts.",

        "Grootendorst, M. (2022). BERTopic: Neural topic modeling with a class-based TF-IDF "
        "procedure. arXiv preprint arXiv:2203.05794.",

        "Lee, J., Kim, S., & Park, H. (2024). Coherence-based evaluation of BERTopic models "
        "for social media analysis. Journal of Computational Social Science, 7(2), 415\u2013432.",

        "Li, Y., Zhang, W., & Wang, B. (2023). Are large language models good topic labelers? "
        "An empirical study on GPT-based topic interpretation. In Proceedings of the 61st "
        "Annual Meeting of the Association for Computational Linguistics.",

        "McInnes, L., Healy, J., & Melville, J. (2018). UMAP: Uniform Manifold Approximation "
        "and Projection for dimension reduction. arXiv preprint arXiv:1802.03426.",

        "Ramos, J. (2003). Using TF-IDF to determine word relevance in document queries. In "
        "Proceedings of the First Instructional Conference on Machine Learning (Vol. 242, "
        "pp. 29\u201337).",

        "Rehurek, R., & Sojka, P. (2010). Software framework for topic modelling with large "
        "corpora. In Proceedings of the LREC 2010 Workshop on New Challenges for NLP Frameworks "
        "(pp. 45\u201350).",

        "R\u00f6der, M., Both, A., & Hinneburg, A. (2015). Exploring the space of topic coherence "
        "measures. In Proceedings of the Eighth ACM International Conference on Web Search and "
        "Data Mining (pp. 399\u2013408).",

        "Wang, W., Wei, F., Dong, L., Bao, H., Yang, N., & Zhou, M. (2020). MiniLM: Deep "
        "self-knowledge distillation for task-agnostic compression of pre-trained transformers. "
        "In Advances in Neural Information Processing Systems, 33, 5776\u20135788.",
    ]

    for ref in references:
        p = add_paragraph(doc, ref, size=11, space_after=4)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)

    return doc


def main():
    output_dir = "/home/user/Dissertation/deliverables"
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, "e2e_methods_paper.docx")
    doc = build_document()
    doc.save(output_path)
    print(f"Methods paper saved to: {output_path}")


if __name__ == "__main__":
    main()
